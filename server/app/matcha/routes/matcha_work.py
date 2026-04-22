"""Matcha Work — chat-driven AI workspace for HR document elements."""

import asyncio
import json
import logging
import math
import os
import urllib.parse

import httpx
import re
from datetime import datetime, timezone
from typing import Optional
from uuid import UUID

from fastapi import APIRouter, Body, Depends, HTTPException, Query, Response, UploadFile, File, status
from fastapi.responses import StreamingResponse

from ...core.models.auth import CurrentUser
from ...core.services.compliance_service import get_location_requirements, get_locations
from ...core.services.storage import get_storage
from ...database import get_connection
from ..dependencies import require_admin_or_client, get_client_company_id, require_feature
from ..services.escalation_service import should_escalate, create_escalation
from ..services.model_pricing import calculate_call_cost
from ..models.matcha_work import (
    CreateThreadRequest,
    CreateThreadResponse,
    DocumentVersionResponse,
    ElementListItem,
    FinalizeResponse,
    HandbookDocument,
    MWMessageOut,
    OnboardingDocument,
    PolicyDocument,
    PresentationDocument,
    ReviewDocument,
    RevertRequest,
    SaveDraftResponse,
    SendMessageRequest,
    SendMessageResponse,
    PinThreadRequest,
    NodeModeRequest,
    ComplianceModeRequest,
    PayerModeRequest,
    ThreadDetailResponse,
    ThreadListItem,
    OfferLetterDocument,
    UsageSummaryResponse,
    UpdateTitleRequest,
    SendReviewRequestsRequest,
    SendReviewRequestsResponse,
    ReviewRequestStatus,
    PublicReviewRequestResponse,
    PublicReviewSubmitRequest,
    PublicReviewSubmitResponse,
    SendHandbookSignaturesRequest,
    SendHandbookSignaturesResponse,
    GeneratePresentationResponse,
    ResumeBatchDocument,
    InventoryDocument,
    ProjectDocument,
    SendInterviewsRequest,
    RejectCandidateRequest,
    WorkbookDocument,
)
from ..services import matcha_work_document as doc_svc
from ..services import billing_service
from ..services import token_budget_service
from ..services.er_document_parser import ERDocumentParser
from ..services.matcha_work_handbook_upload import (
    AuditedLocation,
    MAX_RED_FLAGS,
    MAX_SECTION_PREVIEWS,
    _audit_location_group,
    _severity_rank,
    audit_uploaded_handbook,
    check_handbook_relevance,
    compute_coverage_summaries,
    derive_handbook_title,
    parse_handbook_sections,
)
from ..services.matcha_work_ai import get_ai_provider, _infer_skill_from_state, _build_company_context, compact_conversation, needs_live_web_context, fetch_live_web_context
from ..services.matcha_work_node import build_node_context, build_compliance_context, ComplianceContextResult
from ..services.onboarding_orchestrator import (
    PROVIDER_GOOGLE_WORKSPACE,
    PROVIDER_SLACK,
    start_google_workspace_onboarding,
    start_slack_onboarding,
)
from ...core.services.email import get_email_service
from ...core.services.handbook_service import HandbookService
from ...config import get_settings

logger = logging.getLogger(__name__)


async def _get_rag_context(content: str, company_id, max_tokens: int = 4000) -> str | None:
    """Fetch compliance RAG context for a user question. Returns None on failure."""
    try:
        from ...core.services.embedding_service import EmbeddingService
        from ...core.services.compliance_rag import ComplianceRAGService

        api_key = os.getenv("GEMINI_API_KEY") or get_settings().gemini_api_key
        if not api_key or not content:
            return None
        es = EmbeddingService(api_key=api_key)
        crag = ComplianceRAGService(es)
        async with get_connection() as conn:
            ctx, _ = await crag.get_context_for_question(
                query=content, conn=conn,
                company_id=company_id, max_tokens=max_tokens,
            )
        return ctx or None
    except Exception as e:
        logger.warning("RAG augmentation failed: %s", e)
        return None

router = APIRouter(dependencies=[Depends(require_feature("matcha_work"))])
public_router = APIRouter()
presence_router = APIRouter()

# ── Name resolution SQL fragment (shared with inbox) ──
_PRESENCE_NAME_EXPR = "COALESCE(c.name, CONCAT(e.first_name, ' ', e.last_name), a.name, u.email)"


@presence_router.post("/heartbeat", status_code=204)
async def presence_heartbeat(current_user: CurrentUser = Depends(require_admin_or_client)):
    """Update the user's Matcha Work last-active timestamp."""
    async with get_connection() as conn:
        await conn.execute(
            "UPDATE users SET mw_last_active = NOW() WHERE id = $1",
            current_user.id,
        )


@presence_router.get("/online")
async def get_online_users(current_user: CurrentUser = Depends(require_admin_or_client)):
    """Return users active on Matcha Work within the last 2 minutes."""
    async with get_connection() as conn:
        rows = await conn.fetch(
            f"""
            SELECT u.id, u.email, u.avatar_url, u.mw_last_active,
                   {_PRESENCE_NAME_EXPR} AS name
            FROM users u
            LEFT JOIN clients c ON c.user_id = u.id
            LEFT JOIN employees e ON e.user_id = u.id
            LEFT JOIN admins a ON a.user_id = u.id
            WHERE u.mw_last_active > NOW() - INTERVAL '2 minutes'
              AND u.id != $1
              AND u.is_active = true
            ORDER BY u.mw_last_active DESC
            """,
            current_user.id,
        )
    return [
        {
            "id": str(r["id"]),
            "name": r["name"] or r["email"],
            "email": r["email"],
            "avatar_url": r["avatar_url"],
            "last_active": r["mw_last_active"].isoformat() if r["mw_last_active"] else None,
        }
        for r in rows
    ]

VALID_OFFER_LETTER_FIELDS = set(OfferLetterDocument.model_fields.keys()) - {"company_logo_url"}
VALID_REVIEW_FIELDS = set(ReviewDocument.model_fields.keys())
VALID_WORKBOOK_FIELDS = set(WorkbookDocument.model_fields.keys()) - {"images"}
VALID_ONBOARDING_FIELDS = set(OnboardingDocument.model_fields.keys())
VALID_PRESENTATION_FIELDS = set(PresentationDocument.model_fields.keys()) - {"cover_image_url"}
HANDBOOK_UPLOAD_MANAGED_FIELDS = {
    "handbook_source_type",
    "handbook_upload_status",
    "handbook_uploaded_file_url",
    "handbook_uploaded_filename",
    "handbook_blocking_error",
    "handbook_review_locations",
    "handbook_red_flags",
    "handbook_green_flags",
    "handbook_jurisdiction_summaries",
    "handbook_analysis_generated_at",
    "handbook_strength_score",
    "handbook_strength_label",
    "handbook_analysis_progress",
    "handbook_total_red_flag_count",
}
VALID_HANDBOOK_FIELDS = set(HandbookDocument.model_fields.keys()) - HANDBOOK_UPLOAD_MANAGED_FIELDS
VALID_POLICY_FIELDS = set(PolicyDocument.model_fields.keys())
VALID_RESUME_BATCH_FIELDS = set(ResumeBatchDocument.model_fields.keys())
VALID_INVENTORY_FIELDS = set(InventoryDocument.model_fields.keys())
VALID_PROJECT_FIELDS = set(ProjectDocument.model_fields.keys())

EMAIL_REGEX = re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")
HANDBOOK_UPLOAD_EXTENSIONS = {".pdf", ".doc", ".docx"}
HANDBOOK_UPLOAD_MAX_BYTES = 10 * 1024 * 1024
RESUME_UPLOAD_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}
RESUME_UPLOAD_MAX_BYTES = 10 * 1024 * 1024
RESUME_TEXT_CAP = 15_000
INVENTORY_UPLOAD_EXTENSIONS = {".pdf", ".csv", ".xlsx", ".xls", ".doc", ".docx", ".txt"}
INVENTORY_UPLOAD_MAX_BYTES = 15 * 1024 * 1024
INVENTORY_TEXT_CAP = 15_000


def _validate_updates_for_skill(skill: str, updates: dict) -> dict:
    """Filter AI updates to known fields for the inferred skill."""
    if skill == "offer_letter":
        valid_fields = VALID_OFFER_LETTER_FIELDS
    elif skill == "review":
        valid_fields = VALID_REVIEW_FIELDS
    elif skill == "onboarding":
        valid_fields = VALID_ONBOARDING_FIELDS
    elif skill == "workbook":
        valid_fields = VALID_WORKBOOK_FIELDS
    elif skill == "presentation":
        valid_fields = VALID_PRESENTATION_FIELDS
    elif skill == "handbook":
        valid_fields = VALID_HANDBOOK_FIELDS
    elif skill == "policy":
        valid_fields = VALID_POLICY_FIELDS
    elif skill == "resume_batch":
        valid_fields = VALID_RESUME_BATCH_FIELDS
    elif skill == "inventory":
        valid_fields = VALID_INVENTORY_FIELDS
    elif skill == "project":
        valid_fields = VALID_PROJECT_FIELDS
    elif skill == "blog":
        from ..services.matcha_work_ai import BLOG_FIELDS as _BLOG_FIELDS
        valid_fields = set(_BLOG_FIELDS)
    else:
        return {}
    return {k: v for k, v in updates.items() if k in valid_fields}


def _row_to_message(row: dict) -> MWMessageOut:
    raw_meta = row.get("metadata")
    if isinstance(raw_meta, str):
        try:
            raw_meta = json.loads(raw_meta)
        except (json.JSONDecodeError, TypeError):
            raw_meta = None
    return MWMessageOut(
        id=row["id"],
        thread_id=row["thread_id"],
        role=row["role"],
        content=row["content"],
        version_created=row.get("version_created"),
        metadata=raw_meta,
        created_at=row["created_at"],
    )


def _location_label(location: dict) -> str:
    city = str(location.get("city") or "").strip()
    state = str(location.get("state") or "").strip().upper()
    return f"{city}, {state}" if city else state


def _thread_accepts_handbook_upload(thread: dict) -> bool:
    current_state = thread.get("current_state") or {}
    current_skill = _infer_skill_from_state(current_state)
    if current_skill == "chat":
        return True
    if current_skill != "handbook":
        return False
    # Reject if an analysis is already in progress.
    if current_state.get("handbook_upload_status") == "analyzing":
        return False
    source_type = current_state.get("handbook_source_type")
    # Allow upload if already in upload mode OR if source type hasn't been
    # committed yet (user started chatting about a handbook but can still
    # switch to upload mode via the paperclip button).
    if source_type in ("upload", None):
        return True
    return False


def _build_handbook_block_message(location_labels: list[str]) -> str:
    if not location_labels:
        return (
            "Handbook upload audit is blocked because no active Compliance Locations were found. "
            "Add or sync company locations in /compliance first."
        )
    if len(location_labels) == 1:
        scoped = location_labels[0]
    else:
        scoped = ", ".join(location_labels[:6])
        if len(location_labels) > 6:
            scoped += f", and {len(location_labels) - 6} more"
    return (
        "Handbook upload audit is blocked because these active Compliance Locations are not fully synced: "
        f"{scoped}. Fix /compliance coverage first, then retry the upload."
    )


def _build_handbook_upload_summary(
    *,
    file_name: str,
    reviewed_locations: list[str],
    red_flags: list[dict],
    green_flags: list[dict] | None = None,
    jurisdiction_summaries: list[dict] | None = None,
    blocked_message: Optional[str] = None,
) -> str:
    if blocked_message:
        return blocked_message

    passing_count = len(green_flags or [])
    gap_count = len(red_flags)

    if not red_flags:
        return (
            f"Uploaded {file_name} and reviewed it against {len(reviewed_locations)} active Compliance Location(s). "
            f"{passing_count} requirement(s) covered, no jurisdiction coverage gaps detected."
        )

    counts = {"high": 0, "medium": 0, "low": 0}
    for row in red_flags:
        severity = str(row.get("severity") or "medium").lower()
        if severity in counts:
            counts[severity] += 1

    severity_bits = []
    for severity in ("high", "medium", "low"):
        count = counts[severity]
        if count:
            severity_bits.append(f"{count} {severity}")

    severity_summary = ", ".join(severity_bits) if severity_bits else f"{gap_count} issue(s)"

    # Per-jurisdiction coverage snippet
    jurisdiction_bits: list[str] = []
    for js in (jurisdiction_summaries or []):
        jurisdiction_bits.append(f"{js['location_label']} {js['covered_count']}/{js['total_count']}")
    jurisdiction_snippet = (" | ".join(jurisdiction_bits) + ".") if jurisdiction_bits else ""

    return (
        f"Uploaded {file_name} and reviewed it against {len(reviewed_locations)} active Compliance Location(s). "
        f"{passing_count} passing, {severity_summary} red flag(s). "
        + (f"Coverage: {jurisdiction_snippet} " if jurisdiction_snippet else "")
        + "Review the Preview panel for details."
    )


async def _build_thread_detail_response(thread_id: UUID, company_id: Optional[UUID], *, user_id: UUID | None = None) -> ThreadDetailResponse:
    thread = await doc_svc.get_thread(thread_id, company_id, user_id=user_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    # Use the thread's actual company_id — callers who are collaborators on a
    # thread from another company would otherwise pass their own company_id.
    thread_company_id = thread["company_id"]

    thread["current_state"] = await doc_svc.ensure_matcha_work_thread_storage_scope(
        thread_id,
        thread_company_id,
        thread["current_state"],
    )
    messages = await doc_svc.get_thread_messages(thread_id)

    # Fetch collaborators
    collaborators = []
    async with get_connection() as conn:
        collab_rows = await conn.fetch(
            """
            SELECT tc.user_id, tc.role, tc.created_at,
                   u.email, u.avatar_url,
                   COALESCE(cl.name, CONCAT(e.first_name, ' ', e.last_name), a.name, u.email) AS name
            FROM mw_thread_collaborators tc
            JOIN users u ON u.id = tc.user_id
            LEFT JOIN clients cl ON cl.user_id = tc.user_id
            LEFT JOIN employees e ON e.user_id = tc.user_id
            LEFT JOIN admins a ON a.user_id = tc.user_id
            WHERE tc.thread_id = $1
            ORDER BY tc.created_at
            """,
            thread_id,
        )
        collaborators = [
            {
                "user_id": str(r["user_id"]),
                "name": r["name"],
                "email": r["email"],
                "role": r["role"],
                "avatar_url": r["avatar_url"],
                "added_at": r["created_at"].isoformat() if r["created_at"] else None,
            }
            for r in collab_rows
        ]

    return ThreadDetailResponse(
        id=thread["id"],
        title=thread["title"],
        status=thread["status"],
        current_state=thread["current_state"],
        version=thread["version"],
        task_type=_infer_skill_from_state(thread["current_state"]),
        is_pinned=thread.get("is_pinned", False),
        node_mode=thread.get("node_mode", False),
        compliance_mode=thread.get("compliance_mode", False),
        linked_offer_letter_id=thread.get("linked_offer_letter_id"),
        created_at=thread["created_at"],
        updated_at=thread["updated_at"],
        messages=[_row_to_message(row) for row in messages],
        collaborators=collaborators,
    )


async def _get_affected_employees(
    company_id: UUID,
    metadata: dict,
) -> list[dict] | None:
    """Count employees affected per referenced compliance location.

    Cross-references Gemini's referenced_locations with the compliance
    reasoning chains to find matching business_location IDs, then counts
    employees at those locations (exact match via work_location_id, with
    work_state fallback for employees without a linked location).
    """
    referenced = metadata.get("referenced_locations", [])
    chains = metadata.get("compliance_reasoning", [])
    if not referenced or not chains:
        return None

    label_to_id: dict[str, str] = {c["location_label"]: c["location_id"] for c in chains}

    # Gemini may abbreviate labels — fuzzy match
    loc_ids: list[UUID] = []
    for ref in referenced:
        for label, lid in label_to_id.items():
            if ref == label or ref in label or label.startswith(ref):
                loc_ids.append(UUID(lid))
                break

    if not loc_ids:
        return None

    async with get_connection() as conn:
        rows = await conn.fetch(
            """
            SELECT bl.id as loc_id, bl.name, bl.city, bl.state, COUNT(e.id) as count
            FROM employees e
            JOIN business_locations bl ON bl.id = e.work_location_id
            WHERE e.org_id = $1 AND e.termination_date IS NULL
              AND bl.id = ANY($2::uuid[])
            GROUP BY bl.id, bl.name, bl.city, bl.state
            """,
            company_id, loc_ids,
        )

        matched_loc_ids = {r["loc_id"] for r in rows}
        unmatched = [lid for lid in loc_ids if lid not in matched_loc_ids]

        state_rows: list = []
        if unmatched:
            loc_states = await conn.fetch(
                "SELECT id, state FROM business_locations WHERE id = ANY($1::uuid[])",
                unmatched,
            )
            states = [r["state"] for r in loc_states if r["state"]]
            if states:
                state_rows = await conn.fetch(
                    """
                    SELECT work_state as state, COUNT(*) as count
                    FROM employees
                    WHERE org_id = $1 AND termination_date IS NULL
                      AND work_state = ANY($2::text[])
                      AND (work_location_id IS NULL OR work_location_id != ALL($3::uuid[]))
                    GROUP BY work_state
                    """,
                    company_id, states, list(matched_loc_ids),
                )

    result = []
    for r in rows:
        result.append({
            "location": f"{r['name'] or r['city']}, {r['state']}",
            "count": r["count"],
            "match_type": "exact",
        })
    for r in state_rows:
        result.append({
            "location": r["state"],
            "count": r["count"],
            "match_type": "state",
        })

    return result if result else None


_GAP_KEYWORDS: dict[str, list[str]] = {
    "hipaa_privacy": ["hipaa", "privacy", "phi", "protected health"],
    "workplace_safety": ["safety", "osha", "workplace safety", "injury prevention"],
    "anti_discrimination": ["discrimination", "harassment", "equal employment", "eeo"],
    "sick_leave": ["sick leave", "paid sick", "illness"],
    "leave": ["leave", "fmla", "family leave", "medical leave"],
    "meal_breaks": ["meal", "break", "rest period"],
    "overtime": ["overtime", "hours worked", "flsa"],
    "minimum_wage": ["minimum wage", "wage"],
    "workers_comp": ["workers comp", "work injury", "occupational injury"],
    "cybersecurity": ["cybersecurity", "data security", "breach", "information security"],
    "emergency_preparedness": ["emergency", "disaster", "evacuation"],
    "clinical_safety": ["clinical", "patient safety", "infection control"],
    "billing_integrity": ["billing", "coding", "false claims", "anti-kickback"],
    "telehealth": ["telehealth", "telemedicine", "remote care"],
    "radiation_safety": ["radiation", "radiology", "nuclear"],
}


async def _detect_compliance_gaps(
    company_id: UUID,
    metadata: dict,
) -> list[dict] | None:
    """Detect gaps where jurisdiction requires a written policy but company lacks one."""
    chains = metadata.get("compliance_reasoning", [])
    if not chains:
        return None

    required_categories: set[str] = set()
    for loc in chains:
        for cat in loc.get("categories", []):
            for level in cat.get("all_levels", []):
                if level.get("requires_written_policy") and level.get("is_governing"):
                    required_categories.add(cat["category"])

    if not required_categories:
        return None

    async with get_connection() as conn:
        policies = await conn.fetch(
            "SELECT title FROM policies WHERE company_id = $1 AND status = 'active'",
            company_id,
        )
        handbook_sections = await conn.fetch("""
            SELECT hs.title FROM handbook_sections hs
            JOIN handbook_versions hv ON hv.id = hs.handbook_version_id
            JOIN handbooks h ON h.id = hv.handbook_id
            WHERE h.company_id = $1 AND h.status = 'active'
              AND hv.version_number = h.active_version
        """, company_id)

    all_titles = {
        p["title"].lower() for p in policies if p["title"]
    } | {
        s["title"].lower() for s in handbook_sections if s["title"]
    }

    gaps = []
    for cat in required_categories:
        keywords = _GAP_KEYWORDS.get(cat, [cat.replace("_", " ")])
        has_match = any(any(kw in title for kw in keywords) for title in all_titles)
        if not has_match:
            gaps.append({
                "category": cat,
                "label": cat.replace("_", " ").title(),
                "status": "missing",
            })

    return gaps if gaps else None


def _build_compliance_metadata(
    compliance_result: ComplianceContextResult | None,
    ai_resp,
) -> dict | None:
    """Merge pre-computed jurisdiction reasoning and Gemini's reasoning steps into message metadata."""
    chains = compliance_result.reasoning_chains if compliance_result else None
    ai_steps = ai_resp.compliance_reasoning if ai_resp else None
    if not chains and not ai_steps:
        return None
    metadata: dict = {}
    if chains:
        metadata["compliance_reasoning"] = chains
    if ai_steps:
        metadata["ai_reasoning_steps"] = ai_steps
    if ai_resp and ai_resp.referenced_categories:
        metadata["referenced_categories"] = ai_resp.referenced_categories
    if ai_resp and ai_resp.referenced_locations:
        metadata["referenced_locations"] = ai_resp.referenced_locations
    return metadata


def _sse_data(payload: dict) -> str:
    return f"data: {json.dumps(payload, default=str)}\n\n"


def _append_action_note(reply: str, note: Optional[str]) -> str:
    clean_note = (note or "").strip()
    if not clean_note:
        return reply
    clean_reply = (reply or "").strip()
    if not clean_reply:
        return clean_note
    return f"{clean_reply}\n\n{clean_note}"


# Phrases that reference UI surfaces that DO NOT EXIST for a plain thread:
# there is no project panel / project document / draft panel / canvas for a
# thread. The AI keeps hallucinating claims like "see the full draft in the
# project document" or "I've also initialized a project document" — scrub any
# sentence containing one of these phrases from replies on plain threads.
_PHANTOM_SURFACE_PHRASES = [
    "project document",
    "project panel",
    "project canvas",
    "draft panel",
    "document panel",
    "separate project",
    "new project",
    "a project for",
    "project for you",
    "project for this",
    "project workspace",
]


_PHANTOM_FALLBACK_REPLY = (
    "This is a plain chat — I can't create projects or documents from here. "
    "If you want me to draft something, just tell me what and I'll write it directly in this chat. "
    "Or create a Project from the + button next to Projects in the sidebar for a persistent document workspace."
)


def _scrub_phantom_surface_claims(reply: str, project_id: Optional[UUID]) -> str:
    """Strip sentences that reference a project surface that doesn't exist.

    Plain threads are just chats (no project attached). The AI keeps
    hallucinating it has created/saved something in a 'project document' /
    'project panel' — none of which exist for a plain thread. Strip any
    sentence containing one of those phrases. If scrubbing leaves nothing,
    return a plain-thread fallback so the user sees something actionable
    instead of the hallucinated claim.

    For threads inside a project, leave replies alone — the project's
    surfaces are real.
    """
    if project_id is not None or not reply:
        return reply
    lower_phrases = [p.lower() for p in _PHANTOM_SURFACE_PHRASES]
    out_lines: list[str] = []
    for line in reply.splitlines():
        parts = re.split(r"(?<=[.!?])\s+", line)
        kept = [p for p in parts if not any(ph in p.lower() for ph in lower_phrases)]
        out_lines.append(" ".join(kept))
    cleaned = "\n".join(out_lines)
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    if not cleaned:
        return _PHANTOM_FALLBACK_REPLY
    # If scrubbing produced a stub (short, no real content left), replace it
    # too. Heuristic: under 40 chars with no alphanumeric draft signal.
    if len(cleaned) < 40 and not re.search(r"[A-Za-z]{5,}.*[A-Za-z]{5,}.*[A-Za-z]{5,}", cleaned):
        return _PHANTOM_FALLBACK_REPLY
    return cleaned


def _extract_emails_from_text(text: str) -> list[str]:
    if not text:
        return []
    return doc_svc.normalize_recipient_emails(EMAIL_REGEX.findall(text))


def _looks_like_send_draft_command(content: str) -> bool:
    text = (content or "").lower()
    if not text:
        return False
    has_send = bool(re.search(r"\b(send|email)\b", text))
    has_draft_context = bool(re.search(r"\b(draft|offer(?:\s+letter)?)\b", text))
    has_email = bool(_extract_emails_from_text(content))
    return has_send and has_draft_context and has_email


def _collect_offer_draft_recipients(
    *,
    structured_update: Optional[dict],
    current_state: dict,
    user_message: str,
) -> list[str]:
    collected: list[str] = []

    if isinstance(structured_update, dict):
        raw_recipients = structured_update.get("recipient_emails")
        if isinstance(raw_recipients, list):
            collected.extend(str(v) for v in raw_recipients if str(v).strip())
        elif isinstance(raw_recipients, str) and raw_recipients.strip():
            collected.append(raw_recipients)

        for key in ("candidate_email", "review_recipient_email"):
            raw_email = structured_update.get(key)
            if isinstance(raw_email, str) and raw_email.strip():
                collected.append(raw_email)

    state_recipients = current_state.get("recipient_emails")
    if isinstance(state_recipients, list):
        collected.extend(str(v) for v in state_recipients if str(v).strip())

    candidate_email = current_state.get("candidate_email")
    if isinstance(candidate_email, str) and candidate_email.strip():
        collected.append(candidate_email)

    collected.extend(_extract_emails_from_text(user_message))
    return doc_svc.normalize_recipient_emails(collected)


def _json_object(value) -> dict:
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except json.JSONDecodeError:
            return {}
    return {}


def _coerce_bool(value, default: bool = True) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    if isinstance(value, (int, float)):
        return bool(value)
    if isinstance(value, str):
        normalized = value.strip().lower()
        if normalized in {"1", "true", "yes", "on"}:
            return True
        if normalized in {"0", "false", "no", "off"}:
            return False
    return default


async def _send_mw_provisioning_email(
    *,
    company_id: UUID,
    personal_email: str,
    employee_name: str,
    work_email: str,
    google_result: dict | None,
    slack_result: dict | None,
) -> None:
    """Send a welcome email with provisioning credentials for Matcha Work onboarding."""
    temp_password: str | None = None
    if google_result:
        temp_password = google_result.get("initial_password")

    slack_invite_link: str | None = None
    slack_workspace_name: str | None = None
    if slack_result:
        for step in slack_result.get("steps") or []:
            resp = step.get("last_response") or {}
            if resp.get("invite_link"):
                slack_invite_link = resp["invite_link"]
                break

    google_succeeded = google_result and google_result.get("status") == "completed"
    slack_succeeded = slack_result and slack_result.get("status") == "completed"

    if not google_succeeded and not slack_succeeded:
        return

    async with get_connection() as conn:
        company_name = await conn.fetchval(
            "SELECT name FROM companies WHERE id = $1", company_id,
        ) or "Your Company"
        if slack_succeeded:
            slack_config_row = await conn.fetchval(
                "SELECT config FROM integration_connections WHERE company_id = $1 AND provider = $2",
                company_id, PROVIDER_SLACK,
            )
            if slack_config_row:
                slack_cfg = json.loads(slack_config_row) if isinstance(slack_config_row, str) else slack_config_row
                slack_workspace_name = slack_cfg.get("slack_team_name")

    email_svc = get_email_service()
    await email_svc.send_provisioning_welcome_email(
        to_email=personal_email,
        to_name=employee_name,
        company_name=company_name,
        work_email=work_email if google_succeeded else None,
        temp_password=temp_password,
        slack_workspace_name=slack_workspace_name,
        slack_invite_link=slack_invite_link,
    )


async def _create_onboarding_employees(
    *,
    company_id: UUID,
    triggered_by: UUID,
    employees: list[dict],
) -> list[dict]:
    """Create employee records and trigger provisioning for each. Returns updated employee dicts."""
    results: list[dict] = []

    async with get_connection() as conn:
        # Pre-fetch integration config once for the company
        google_workspace_auto = False
        slack_auto = False
        try:
            integration_rows = await conn.fetch(
                """
                SELECT provider, config
                FROM integration_connections
                WHERE company_id = $1 AND status = 'connected'
                """,
                company_id,
            )
            for irow in integration_rows:
                cfg = _json_object(irow["config"])
                if irow["provider"] == PROVIDER_GOOGLE_WORKSPACE:
                    google_workspace_auto = _coerce_bool(cfg.get("auto_provision_on_employee_create"), True)
                elif irow["provider"] == PROVIDER_SLACK:
                    slack_auto = _coerce_bool(cfg.get("auto_invite_on_employee_create"), True)
        except Exception:
            logger.exception("Unable to query integration connections for company %s", company_id)

        for emp in employees:
            emp_status = (emp.get("status") or "").strip().lower()
            if emp_status in ("created", "done", "error"):
                results.append(emp)
                continue

            first_name = (emp.get("first_name") or "").strip()
            last_name = (emp.get("last_name") or "").strip()
            work_email = (emp.get("work_email") or "").strip().lower()

            if not first_name or not last_name or not work_email:
                emp["status"] = "error"
                emp["error"] = "Missing required fields: first_name, last_name, work_email"
                results.append(emp)
                continue

            try:
                existing = await conn.fetchval(
                    "SELECT id FROM employees WHERE org_id = $1 AND email = $2",
                    company_id, work_email,
                )
                if existing:
                    emp["status"] = "error"
                    emp["error"] = f"Employee with email {work_email} already exists"
                    results.append(emp)
                    continue

                start_date = None
                raw_start = (emp.get("start_date") or "").strip()
                if raw_start:
                    try:
                        start_date = datetime.strptime(raw_start, "%Y-%m-%d").date()
                    except ValueError:
                        pass  # leave as None

                personal_email = (emp.get("personal_email") or "").strip().lower() or None
                address = (emp.get("address") or "").strip() or None

                row = await conn.fetchrow(
                    """
                    INSERT INTO employees (org_id, email, personal_email, first_name, last_name,
                                           work_state, employment_type, start_date, address)
                    VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9)
                    RETURNING id
                    """,
                    company_id,
                    work_email,
                    personal_email,
                    first_name,
                    last_name,
                    (emp.get("work_state") or "").strip() or None,
                    (emp.get("employment_type") or "").strip() or None,
                    start_date,
                    address,
                )

                emp["status"] = "created"
                emp["employee_id"] = str(row["id"])
                emp["error"] = None

                # Trigger provisioning
                prov = {}
                google_result: dict | None = None
                slack_result: dict | None = None

                if google_workspace_auto:
                    try:
                        google_result = await start_google_workspace_onboarding(
                            company_id=company_id,
                            employee_id=row["id"],
                            triggered_by=triggered_by,
                            trigger_source="matcha_work_onboarding",
                        )
                        prov["google_workspace"] = "triggered"
                    except Exception as gex:
                        logger.exception("Google Workspace provisioning failed for %s", work_email)
                        prov["google_workspace"] = f"error: {gex}"

                if slack_auto:
                    try:
                        slack_result = await start_slack_onboarding(
                            company_id=company_id,
                            employee_id=row["id"],
                            triggered_by=triggered_by,
                            trigger_source="matcha_work_onboarding",
                        )
                        prov["slack"] = "triggered"
                    except Exception as sex:
                        logger.exception("Slack provisioning failed for %s", work_email)
                        prov["slack"] = f"error: {sex}"

                if prov:
                    emp["provisioning_results"] = prov

                # Send provisioning welcome email
                if personal_email and (google_result or slack_result):
                    try:
                        await _send_mw_provisioning_email(
                            company_id=company_id,
                            personal_email=personal_email,
                            employee_name=f"{first_name} {last_name}".strip(),
                            work_email=work_email,
                            google_result=google_result,
                            slack_result=slack_result,
                        )
                    except Exception:
                        logger.exception("Failed to send provisioning email for %s", work_email)

            except Exception as e:
                logger.exception("Failed to create employee %s %s: %s", first_name, last_name, e)
                emp["status"] = "error"
                emp["error"] = str(e)

            results.append(emp)

    return results


def _inject_slide_context(msg_dicts: list[dict], current_state: dict, slide_index: Optional[int]) -> None:
    """Prepend selected slide content to the last user message so the AI sees what it's editing.

    Modifies msg_dicts in-place. Only the AI-facing message list is changed — the saved
    database message remains the user's original text.
    """
    if slide_index is None or not msg_dicts:
        return

    # Find slides in top-level or under presentation (workbook)
    slides = current_state.get("slides") or []
    if not slides:
        pres = current_state.get("presentation")
        if isinstance(pres, dict):
            slides = pres.get("slides") or []

    if not slides or not (0 <= slide_index < len(slides)):
        return

    slide = slides[slide_index]
    if not isinstance(slide, dict):
        return

    total = len(slides)
    title = slide.get("title", "Untitled")
    bullets = slide.get("bullets") or []
    speaker_notes = slide.get("speaker_notes", "")

    context_lines = [
        f"[EDITING Slide {slide_index + 1}/{total}: \"{title}\"]",
        "Below is the slide's CURRENT content. The user's request describes changes they want applied to THIS content.",
        "You MUST modify the slide to reflect their request — do NOT return the current content unchanged.",
        f"- Title: {title}",
        f"- Bullets: {json.dumps(bullets)}",
    ]
    if speaker_notes:
        context_lines.append(f"- Speaker Notes: {speaker_notes}")
    context_lines.append("")
    context_lines.append("Apply the following change:")

    context_block = "\n".join(context_lines)

    # Find last user message and prepend context
    for i in range(len(msg_dicts) - 1, -1, -1):
        if msg_dicts[i]["role"] == "user":
            original = msg_dicts[i]["content"]
            msg_dicts[i] = {
                "role": "user",
                "content": f"{context_block}\n{original}",
            }
            break


def _scope_slide_update(ai_resp, current_state: dict, slide_index: Optional[int]) -> None:
    """When a specific slide was targeted, restrict the AI's update to only that slide.

    The AI is instructed to return the full slides array, but it may inadvertently modify
    other slides or presentation-level fields. This function:
    1) Strips non-slide keys (title, subtitle, theme, etc.) from the update
    2) Merges only the targeted slide from the AI response into the existing slides array
    3) Handles workbook presentations where slides live under presentation.slides
    """
    if slide_index is None:
        return
    if not isinstance(ai_resp.structured_update, dict):
        return

    # Strip non-slide presentation fields — only slide content should change
    for key in ("presentation_title", "subtitle", "theme", "cover_image_url", "generated_at"):
        ai_resp.structured_update.pop(key, None)

    ai_slides = ai_resp.structured_update.get("slides")

    # Determine where current slides live
    current_slides = list(current_state.get("slides") or [])

    # Fallback: workbook presentations store slides under presentation.slides
    if not current_slides:
        pres = current_state.get("presentation")
        if isinstance(pres, dict):
            current_slides = list(pres.get("slides") or [])

    if not isinstance(ai_slides, list) or not current_slides:
        return
    if not (0 <= slide_index < len(ai_slides) and 0 <= slide_index < len(current_slides)):
        return

    merged = list(current_slides)
    merged[slide_index] = ai_slides[slide_index]
    ai_resp.structured_update["slides"] = merged


_ACTION_ITEMS_RE = re.compile(
    r"\n?\s*ACTION ITEMS DETECTED:\s*\n((?:\s*[-*•]\s*.+\n?)+)",
    re.IGNORECASE,
)


def _extract_action_items(reply: str) -> tuple[str, list[str]]:
    """Pull the 'ACTION ITEMS DETECTED:' block out of the reply. Returns
    (cleaned_reply, items). If no block, returns (reply, [])."""
    if not reply:
        return reply, []
    match = _ACTION_ITEMS_RE.search(reply)
    if not match:
        return reply, []
    block = match.group(1)
    items: list[str] = []
    for line in block.splitlines():
        stripped = line.strip().lstrip("-*•").strip()
        if stripped:
            items.append(stripped)
    cleaned = _ACTION_ITEMS_RE.sub("", reply).rstrip()
    return cleaned, items


async def _inject_consultation_context(ctx: str, row) -> str:
    """Inject context for a consultation project so the AI can act as a
    freelancer's CRM sidekick: recall prior sessions, surface action items,
    draft communications grounded in the engagement data."""
    project_data = json.loads(row["project_data"]) if isinstance(row["project_data"], str) else (row["project_data"] or {})
    client = project_data.get("client") or {}
    engagement = project_data.get("engagement") or {}
    stage = project_data.get("stage") or "active"
    tags = project_data.get("tags") or []
    sessions = project_data.get("sessions") or []
    action_items = project_data.get("action_items") or []
    deliverables = project_data.get("deliverables") or []

    pricing = engagement.get("pricing_model") or "hourly"
    rate_line = ""
    if pricing == "hourly" and engagement.get("rate_cents_per_hour"):
        rate_line = f" · Rate: ${engagement['rate_cents_per_hour']/100:.2f}/hr"
    elif pricing == "retainer" and engagement.get("monthly_retainer_cents"):
        rate_line = f" · Retainer: ${engagement['monthly_retainer_cents']/100:.2f}/mo"
    elif pricing == "fixed" and engagement.get("fixed_fee_cents"):
        rate_line = f" · Fixed fee: ${engagement['fixed_fee_cents']/100:.2f}"

    client_name = client.get("name") or row["title"] or "this client"
    client_org = client.get("org")
    pc = client.get("primary_contact") or {}
    pc_line = ""
    if pc.get("name") or pc.get("email"):
        pc_line = f"\nPrimary contact: {pc.get('name') or ''} <{pc.get('email') or ''}>".strip()

    # Sessions: last 3 most recent, most-recent first
    sorted_sessions = sorted(
        [s for s in sessions if isinstance(s, dict)],
        key=lambda s: s.get("at") or "",
        reverse=True,
    )[:3]
    session_lines: list[str] = []
    for s in sorted_sessions:
        dur = s.get("duration_min")
        dur_part = f"{dur} min" if dur else "duration n/a"
        billable = ", billable" if s.get("billable") else ""
        notes = (s.get("notes") or "").strip().replace("\n", " ")
        if len(notes) > 240:
            notes = notes[:237] + "..."
        at = (s.get("at") or "")[:10]
        session_lines.append(f"  - {at} ({dur_part}{billable}): {notes or '(no notes)'}")

    open_items = [a for a in action_items if isinstance(a, dict) and not a.get("completed") and not a.get("pending_confirmation")]
    item_lines = [f"  - [ ] {a.get('text','')}" for a in open_items[:10]]

    active_delivs = [d for d in deliverables if isinstance(d, dict) and (d.get("status") or "").lower() in ("planned", "in_progress")]
    deliv_lines = []
    for d in active_delivs[:10]:
        status = d.get("status") or "planned"
        due = f", due {d.get('due_date')}" if d.get("due_date") else ""
        deliv_lines.append(f"  - {d.get('title','(untitled)')} ({status}{due})")

    unbilled_cents = 0
    for s in sessions:
        if not isinstance(s, dict) or not s.get("billable") or s.get("invoice_id"):
            continue
        dur = s.get("duration_min") or 0
        rate = s.get("rate_cents_override") or engagement.get("rate_cents_per_hour") or 0
        unbilled_cents += int(dur * rate / 60)

    ctx += f"""

=== CONSULTATION CONTEXT ===
This chat is tied to the consultation with {client_name}{f' ({client_org})' if client_org else ''}.
Stage: {stage}{rate_line}{(' · Started: ' + engagement['start_date']) if engagement.get('start_date') else ''}{(' · Tags: ' + ', '.join(tags)) if tags else ''}{pc_line}

Recent sessions (most recent first):
{chr(10).join(session_lines) if session_lines else '  (no sessions logged yet)'}

Active deliverables:
{chr(10).join(deliv_lines) if deliv_lines else '  (none)'}

Open action items:
{chr(10).join(item_lines) if item_lines else '  (none)'}

Unbilled this period: ${unbilled_cents/100:.2f}
"""
    return ctx


def _format_blog_mode_state(row) -> str:
    """Return a plain-text block describing the current blog draft (title,
    tone, audience, sections with ids, word counts, etc.). This is passed as
    `blog_mode_state` to the AI provider, which uses a dedicated blog-only
    system prompt — no generic multi-skill prompt and no risk of the AI
    invoking project/workbook/etc skills."""
    project_data = json.loads(row["project_data"]) if isinstance(row["project_data"], str) else (row["project_data"] or {})
    sections = json.loads(row["sections"]) if isinstance(row["sections"], str) else (row["sections"] or [])
    title = row["title"] or "(untitled)"
    slug = project_data.get("slug") or "(auto)"
    tone = project_data.get("tone") or "expert-casual"
    audience = project_data.get("audience") or "(not set — ask the user before drafting)"
    tags = project_data.get("tags") or []
    status = project_data.get("status") or "draft"
    excerpt = project_data.get("excerpt") or "(none)"
    stats = project_data.get("stats") or {}

    section_lines: list[str] = []
    for i, s in enumerate(sections, start=1):
        if not isinstance(s, dict):
            continue
        sid = s.get("id") or "?"
        st = (s.get("title") or "(untitled)").strip()
        content = s.get("content") or ""
        wc = len([w for w in re.split(r"\s+", content.strip()) if w])
        section_lines.append(f'  {i}. id={sid} · "{st}" ({wc} words)')
    sec_block = "\n".join(section_lines) if section_lines else "  (no sections yet — if the user wants to start, emit blog_outline to seed them)"

    return (
        f"Title: {title}\n"
        f"Slug: {slug}\n"
        f"Status: {status}\n"
        f"Audience: {audience}\n"
        f"Tone: {tone}\n"
        f"Tags: {', '.join(tags) if tags else '(none)'}\n"
        f"Word count: {stats.get('word_count', 0)} · Reading time: {stats.get('read_minutes', 0)} min\n"
        f"Excerpt: {excerpt}\n\n"
        f"Sections (in order — use these exact ids when emitting blog_section_draft/blog_section_revision):\n"
        f"{sec_block}"
    )


async def _fetch_project_meta(project_id: Optional[UUID]) -> Optional[dict]:
    """Fetch the project row needed by AI-turn helpers (skill routing + blog
    prompt). One query per turn — callers pass the result to both
    `_apply_ai_updates_and_operations` and `_blog_mode_state_from_meta`.
    """
    if not project_id:
        return None
    async with get_connection() as conn:
        row = await conn.fetchrow(
            "SELECT title, project_type, sections, project_data FROM mw_projects WHERE id = $1",
            project_id,
        )
    return dict(row) if row else None


def _blog_mode_state_from_meta(project_meta: Optional[dict]) -> Optional[str]:
    """If the project meta describes a blog, return the formatted state block
    for the dedicated blog system prompt. Otherwise None."""
    if not project_meta or project_meta.get("project_type") != "blog":
        return None
    return _format_blog_mode_state(project_meta)


async def _inject_recruiting_project_context(ctx: str, thread: dict, current_state: dict) -> str:
    """If this thread belongs to a recruiting project, inject context so the AI
    generates posting sections instead of creating a new project. If it's a
    consultation project, dispatch to the consultation injector instead."""
    project_id = thread.get("project_id")
    if not project_id:
        return ctx

    from ..services import project_service as proj_svc
    async with get_connection() as conn:
        row = await conn.fetchrow(
            "SELECT title, project_type, sections, project_data FROM mw_projects WHERE id = $1",
            project_id,
        )
    if not row:
        return ctx
    if row["project_type"] == "consultation":
        return await _inject_consultation_context(ctx, row)
    if row["project_type"] == "blog":
        # Blog projects use a dedicated system prompt (built separately via
        # _fetch_blog_mode_state) — skip the generic context injection.
        return ctx
    if row["project_type"] != "recruiting":
        return ctx

    sections = json.loads(row["sections"]) if isinstance(row["sections"], str) else (row["sections"] or [])
    project_data = json.loads(row["project_data"]) if isinstance(row["project_data"], str) else (row["project_data"] or {})
    posting = project_data.get("posting") or {}
    is_finalized = bool(posting.get("finalized"))
    candidates = project_data.get("candidates") or []
    candidates_count = len(candidates)
    section_count = len(sections)
    position_title = row["title"] or "the open role"

    # Build a compact candidate roster for the AI so it can auto-fill offer
    # letters without asking the user to re-enter details that are already in
    # the project. Include id so the AI can disambiguate by explicit id if
    # the user says "the first candidate" etc.
    shortlist_ids = set(project_data.get("shortlist_ids") or [])
    dismissed_ids = set(project_data.get("dismissed_ids") or [])
    roster_lines: list[str] = []
    for c in candidates[:50]:
        if c.get("id") in dismissed_ids:
            status_tag = "dismissed"
        elif c.get("id") in shortlist_ids:
            status_tag = "shortlisted"
        else:
            status_tag = c.get("status") or "pending"
        parts = [
            f"id={c.get('id', '?')}",
            f"name={c.get('name') or '(unknown)'}",
            f"email={c.get('email') or '(no email)'}",
        ]
        if c.get("current_title"):
            parts.append(f"current_title={c['current_title']}")
        if c.get("location"):
            parts.append(f"location={c['location']}")
        parts.append(f"status={status_tag}")
        roster_lines.append("- " + ", ".join(parts))
    roster_block = "\n".join(roster_lines) if roster_lines else "(no candidates uploaded yet)"

    ctx += f"""

=== RECRUITING PROJECT CONTEXT ===
This chat is part of a RECRUITING project. You are helping the user hire for a role.
- Position title (use this for offer letters): {position_title}
- Posting finalized: {is_finalized}
- Existing posting sections: {section_count}
- Candidates uploaded: {candidates_count}

CANDIDATE ROSTER (authoritative — use these exact values, do NOT ask the user to re-enter them):
{roster_block}

CRITICAL RULES FOR RECRUITING PROJECTS:
1. When the user describes a role or asks you to create a job posting, first determine whether you have enough information to draft.
   Required signals before drafting: location (city + remote/hybrid/on-site), employment type (FT/PT/contract + hours), ballpark compensation or explicit "open to discussion", and 2–3 role-specific responsibilities or must-haves.
   If any of those are missing, DO NOT emit project_sections yet and DO NOT claim a posting has been drafted.
   Instead respond conversationally (mode="general", operation="none", empty updates) with 2–4 concise clarifying questions covering the missing signals.
   Once you have the signals (either gathered across turns or supplied in one message), generate the full posting using the "project" skill with project_sections like "About the Role", "Responsibilities", "Requirements", "Compensation & Benefits". The sections will automatically appear in the project's Posting panel on the right.
   Never say "I've drafted the posting" in the same reply where project_sections is empty — only confirm a draft when you are actually emitting the sections.
2. Do NOT create a new project. The project already exists — when ready, just output project_sections with the posting content.
3. NEVER output raw JSON, code, SVG, or internal state in your responses. Always respond in clear, human-readable language.
4. To send interviews: tell the user to select candidates in the pipeline panel and click "Send Interviews".
5. To upload resumes: tell the user to click the paperclip icon or drag-and-drop PDF resumes into the chat.
6. To analyze candidates: tell the user to click "Analyze" in the Candidates tab of the pipeline panel.
7. Keep responses concise and actionable — guide the user through the recruiting workflow step by step.

OFFER LETTER AUTO-FILL (important):
8. When the user asks to generate/draft/create an offer letter for a specific candidate — whether they name them ("draft an offer for Mark"), point at them ("offer to the shortlisted one", "to the first candidate", "to this one"), or say they've decided to hire someone — you MUST pull candidate_name, candidate_email, and position_title directly from the CANDIDATE ROSTER above. Do NOT ask the user to re-type these fields.
9. If the user's reference is ambiguous (e.g. multiple shortlisted candidates), ask a single clarifying question naming the options from the roster — do not ask them to re-enter contact details.
10. Only prompt the user for fields that are NOT in the roster: salary, start_date, employment_type, benefits, etc.
11. When in doubt, prefer the shortlisted candidate. If there's exactly one candidate in the roster and the user says "generate the offer letter", use that candidate without asking who.
"""
    return ctx


async def _apply_ai_updates_and_operations(
    *,
    thread_id: UUID,
    company_id: UUID,
    ai_resp,
    current_state: dict,
    current_version: int,
    user_message: str,
    current_user_id: Optional[UUID] = None,
    project_id: Optional[UUID] = None,
    project_meta: Optional[dict] = None,
) -> tuple[dict, int, Optional[str], bool, str]:
    """Apply structured updates, execute supported operations, and return updated response state.

    `project_meta` (when supplied by the caller) contains at least
    `project_type` for the thread's project. We rely on it to drive the
    blog-skill routing below without a redundant DB fetch.
    """
    project_type_hint = (project_meta or {}).get("project_type") if project_id else None

    skill = ai_resp.skill or _infer_skill_from_state(current_state)
    # Blog projects always route to the blog skill — covers both legacy threads
    # whose state inference would otherwise return "project", and any future
    # prompt drift. With the dedicated blog system prompt in place, this is a
    # safety net rather than the primary defence.
    if project_type_hint == "blog":
        skill = "blog"
    # If skill is not a known document type (e.g. "none" or "chat"), fall back to
    # inferring from the update keys themselves so workbook/review/etc. updates
    # created on a fresh thread aren't silently dropped.
    elif skill not in ("offer_letter", "review", "workbook", "onboarding", "presentation", "handbook", "policy", "resume_batch", "inventory", "project", "blog") and isinstance(ai_resp.structured_update, dict) and ai_resp.structured_update:
        skill_from_updates = _infer_skill_from_state(ai_resp.structured_update)
        if skill_from_updates != "chat":
            skill = skill_from_updates
    initial_version = int(current_version)
    pdf_url: Optional[str] = None
    assistant_reply = ai_resp.assistant_reply
    changed = False

    should_execute_skill = bool(
        ai_resp.mode == "skill"
        and (ai_resp.confidence >= 0.65 or not ai_resp.missing_fields)
    )
    logger.info(
        "[MW] ai mode=%s skill=%s op=%s conf=%.2f missing=%s update_keys=%s should_exec=%s",
        ai_resp.mode, skill, ai_resp.operation, ai_resp.confidence,
        ai_resp.missing_fields,
        list(ai_resp.structured_update.keys()) if isinstance(ai_resp.structured_update, dict) else None,
        should_execute_skill,
    )
    force_send_draft = (
        skill == "offer_letter"
        and _looks_like_send_draft_command(user_message)
    )
    can_execute_operation = should_execute_skill or force_send_draft
    blog_directives: dict = {}

    if should_execute_skill and ai_resp.structured_update:
        safe_updates = _validate_updates_for_skill(skill, ai_resp.structured_update)
        if skill == "handbook" and current_state.get("handbook_source_type") == "upload":
            safe_updates = {}

        # Blog directives aren't part of the thread document schema — strip them
        # out of safe_updates and process them in the blog directive handler.
        # _validate_updates_for_skill("blog", ...) already whitelists BLOG_FIELDS
        # only, so nothing else survives for blog chats.
        if skill == "blog":
            from ..services.matcha_work_ai import BLOG_FIELDS as _BLOG_FIELDS
            for _bk in _BLOG_FIELDS:
                if _bk in safe_updates:
                    blog_directives[_bk] = safe_updates.pop(_bk)

        # No-project guard: don't persist project_title/project_sections on
        # threads with no linked project. Otherwise the state accumulates a
        # phantom project the user can't see and the AI keeps "updating" it.
        if not project_id:
            for _k in ("project_title", "project_sections", "project_status"):
                safe_updates.pop(_k, None)
        if safe_updates:
            result = await doc_svc.apply_update(thread_id, safe_updates)
            current_version = result["version"]
            current_state = result["current_state"]
            changed = changed or current_version != initial_version

            # Auto-sync AI-generated project_sections to the project's sections table.
            # Match by title (case-insensitive) so regenerations update existing
            # sections in place instead of appending duplicates. Sections added
            # manually by the user keep their own titles and are preserved.
            # Skip this sync for consultation and blog projects:
            #   - Consultations are CRM engagements, not document drafts
            #   - Blog drafts use blog_outline / blog_section_draft / blog_sections_replace
            #     directives; a stray project_sections emission would silently
            #     append to the blog's sections (was a real bug, see commit
            #     history). Use project_meta to decide without another DB fetch.
            _project_type = (project_meta or {}).get("project_type") if project_id else None
            _skip_project_sections_sync = _project_type in ("consultation", "blog")

            if project_id and not _skip_project_sections_sync and "project_sections" in safe_updates:
                from ..services import project_service as proj_svc
                new_sections = safe_updates.get("project_sections") or []
                if new_sections:
                    try:
                        existing = list(await proj_svc.get_sections(project_id))
                        existing_by_title: dict[str, int] = {}
                        for idx, s in enumerate(existing):
                            title_key = (s.get("title") or "").strip().lower()
                            if title_key and title_key not in existing_by_title:
                                existing_by_title[title_key] = idx
                        changed_sections = False
                        for section in new_sections:
                            if not isinstance(section, dict):
                                continue
                            content = (section.get("content") or "").strip()
                            if not content:
                                continue
                            title = (section.get("title") or "").strip()
                            title_key = title.lower()
                            existing_idx = existing_by_title.get(title_key) if title_key else None
                            if existing_idx is not None:
                                # Update existing section in place
                                merged = {**existing[existing_idx], "content": content}
                                if title:
                                    merged["title"] = title
                                existing[existing_idx] = merged
                                changed_sections = True
                            else:
                                # Append a new section with a fresh id
                                new_entry = {
                                    "id": os.urandom(8).hex(),
                                    "title": title or None,
                                    "content": content,
                                    "source_message_id": None,
                                }
                                existing.append(new_entry)
                                if title_key:
                                    existing_by_title[title_key] = len(existing) - 1
                                changed_sections = True
                        if changed_sections:
                            await proj_svc._update_sections(project_id, existing)
                            # Mirror the section content into project_data.posting.content
                            # so the Posting tab actually renders what the AI wrote.
                            # Without this, sections live on project.sections but the
                            # posting tab (which reads posting.content) stays empty —
                            # so the AI says "drafted it" but the user sees nothing.
                            try:
                                composed_parts: list[str] = []
                                for s in existing:
                                    if not isinstance(s, dict):
                                        continue
                                    content = (s.get("content") or "").strip()
                                    if not content:
                                        continue
                                    title = (s.get("title") or "").strip()
                                    if title:
                                        composed_parts.append(f"## {title}\n\n{content}")
                                    else:
                                        composed_parts.append(content)
                                composed = "\n\n".join(composed_parts)
                                if composed:
                                    async with get_connection() as _pconn:
                                        row = await _pconn.fetchrow(
                                            "SELECT project_data FROM mw_projects WHERE id = $1", project_id
                                        )
                                    if row:
                                        existing_data = row["project_data"]
                                        if isinstance(existing_data, str):
                                            existing_data = json.loads(existing_data or "{}")
                                        existing_data = existing_data or {}
                                        prior_posting = dict(existing_data.get("posting") or {})
                                        # Don't trample a finalized posting — operator may
                                        # have locked the content on purpose.
                                        if not prior_posting.get("finalized"):
                                            prior_posting["content"] = composed
                                            await proj_svc.update_project_data(
                                                project_id, {"posting": prior_posting}
                                            )
                            except Exception:
                                logger.warning(
                                    "Failed to mirror sections into posting for project %s",
                                    project_id, exc_info=True,
                                )
                    except Exception:
                        logger.warning("Failed to sync project_sections to project %s", project_id, exc_info=True)

            inferred = _infer_skill_from_state(current_state)
            if inferred == "offer_letter":
                pdf_url = await doc_svc.generate_pdf(
                    current_state,
                    thread_id,
                    current_version,
                    is_draft=True,
                    company_id=company_id,
                )
            elif inferred == "presentation" and not current_state.get("cover_image_url"):
                # Generate a cover image the first time slides are created
                cover_url = await doc_svc.generate_cover_image(
                    presentation_title=str(current_state.get("presentation_title") or "Presentation"),
                    subtitle=current_state.get("subtitle"),
                    company_id=company_id,
                    thread_id=thread_id,
                )
                if cover_url:
                    cover_result = await doc_svc.apply_update(thread_id, {"cover_image_url": cover_url})
                    current_version = cover_result["version"]
                    current_state = cover_result["current_state"]

    # Apply blog directives atomically under a row lock (prevents lost updates
    # from concurrent manual edits). blog_title_suggestions is informational —
    # AI describes options in reply text, no persistence.
    blog_changes_applied = False
    if blog_directives and project_id:
        from ..services import project_service as _blog_proj_svc
        try:
            _, blog_secs_changed = await _blog_proj_svc.apply_blog_directives(
                project_id,
                outline=blog_directives.get("blog_outline") if isinstance(blog_directives.get("blog_outline"), list) else None,
                draft=blog_directives.get("blog_section_draft") if isinstance(blog_directives.get("blog_section_draft"), dict) else None,
                revision=blog_directives.get("blog_section_revision") if isinstance(blog_directives.get("blog_section_revision"), dict) else None,
                replace=blog_directives.get("blog_sections_replace") if isinstance(blog_directives.get("blog_sections_replace"), list) else None,
            )
            if blog_secs_changed:
                changed = True
                blog_changes_applied = True
        except Exception:
            logger.warning("Failed to apply blog section directives for project %s", project_id, exc_info=True)

    # "Say it, do it" enforcement for blog chats. If the AI claimed to have
    # added/drafted/written/updated something but no directive actually
    # applied, strip the claim from the reply and replace it with an honest
    # propose-then-act line. Eliminates the pattern where the AI says
    # "I've put together an outline in the Write tab" while the sections
    # panel stays empty.
    if project_type_hint == "blog" and not blog_changes_applied and assistant_reply:
        claim_patterns = [
            r"\bI['’]ve\s+(?:put\s+together|added|drafted|written|written\s+up|updated|revised|structured|consolidated|created|outlined|put\s+down|laid\s+out|pulled\s+together|organized|compiled|assembled|set\s+up|composed)\b[^.!?]*[.!?]",
            r"\b(?:here['’]s|I['’]ve\s+got)\s+(?:an?|the|your)\s+(?:initial\s+)?outline[^.!?]*[.!?]",
            r"\bin\s+the\s+(?:Write|Preview|Publish)\s+tab[^.!?]*[.!?]",
            r"\bthe\s+outline\s+is\s+(?:now\s+)?in[^.!?]*[.!?]",
            r"\b(?:added|updated|populated|seeded|filled)\s+(?:the\s+)?(?:outline|sections?|draft)[^.!?]*[.!?]",
        ]
        import re as _re_claim
        stripped_reply = assistant_reply
        found_claim = False
        for pat in claim_patterns:
            if _re_claim.search(pat, stripped_reply, flags=_re_claim.IGNORECASE):
                found_claim = True
                stripped_reply = _re_claim.sub(pat, "", stripped_reply, flags=_re_claim.IGNORECASE)
        if found_claim:
            stripped_reply = _re_claim.sub(r"[ \t]+", " ", stripped_reply)
            stripped_reply = _re_claim.sub(r"\s+([.!?,;])", r"\1", stripped_reply)
            stripped_reply = _re_claim.sub(r"\n{3,}", "\n\n", stripped_reply).strip()
            honest_tail = (
                "\n\n(Note: I haven't actually written that to the Write tab yet — "
                "tell me to go ahead and I'll draft the outline now.)"
            )
            assistant_reply = (stripped_reply or "Ready when you are.") + honest_tail
            logger.warning(
                "[MW-blog] scrubbed phantom outline claim from reply for project %s thread %s",
                project_id, thread_id,
            )

    operation = str(ai_resp.operation or "none").strip().lower()
    if force_send_draft and operation in {"none", "create", "update", "track"}:
        operation = "send_draft"

    if can_execute_operation and operation not in {"none", "create", "update", "track"}:
        action_note: Optional[str] = None
        try:
            if operation == "save_draft":
                if skill != "offer_letter":
                    action_note = "Save draft is only available for offer letters."
                else:
                    saved = await doc_svc.save_offer_letter_draft(thread_id, company_id)
                    action_note = f"Saved draft successfully ({saved['offer_status']})."
            elif operation == "send_draft":
                if skill != "offer_letter":
                    action_note = "Draft sending is only available for offer letters."
                else:
                    recipients = _collect_offer_draft_recipients(
                        structured_update=ai_resp.structured_update
                        if isinstance(ai_resp.structured_update, dict)
                        else None,
                        current_state=current_state,
                        user_message=user_message,
                    )
                    send_result = await doc_svc.send_offer_letter_draft(
                        thread_id=thread_id,
                        company_id=company_id,
                        recipient_emails=recipients,
                    )
                    if send_result.get("pdf_url"):
                        pdf_url = send_result["pdf_url"]
                    action_note = (
                        f"Draft email send complete: {send_result['sent_count']} sent, "
                        f"{send_result['failed_count']} failed."
                    )
                    if send_result["failed_count"] > 0:
                        failed_recipients = [
                            str(row.get("email"))
                            for row in send_result.get("recipients", [])
                            if row.get("status") != "sent" and row.get("email")
                        ]
                        if failed_recipients:
                            action_note += f" Failed recipient(s): {', '.join(failed_recipients[:3])}."
            elif operation == "send_requests":
                if skill != "review":
                    action_note = "Sending review requests is only available in review threads."
                else:
                    recipient_emails: list[str] = []
                    if isinstance(ai_resp.structured_update, dict):
                        raw_emails = ai_resp.structured_update.get("recipient_emails")
                        if isinstance(raw_emails, list):
                            recipient_emails = [str(v) for v in raw_emails if str(v).strip()]

                    send_result = await doc_svc.send_review_requests(
                        thread_id=thread_id,
                        company_id=company_id,
                        recipient_emails=recipient_emails,
                    )
                    refreshed = await doc_svc.get_thread(thread_id, company_id)
                    if refreshed:
                        current_state = refreshed["current_state"]
                        current_version = refreshed["version"]
                        changed = changed or current_version != initial_version
                    action_note = (
                        f"Sent {send_result['sent_count']} request(s), "
                        f"{send_result['failed_count']} failed. "
                        f"Received {send_result['received_responses']}/{send_result['expected_responses']}."
                    )
            elif operation == "finalize":
                finalized = await doc_svc.finalize_thread(thread_id, company_id)
                refreshed = await doc_svc.get_thread(thread_id, company_id)
                if refreshed:
                    current_state = refreshed["current_state"]
                    current_version = refreshed["version"]
                if finalized.get("pdf_url"):
                    pdf_url = finalized["pdf_url"]
                action_note = "Thread finalized."
            elif operation == "create_employees":
                if skill != "onboarding":
                    action_note = "Employee creation is only available in onboarding threads."
                else:
                    raw_employees = current_state.get("employees")
                    if not isinstance(raw_employees, list) or not raw_employees:
                        action_note = "No employees found to create. Please add employee details first."
                    else:
                        results = await _create_onboarding_employees(
                            company_id=company_id,
                            triggered_by=current_user_id,
                            employees=[dict(e) for e in raw_employees],
                        )
                        current_state["employees"] = results
                        current_state["batch_status"] = "complete"
                        created = sum(1 for e in results if e.get("status") == "created")
                        errors = sum(1 for e in results if e.get("status") == "error")
                        result = await doc_svc.apply_update(
                            thread_id,
                            current_state,
                            diff_summary=f"Created {created} employee(s), {errors} error(s)",
                        )
                        current_version = result["version"]
                        current_state = result["current_state"]
                        changed = True
                        action_note = f"Created {created} employee(s)"
                        if errors:
                            error_names = [
                                f"{e.get('first_name', '?')} {e.get('last_name', '?')}: {e.get('error', 'unknown')}"
                                for e in results if e.get("status") == "error"
                            ]
                            action_note += f", {errors} failed: " + "; ".join(error_names[:5])
                        action_note += "."
            elif operation == "generate_presentation":
                if skill != "workbook":
                    action_note = "Presentation generation is only available in workbook threads."
                else:
                    generated = await doc_svc.generate_workbook_presentation(
                        thread_id=thread_id,
                        company_id=company_id,
                    )
                    current_state = generated["current_state"]
                    current_version = generated["version"]
                    changed = True
                    action_note = (
                        f"Generated a presentation with {generated['slide_count']} slides. "
                        "Open Preview to review and download."
                    )
            elif operation == "generate_handbook":
                if skill != "handbook":
                    action_note = "Handbook generation is only available in handbook threads."
                elif current_state.get("handbook_source_type") == "upload":
                    action_note = "Handbook generation is unavailable in upload review mode. Start a new template handbook thread instead."
                else:
                    title = current_state.get("handbook_title")
                    states = current_state.get("handbook_states") or []
                    legal_name = current_state.get("handbook_legal_name")
                    ceo = current_state.get("handbook_ceo")

                    if not title or not states or not legal_name or not ceo:
                        missing = []
                        if not title: missing.append("handbook_title")
                        if not states: missing.append("handbook_states")
                        if not legal_name: missing.append("handbook_legal_name")
                        if not ceo: missing.append("handbook_ceo")
                        action_note = f"Cannot generate — missing required fields: {', '.join(missing)}"
                    else:
                        await doc_svc.apply_update(thread_id, {"handbook_status": "generating"})

                        from ...core.models.handbook import (
                            HandbookCreateRequest, HandbookScopeInput,
                            CompanyHandbookProfileInput, HandbookSectionInput,
                        )
                        profile_data = current_state.get("handbook_profile") or {}
                        raw = profile_data if isinstance(profile_data, dict) else {}

                        scopes = [HandbookScopeInput(state=s.upper()) for s in states]
                        mode = "single_state" if len(scopes) == 1 else "multi_state"

                        profile = CompanyHandbookProfileInput(
                            legal_name=legal_name,
                            dba=current_state.get("handbook_dba"),
                            ceo_or_president=ceo,
                            headcount=current_state.get("handbook_headcount"),
                            remote_workers=raw.get("remote_workers", False),
                            minors=raw.get("minors", False),
                            tipped_employees=raw.get("tipped_employees", False),
                            tip_pooling=raw.get("tip_pooling", False),
                            union_employees=raw.get("union_employees", False),
                            federal_contracts=raw.get("federal_contracts", False),
                            group_health_insurance=raw.get("group_health_insurance", False),
                            background_checks=raw.get("background_checks", False),
                            hourly_employees=raw.get("hourly_employees", True),
                            salaried_employees=raw.get("salaried_employees", False),
                            commissioned_employees=raw.get("commissioned_employees", False),
                        )

                        custom_sections_raw = current_state.get("handbook_custom_sections") or []
                        custom_sections = [
                            HandbookSectionInput(
                                section_key=f"custom_{i}",
                                title=s.get("title", ""),
                                content=s.get("content", ""),
                                section_order=900 + i,
                                section_type="custom",
                            )
                            for i, s in enumerate(custom_sections_raw)
                            if isinstance(s, dict) and s.get("title")
                        ]

                        req = HandbookCreateRequest(
                            title=title,
                            mode=mode,
                            source_type="template",
                            industry=current_state.get("handbook_industry"),
                            scopes=scopes,
                            profile=profile,
                            custom_sections=custom_sections,
                            guided_answers=current_state.get("handbook_guided_answers") or {},
                            create_from_template=True,
                        )

                        handbook = await HandbookService.create_handbook(
                            company_id=str(company_id),
                            data=req,
                            created_by=str(current_user_id) if current_user_id else None,
                        )

                        section_previews = [
                            {"section_key": s.section_key, "title": s.title,
                             "content": (s.content or "")[:500], "section_type": s.section_type}
                            for s in (handbook.sections or [])
                        ]

                        coverage_updates = {}
                        try:
                            async with get_connection() as cov_conn:
                                industry_key = await cov_conn.fetchval(
                                    "SELECT industry FROM companies WHERE id = $1", str(company_id)
                                )
                            coverage = HandbookService.compute_coverage(handbook, industry_key or "")
                            coverage_updates = {
                                "handbook_strength_score": coverage.strength_score,
                                "handbook_strength_label": coverage.strength_label,
                            }
                        except Exception as cov_err:
                            logger.warning("Coverage computation failed: %s", cov_err)

                        result = await doc_svc.apply_update(thread_id, {
                            "handbook_status": "created",
                            "handbook_id": str(handbook.id),
                            "handbook_mode": mode,
                            "handbook_sections": section_previews,
                            **coverage_updates,
                        })
                        current_version = result["version"]
                        current_state = result["current_state"]
                        changed = True

                        section_count = len(handbook.sections or [])
                        score_note = ""
                        if coverage_updates.get("handbook_strength_score") is not None:
                            score_note = f" Coverage score: {coverage_updates['handbook_strength_score']}/100 ({coverage_updates['handbook_strength_label']})."
                        action_note = f"Handbook created with {section_count} sections.{score_note}"
            elif operation == "generate_policy":
                if skill != "policy":
                    action_note = "Policy generation is only available in policy threads."
                else:
                    policy_type = current_state.get("policy_type")
                    location_names = current_state.get("policy_location_names") or []

                    if not policy_type:
                        action_note = "Cannot generate — missing required field: policy_type"
                    elif not location_names:
                        action_note = "Cannot generate — please specify at least one location (city, state)."
                    else:
                        await doc_svc.apply_update(thread_id, {"policy_status": "generating"})

                        # Resolve location names to location IDs from business_locations
                        location_ids: list[str] = []
                        async with get_connection() as conn:
                            for loc_name in location_names:
                                parts = [p.strip() for p in loc_name.split(",")]
                                if len(parts) == 2:
                                    city, state = parts[0], parts[1]
                                    row = await conn.fetchrow(
                                        "SELECT id FROM business_locations WHERE company_id = $1 AND city ILIKE $2 AND state ILIKE $3",
                                        str(company_id), city, state,
                                    )
                                    if row:
                                        location_ids.append(str(row["id"]))

                        from ...core.services.policy_draft_service import generate_policy_draft_stream, PolicyDraftRequest

                        draft_request = PolicyDraftRequest(
                            policy_type=policy_type,
                            location_ids=location_ids if location_ids else None,
                            additional_context=current_state.get("policy_additional_context"),
                        )

                        # Collect the full generated content
                        policy_content = ""
                        async for event in generate_policy_draft_stream(str(company_id), draft_request):
                            if event.get("type") == "content":
                                policy_content += event.get("text", "")
                            elif event.get("type") == "error":
                                raise ValueError(event.get("message", "Policy generation failed"))

                        result = await doc_svc.apply_update(thread_id, {
                            "policy_status": "created",
                            "policy_content": policy_content,
                        })
                        current_version = result["version"]
                        current_state = result["current_state"]
                        changed = True
                        action_note = "Policy draft generated. Review in the Preview panel, then edit or save."
            else:
                action_note = f"The action '{operation}' is not supported yet."
        except ValueError as e:
            action_note = str(e)
        except Exception as e:
            logger.error(
                "Failed Matcha Work operation '%s' for thread %s: %s",
                operation,
                thread_id,
                e,
                exc_info=True,
            )
            action_note = "I understood the command, but couldn't complete it right now."

        assistant_reply = _append_action_note(assistant_reply, action_note)

    assistant_reply = _scrub_phantom_surface_claims(assistant_reply, project_id)

    return current_state, current_version, pdf_url, changed, assistant_reply


@router.post("/threads", response_model=CreateThreadResponse, status_code=201)
async def create_thread(
    body: CreateThreadRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Create a new chat thread (optionally with an initial message)."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=400, detail="No company associated with this account")

    title = body.title or "New Chat"
    thread = await doc_svc.create_thread(company_id, current_user.id, title)
    thread_id = thread["id"]

    assistant_reply = None
    pdf_url = None

    if body.initial_message:
        # Save user message
        await doc_svc.add_message(thread_id, "user", body.initial_message)

        # Call AI with company context (profile already fetched during create_thread)
        ai_provider = get_ai_provider()
        profile = await doc_svc.get_company_profile_for_ai(company_id)
        ctx = _build_company_context(profile)
        messages = [{"role": "user", "content": body.initial_message}]
        ai_resp = await ai_provider.generate(messages, thread["current_state"], company_context=ctx)
        final_usage = ai_resp.token_usage

        new_version = thread["version"]
        (
            updated_state,
            new_version,
            pdf_url,
            changed,
            assistant_reply_text,
        ) = await _apply_ai_updates_and_operations(
            thread_id=thread_id,
            company_id=company_id,
            ai_resp=ai_resp,
            current_state=thread["current_state"],
            current_version=new_version,
            user_message=body.initial_message,
            current_user_id=current_user.id,
            project_id=thread.get("project_id"),
        )
        thread["current_state"] = updated_state

        await doc_svc.add_message(
            thread_id,
            "assistant",
            assistant_reply_text,
            version_created=new_version if changed else None,
        )
        try:
            await doc_svc.log_token_usage_event(
                company_id=company_id,
                user_id=current_user.id,
                thread_id=thread_id,
                token_usage=final_usage,
                operation="send_message",
            )
        except Exception as e:
            logger.warning("Failed to log Matcha Work token usage for thread %s: %s", thread_id, e)

        thread["version"] = new_version
        assistant_reply = assistant_reply_text

    return CreateThreadResponse(
        id=thread_id,
        title=thread["title"],
        status=thread["status"],
        current_state=thread["current_state"],
        version=thread["version"],
        task_type=_infer_skill_from_state(thread["current_state"]),
        is_pinned=thread.get("is_pinned", False),
        node_mode=thread.get("node_mode", False),
        compliance_mode=thread.get("compliance_mode", False),
        created_at=thread["created_at"],
        assistant_reply=assistant_reply,
        pdf_url=pdf_url,
    )


@router.post("/threads/{thread_id}/logo")
async def upload_thread_logo(
    thread_id: UUID,
    file: UploadFile = File(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Upload a company logo for an offer-letter thread and save it for the company."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    if _infer_skill_from_state(thread.get("current_state") or {}) != "offer_letter":
        raise HTTPException(status_code=400, detail="Logo upload is only available for offer letters")

    # Validate file type
    if not file.content_type or not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="Only image files are allowed")

    try:
        content = await file.read()
        if len(content) > 5 * 1024 * 1024:  # 5MB limit
            raise HTTPException(status_code=400, detail="File size exceeds 5MB limit")

        # Upload to storage
        logo_url = await get_storage().upload_file(
            content,
            file.filename or "logo.png",
            prefix=f"company-logos/{company_id}",
            content_type=file.content_type,
        )

        # Update thread state
        await doc_svc.apply_update(
            thread_id,
            {"company_logo_url": logo_url},
            diff_summary="Uploaded company logo",
        )

        # Also update company logo for future use
        async with get_connection() as conn:
            await conn.execute(
                "UPDATE companies SET logo_url = $1 WHERE id = $2",
                logo_url,
                company_id,
            )

        return {"logo_url": logo_url}

    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to upload logo for thread %s: %s", thread_id, e, exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to upload logo. Please try again.")


@router.post("/threads/{thread_id}/handbook/upload")
async def upload_thread_handbook(
    thread_id: UUID,
    file: UploadFile = File(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Upload an existing handbook and audit it against synced company jurisdictions.

    Returns JSON (ThreadDetailResponse) for blocked/rejected uploads, or an SSE
    stream with quarterly progress events for the happy-path analysis.
    """
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    if not _thread_accepts_handbook_upload(thread):
        raise HTTPException(
            status_code=400,
            detail="Handbook upload is only available in a new chat or an upload-mode handbook thread.",
        )

    filename = (file.filename or "handbook.pdf").strip() or "handbook.pdf"
    extension = os.path.splitext(filename)[1].lower()
    if extension not in HANDBOOK_UPLOAD_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and DOC handbooks are supported")

    active_locations = [
        loc for loc in await get_locations(company_id)
        if loc.get("is_active", True)
    ]
    active_location_labels = [_location_label(loc) for loc in active_locations if _location_label(loc)]
    unsynced_labels = [
        _location_label(loc)
        for loc in active_locations
        if loc.get("data_status") != "synced" and _location_label(loc)
    ]
    if not active_locations or unsynced_labels:
        blocking_message = _build_handbook_block_message(
            unsynced_labels if unsynced_labels else active_location_labels
        )
        result = await doc_svc.apply_update(
            thread_id,
            {
                "handbook_source_type": "upload",
                "handbook_upload_status": "blocked",
                "handbook_title": thread.get("current_state", {}).get("handbook_title") or "Uploaded Employee Handbook",
                "handbook_status": "error",
                "handbook_uploaded_file_url": None,
                "handbook_uploaded_filename": None,
                "handbook_blocking_error": blocking_message,
                "handbook_review_locations": active_location_labels,
                "handbook_red_flags": [],
                "handbook_sections": [],
                "handbook_analysis_generated_at": datetime.now(timezone.utc).isoformat(),
                "handbook_error": None,
            },
            diff_summary="Blocked handbook upload audit",
        )
        await doc_svc.add_message(
            thread_id,
            "system",
            f"Handbook upload attempted for {filename}.",
            version_created=result["version"],
        )
        await doc_svc.add_message(
            thread_id,
            "assistant",
            blocking_message,
            version_created=result["version"],
        )
        return await _build_thread_detail_response(thread_id, company_id)

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded handbook file is empty")
    if len(content) > HANDBOOK_UPLOAD_MAX_BYTES:
        raise HTTPException(status_code=400, detail="Handbook file exceeds the 10 MB limit")

    try:
        extracted_text, _page_count = ERDocumentParser().extract_text_from_bytes(content, filename)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        logger.error("Failed to extract handbook upload text for thread %s: %s", thread_id, exc, exc_info=True)
        raise HTTPException(status_code=400, detail="Failed to read the uploaded handbook file") from exc

    if not extracted_text or not extracted_text.strip():
        raise HTTPException(status_code=400, detail="No readable handbook text was found in the uploaded file")

    # Quick relevance check — reject clearly wrong documents before expensive work.
    is_handbook, rejection_reason = await check_handbook_relevance(extracted_text, get_ai_provider().client)
    if not is_handbook:
        blocking_message = rejection_reason or (
            "This document does not appear to be an employee handbook. "
            "Please upload your company's employee handbook and try again."
        )
        result = await doc_svc.apply_update(
            thread_id,
            {
                "handbook_source_type": "upload",
                "handbook_upload_status": "blocked",
                "handbook_title": thread.get("current_state", {}).get("handbook_title") or "Uploaded Employee Handbook",
                "handbook_status": "error",
                "handbook_uploaded_file_url": None,
                "handbook_uploaded_filename": None,
                "handbook_blocking_error": blocking_message,
                "handbook_review_locations": [],
                "handbook_red_flags": [],
                "handbook_green_flags": [],
                "handbook_jurisdiction_summaries": [],
                "handbook_sections": [],
                "handbook_analysis_generated_at": datetime.now(timezone.utc).isoformat(),
                "handbook_error": None,
            },
            diff_summary="Rejected non-handbook upload",
        )
        await doc_svc.add_message(
            thread_id,
            "system",
            f"Uploaded file: {filename}.",
            version_created=result["version"],
        )
        await doc_svc.add_message(
            thread_id,
            "assistant",
            blocking_message,
            version_created=result["version"],
        )
        return await _build_thread_detail_response(thread_id, company_id)

    # --- Happy path: upload to S3, then stream incremental analysis via SSE ---

    storage = get_storage()
    uploaded_file_url = await storage.upload_file(
        content,
        filename,
        prefix=doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "handbooks"),
        content_type=file.content_type,
    )
    await storage.upload_file(
        extracted_text.encode("utf-8"),
        f"{os.path.splitext(filename)[0] or 'handbook'}-extracted.txt",
        prefix=doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "handbook-analysis"),
        content_type="text/plain",
    )

    # Pre-compute handbook metadata that stays constant across quarters.
    parsed_sections = parse_handbook_sections(extracted_text)
    if not parsed_sections:
        raise HTTPException(status_code=400, detail="No readable handbook text found in the uploaded file")

    handbook_title = derive_handbook_title(filename)
    all_states = sorted({str(loc.get("state") or "").strip().upper() for loc in active_locations if loc.get("state")})
    handbook_mode = "single_state" if len(set(all_states)) <= 1 else "multi_state"
    total_location_count = len(active_locations)
    all_location_labels = [_location_label(loc) for loc in active_locations if _location_label(loc)]
    section_previews = [
        {
            "section_key": section.section_key,
            "title": section.title,
            "content": section.content[:500],
            "section_type": section.section_type,
        }
        for section in parsed_sections[:MAX_SECTION_PREVIEWS]
    ]

    # Split locations into up to 4 quarter groups for incremental analysis.
    quarter_size = math.ceil(total_location_count / 4) if total_location_count else 1
    location_quarters: list[list[dict]] = [
        active_locations[i : i + quarter_size]
        for i in range(0, total_location_count, quarter_size)
    ]

    async def event_stream():
        try:
            # Mark thread as analyzing.
            await doc_svc.apply_update(
                thread_id,
                {
                    "handbook_source_type": "upload",
                    "handbook_upload_status": "analyzing",
                    "handbook_analysis_progress": 0,
                    "handbook_title": handbook_title,
                    "handbook_uploaded_file_url": uploaded_file_url,
                    "handbook_uploaded_filename": filename,
                    "handbook_mode": handbook_mode,
                    "handbook_states": all_states,
                    "handbook_sections": section_previews,
                    "handbook_review_locations": all_location_labels,
                    "handbook_blocking_error": None,
                    "handbook_error": None,
                },
                diff_summary="Started handbook analysis",
            )
            yield _sse_data({"type": "handbook_progress", "progress": 0, "status": "analyzing"})

            seen_flag_keys: set[str] = set()
            accumulated_red_flags: list[dict] = []
            accumulated_green_flags: list[dict] = []
            accumulated_coverage: dict[str, dict[str, set[str]]] = {}
            num_quarters = len(location_quarters)

            for q_idx, quarter_locs in enumerate(location_quarters, 1):
                # Fetch requirements for this quarter's locations sequentially
                # to avoid connection pool exhaustion.
                audited_locs: list[AuditedLocation] = []
                for loc in quarter_locs:
                    if not loc.get("id"):
                        continue
                    try:
                        requirements = await get_location_requirements(UUID(str(loc["id"])), company_id)
                    except Exception:
                        logger.error(
                            "Failed to load location requirements for handbook upload audit thread %s location %s",
                            thread_id,
                            loc.get("id"),
                            exc_info=True,
                        )
                        yield _sse_data({"type": "error", "message": "Failed to load synced compliance requirements."})
                        return
                    audited_locs.append(
                        AuditedLocation(
                            id=UUID(str(loc["id"])),
                            label=_location_label(loc),
                            state=str(loc.get("state") or "").strip().upper(),
                            city=str(loc.get("city") or "").strip() or None,
                            requirements=list(requirements),
                        )
                    )

                # Audit this quarter's locations.
                q_red, q_green, q_coverage = _audit_location_group(
                    parsed_sections=parsed_sections,
                    locations_subset=audited_locs,
                    all_states=all_states,
                    total_location_count=total_location_count,
                    seen_flag_keys=seen_flag_keys,
                )

                # Accumulate results.
                accumulated_red_flags.extend(q_red)
                accumulated_green_flags.extend(q_green)
                for loc_key, info in q_coverage.items():
                    if loc_key in accumulated_coverage:
                        accumulated_coverage[loc_key]["covered"] |= info["covered"]
                        accumulated_coverage[loc_key]["total"] |= info["total"]
                        accumulated_coverage[loc_key]["state"] |= info["state"]
                        accumulated_coverage[loc_key]["city"] |= info["city"]
                    else:
                        accumulated_coverage[loc_key] = info

                # Sort and cap red flags.
                sorted_red = sorted(
                    accumulated_red_flags,
                    key=lambda item: (_severity_rank(item["severity"]), item["jurisdiction"], item["section_title"]),
                )
                total_red_count = len(accumulated_red_flags)
                sorted_red = sorted_red[:MAX_RED_FLAGS]

                # Compute running summaries.
                jurisdiction_summaries, strength_score, strength_label = compute_coverage_summaries(accumulated_coverage)
                progress = q_idx / num_quarters

                partial_state = {
                    "handbook_source_type": "upload",
                    "handbook_upload_status": "analyzing",
                    "handbook_analysis_progress": progress,
                    "handbook_title": handbook_title,
                    "handbook_mode": handbook_mode,
                    "handbook_states": all_states,
                    "handbook_uploaded_file_url": uploaded_file_url,
                    "handbook_uploaded_filename": filename,
                    "handbook_blocking_error": None,
                    "handbook_error": None,
                    "handbook_sections": section_previews,
                    "handbook_review_locations": all_location_labels,
                    "handbook_red_flags": sorted_red,
                    "handbook_green_flags": accumulated_green_flags,
                    "handbook_jurisdiction_summaries": jurisdiction_summaries,
                    "handbook_strength_score": strength_score,
                    "handbook_strength_label": strength_label,
                    "handbook_analysis_generated_at": datetime.now(timezone.utc).isoformat(),
                    "handbook_total_red_flag_count": total_red_count,
                }

                await doc_svc.apply_update(
                    thread_id,
                    partial_state,
                    diff_summary=f"Handbook analysis quarter {q_idx}/{num_quarters}",
                )
                yield _sse_data({"type": "handbook_progress", "progress": progress, "partial_state": partial_state})

            # Final: mark as reviewed and add messages.
            final_state = {
                **partial_state,
                "handbook_upload_status": "reviewed",
                "handbook_analysis_progress": 1.0,
                "handbook_status": "ready",
            }
            result = await doc_svc.apply_update(
                thread_id,
                final_state,
                diff_summary=f"Uploaded handbook audit: {filename}",
            )

            summary_message = _build_handbook_upload_summary(
                file_name=filename,
                reviewed_locations=all_location_labels,
                red_flags=sorted_red,
                green_flags=accumulated_green_flags,
                jurisdiction_summaries=jurisdiction_summaries,
            )
            await doc_svc.add_message(
                thread_id,
                "system",
                f"Uploaded handbook file: {filename}.",
                version_created=result["version"],
            )
            await doc_svc.add_message(
                thread_id,
                "assistant",
                summary_message,
                version_created=result["version"],
            )

            detail = await _build_thread_detail_response(thread_id, company_id)
            yield _sse_data({"type": "complete", "data": detail.model_dump(mode="json")})
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error("Handbook upload stream failed for thread %s: %s", thread_id, e, exc_info=True)
            yield _sse_data({"type": "error", "message": "Handbook analysis failed. Please try again."})

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/threads/{thread_id}/images")
async def upload_thread_images(
    thread_id: UUID,
    files: list[UploadFile] = File(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Upload up to 4 images for a workbook thread's presentation (10 MB each)."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    try:
        thread["current_state"] = await doc_svc.ensure_matcha_work_thread_storage_scope(
            thread_id,
            company_id,
            thread.get("current_state") or {},
        )
        existing: list[str] = (thread.get("current_state") or {}).get("images") or []
        if len(existing) + len(files) > 4:
            raise HTTPException(
                status_code=400,
                detail=f"Maximum 4 images allowed. You already have {len(existing)}.",
            )

        uploaded_urls: list[str] = []
        for file in files:
            if not file.content_type or not file.content_type.startswith("image/"):
                raise HTTPException(status_code=400, detail=f"'{file.filename}' is not an image")
            content = await file.read()
            if len(content) > 10 * 1024 * 1024:
                raise HTTPException(status_code=400, detail=f"'{file.filename}' exceeds 10 MB limit")
            url = await get_storage().upload_file(
                content,
                file.filename or "image.jpg",
                prefix=doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "images"),
                content_type=file.content_type,
            )
            uploaded_urls.append(url)

        new_images = existing + uploaded_urls
        await doc_svc.apply_update(thread_id, {"images": new_images}, diff_summary="Added presentation images")
        return {"images": new_images, "uploaded_count": len(uploaded_urls)}

    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to upload images for thread %s: %s", thread_id, e, exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to upload images. Please try again.")


@router.delete("/threads/{thread_id}/images")
async def remove_thread_image(
    thread_id: UUID,
    url: str = Query(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Remove a single image from a workbook thread by URL."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    current_images: list[str] = (thread.get("current_state") or {}).get("images") or []
    if url not in current_images:
        raise HTTPException(status_code=404, detail="Image not found in thread")

    updated_images = [u for u in current_images if u != url]
    await doc_svc.apply_update(thread_id, {"images": updated_images}, diff_summary="Removed presentation image")

    # Best-effort S3 deletion — don't fail the request if this errors
    try:
        await get_storage().delete_file(url)
    except Exception:
        pass

    return {"images": updated_images}


RESUME_EXTRACT_PROMPT = """Extract candidate information from this resume. Return ONLY valid JSON with these fields:
{"name":"...","email":"...","phone":"...","location":"...","current_title":"...","experience_years":0,"skills":["..."],"education":"highest degree - school","certifications":["..."],"summary":"1-2 sentence professional summary","strengths":["top 3 strengths"],"flags":["any concerns or gaps"]}

Resume text:
---
%s
---"""


@router.post("/threads/{thread_id}/resume/upload")
async def upload_thread_resume(
    thread_id: UUID,
    files: list[UploadFile] = File(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Upload one or more resumes, extract structured candidate data, and stream batch insights."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    if thread["status"] in ("finalized", "archived"):
        raise HTTPException(status_code=400, detail=f"Cannot upload to a {thread['status']} thread")

    if current_user.role != "admin":
        await token_budget_service.check_token_budget(company_id)

    # Validate all files upfront
    parsed_files: list[tuple[str, bytes, str]] = []  # (filename, content, content_type)
    for f in files:
        fname = f.filename or "resume"
        ext = os.path.splitext(fname)[1].lower()
        if ext not in RESUME_UPLOAD_EXTENSIONS:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {fname}")
        raw = await f.read()
        if len(raw) > RESUME_UPLOAD_MAX_BYTES:
            raise HTTPException(status_code=400, detail=f"File exceeds 10 MB limit: {fname}")
        parsed_files.append((fname, raw, f.content_type or "application/octet-stream"))

    file_count = len(parsed_files)
    filenames = [f[0] for f in parsed_files]

    # Save a single user message for the batch
    user_content = f"[Resume batch: {file_count} files]\n" + "\n".join(f"- {fn}" for fn in filenames)
    user_msg = await doc_svc.add_message(thread_id, "user", user_content)

    async def event_stream():
        try:
            from ..services.resume_parser import (
                extract_resume_text,
                parse_resume_text,
                ResumeParseError,
            )

            existing_candidates = list((thread.get("current_state") or {}).get("candidates") or [])
            new_candidates = []
            errors = []

            for idx, (fname, raw, ct) in enumerate(parsed_files, 1):
                yield _sse_data({"type": "status", "message": f"Extracting text from {fname} ({idx}/{file_count})..."})

                # Extract text via the shared helper
                try:
                    text = await extract_resume_text(raw, fname)
                except Exception:
                    errors.append(fname)
                    continue

                # Upload raw file to S3 (best-effort)
                resume_url = None
                try:
                    resume_url = await get_storage().upload_file(
                        raw, fname,
                        prefix=doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "resumes"),
                        content_type=ct,
                    )
                except Exception:
                    pass

                # Structured extraction via shared helper
                yield _sse_data({"type": "status", "message": f"Analyzing {fname} ({idx}/{file_count})..."})
                candidate_id = os.urandom(8).hex()
                try:
                    data = await parse_resume_text(text)
                    candidate = {
                        "id": candidate_id,
                        "filename": fname,
                        "resume_url": resume_url,
                        "name": data.get("name"),
                        "email": data.get("email"),
                        "phone": data.get("phone"),
                        "location": data.get("location"),
                        "current_title": data.get("current_title"),
                        "experience_years": data.get("experience_years"),
                        "skills": data.get("skills"),
                        "education": data.get("education"),
                        "certifications": data.get("certifications"),
                        "summary": data.get("summary"),
                        "strengths": data.get("strengths"),
                        "flags": data.get("flags"),
                        "status": "analyzed",
                    }
                except ResumeParseError as e:
                    logger.warning("AI extraction failed for %s: %s", fname, e)
                    candidate = {
                        "id": candidate_id,
                        "filename": fname,
                        "resume_url": resume_url,
                        "status": "error",
                    }

                new_candidates.append(candidate)

                # Update name in status if available
                name = candidate.get("name") or fname
                yield _sse_data({"type": "status", "message": f"Analyzed {name} ({idx}/{file_count})"})

            # Accumulate into state
            all_candidates = existing_candidates + new_candidates
            analyzed = sum(1 for c in all_candidates if c.get("status") == "analyzed")
            result = await doc_svc.apply_update(thread_id, {
                "candidates": all_candidates,
                "batch_status": "ready",
                "total_count": len(all_candidates),
                "analyzed_count": analyzed,
            })
            current_state = result["current_state"]
            current_version = result["version"]

            # Generate batch summary
            yield _sse_data({"type": "status", "message": "Generating batch insights..."})
            summaries = []
            for c in new_candidates:
                if c.get("status") != "analyzed":
                    continue
                summaries.append(
                    f"- {c.get('name', 'Unknown')}: {c.get('current_title', 'N/A')}, "
                    f"{c.get('experience_years', '?')} yrs, {c.get('location', 'N/A')}. "
                    f"{c.get('summary', '')}"
                )

            if summaries:
                batch_prompt = (
                    f"I just uploaded {len(new_candidates)} resumes. Here are the candidates:\n\n"
                    + "\n".join(summaries)
                    + "\n\nProvide a brief batch overview:\n"
                    "1. Quick summary of the candidate pool (experience range, common skills, locations)\n"
                    "2. Top standout candidates and why\n"
                    "3. Any common gaps or concerns\n"
                    "Keep it concise — 2-3 short paragraphs max."
                )
                ai_provider = get_ai_provider()
                profile = await doc_svc.get_company_profile_for_ai(company_id)
                ctx = _build_company_context(profile)
                ai_resp = await ai_provider.generate(
                    [{"role": "user", "content": batch_prompt}],
                    current_state,
                    company_context=ctx,
                )
                batch_reply = ai_resp.assistant_reply
            else:
                batch_reply = f"Uploaded {file_count} files."
                if errors:
                    batch_reply += f" Could not process: {', '.join(errors)}."

            assistant_msg = await doc_svc.add_message(thread_id, "assistant", batch_reply)

            response = SendMessageResponse(
                user_message=_row_to_message(user_msg),
                assistant_message=_row_to_message(assistant_msg),
                current_state=current_state,
                version=current_version,
                task_type=_infer_skill_from_state(current_state),
                pdf_url=None,
                token_usage=None,
            )
            yield _sse_data({"type": "complete", "data": response.model_dump(mode="json")})
        except Exception as e:
            logger.error("Resume batch failed for thread %s: %s", thread_id, e, exc_info=True)
            yield _sse_data({"type": "error", "message": "Failed to process resumes. Please try again."})
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/threads/{thread_id}/resume/send-interviews")
async def send_resume_batch_interviews(
    thread_id: UUID,
    body: SendInterviewsRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Create screening interviews for selected candidates and send invite emails."""
    import secrets as _secrets
    from app.core.services.email import EmailService

    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    current_state = thread.get("current_state") or {}
    candidates = current_state.get("candidates") or []
    if not candidates:
        raise HTTPException(status_code=400, detail="No candidates in this batch")

    async with get_connection() as conn:
        company_row = await conn.fetchrow("SELECT name FROM companies WHERE id = $1", company_id)
    company_name = company_row["name"] if company_row else "the company"

    position_title = body.position_title or "Open Position"
    email_svc = EmailService()
    settings = get_settings()

    sent = []
    failed = []
    updated_candidates = list(candidates)

    for cid in body.candidate_ids:
        candidate = None
        candidate_idx = None
        for idx, c in enumerate(updated_candidates):
            if c.get("id") == cid:
                candidate = c
                candidate_idx = idx
                break

        if candidate is None:
            failed.append({"id": cid, "error": "Candidate not found in batch"})
            continue

        email = candidate.get("email")
        name = candidate.get("name") or candidate.get("filename", "Candidate")
        if not email:
            failed.append({"id": cid, "error": f"No email for {name}"})
            continue

        try:
            invite_token = _secrets.token_urlsafe(32)
            interview_data = json.dumps({
                "invite_token": invite_token,
                "candidate_name": name,
                "candidate_email": email,
                "position_title": position_title,
                "company_name": company_name,
                "resume_batch_thread_id": str(thread_id),
                "resume_batch_candidate_id": cid,
            })

            async with get_connection() as interview_conn:
                interview_row = await interview_conn.fetchrow(
                    """
                    INSERT INTO interviews (id, company_id, interviewer_name, interview_type, raw_culture_data, status, created_at)
                    VALUES (gen_random_uuid(), $1, $2, 'screening', $3, 'pending', NOW())
                    RETURNING id
                    """,
                    company_id,
                    name,
                    interview_data,
                )
            interview_id = interview_row["id"]

            invite_url = f"{settings.app_base_url}/candidate-interview/{invite_token}"
            email_sent = await email_svc.send_candidate_interview_invite_email(
                to_email=email,
                to_name=name,
                company_name=company_name,
                position_title=position_title,
                invite_url=invite_url,
                custom_message=body.custom_message,
            )

            updated_candidates[candidate_idx] = {
                **candidate,
                "status": "interview_sent",
                "interview_id": str(interview_id),
            }

            sent.append({
                "id": cid,
                "name": name,
                "email": email,
                "interview_id": str(interview_id),
                "email_sent": email_sent,
            })
        except Exception as e:
            logger.error("Failed to create interview for candidate %s: %s", cid, e, exc_info=True)
            failed.append({"id": cid, "error": str(e)})

    if sent:
        await doc_svc.apply_update(thread_id, {"candidates": updated_candidates})

    return {"sent": sent, "failed": failed}


@router.post("/threads/{thread_id}/resume/sync-interviews")
async def sync_resume_batch_interviews(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Sync interview statuses back into the resume batch candidates."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    current_state = thread.get("current_state") or {}
    candidates = current_state.get("candidates") or []

    # Collect interview IDs from candidates
    interview_ids = [c.get("interview_id") for c in candidates if c.get("interview_id")]
    if not interview_ids:
        return {"updated": 0}

    # Fetch interview statuses + screening analysis
    async with get_connection() as conn:
        rows = await conn.fetch(
            """
            SELECT id, status, screening_analysis
            FROM interviews
            WHERE id = ANY($1::uuid[])
            """,
            interview_ids,
        )

    interview_map = {str(r["id"]): r for r in rows}
    updated = 0
    updated_candidates = list(candidates)

    for idx, c in enumerate(updated_candidates):
        iid = c.get("interview_id")
        if not iid or iid not in interview_map:
            continue
        row = interview_map[iid]
        new_status = c.get("status")
        interview_status = row["status"]

        # Treat analyzing as completed for UI purposes — interview is over,
        # only the AI analysis is still running. The review modal handles the
        # "no analysis yet" state gracefully.
        if interview_status in ("completed", "analyzed", "analyzing") and c.get("status") != "interview_completed":
            new_status = "interview_completed"
        elif interview_status == "in_progress" and c.get("status") == "interview_sent":
            new_status = "interview_in_progress"

        # Extract score from screening_analysis if available
        score = None
        summary = None
        analysis = row.get("screening_analysis")
        if analysis:
            if isinstance(analysis, str):
                analysis = json.loads(analysis)
            score = analysis.get("overall_score") or analysis.get("score")
            summary = analysis.get("summary") or analysis.get("overall_assessment")

        if new_status != c.get("status") or score or summary:
            updated_candidates[idx] = {
                **c,
                "status": new_status,
                "interview_status": interview_status,
                "interview_score": score,
                "interview_summary": summary,
            }
            updated += 1

    if updated > 0:
        await doc_svc.apply_update(thread_id, {"candidates": updated_candidates})

    return {"updated": updated}


async def _convert_svgs_to_images(content: str, company_id, project_id) -> tuple[str, list[dict]]:
    """Convert inline SVGs to uploaded <img> tags. Returns (html, diagram_data_list)."""
    import re as _re

    storage = get_storage()
    prefix = f"matcha-work/{company_id}/{project_id}/diagrams"
    counter = 0
    diagrams: list[dict] = []

    async def _upload_svg(svg_str: str) -> str:
        nonlocal counter
        counter += 1
        svg_bytes = svg_str.encode("utf-8")
        url = await storage.upload_file(svg_bytes, f"diagram-{counter}.svg", prefix=prefix, content_type="image/svg+xml")
        diagrams.append({"svg_source": svg_str, "storage_url": url, "created_from": "ai_generation"})
        return f'<img src="{url}" alt="Diagram" data-diagram-index="{len(diagrams) - 1}" style="max-width:100%;margin:8px 0;" />'

    result = content

    # 1. Fenced code blocks: ```svg ... ``` or ```xml ... ``` containing <svg
    fenced_pattern = _re.compile(r'```(?:svg|xml)\s*\n([\s\S]*?)\n```', _re.IGNORECASE)
    fenced_matches = list(fenced_pattern.finditer(result))
    for match in reversed(fenced_matches):
        inner = match.group(1).strip()
        if '<svg' in inner.lower():
            img_tag = await _upload_svg(inner)
            result = result[:match.start()] + img_tag + result[match.end():]

    # 2. Inline <svg>...</svg> blocks
    svg_pattern = _re.compile(r'<svg[\s\S]*?</svg>', _re.IGNORECASE)
    svg_matches = list(svg_pattern.finditer(result))
    for match in reversed(svg_matches):
        img_tag = await _upload_svg(match.group(0))
        result = result[:match.start()] + img_tag + result[match.end():]

    # 3. Markdown image references to data URIs: ![...](data:image/svg+xml;base64,...)
    data_uri_pattern = _re.compile(r'!\[([^\]]*)\]\(data:image/svg\+xml;base64,([A-Za-z0-9+/=]+)\)')
    data_matches = list(data_uri_pattern.finditer(result))
    for match in reversed(data_matches):
        import base64
        try:
            svg_str = base64.b64decode(match.group(2)).decode("utf-8")
            img_tag = await _upload_svg(svg_str)
            result = result[:match.start()] + img_tag + result[match.end():]
        except Exception:
            pass

    return result, diagrams


def _strip_markdown(text: str) -> str:
    """Strip common markdown syntax to produce clean plain text for project sections."""
    import re as _re
    t = text
    t = _re.sub(r'\*\*(.+?)\*\*', r'\1', t)       # **bold**
    t = _re.sub(r'__(.+?)__', r'\1', t)             # __bold__
    t = _re.sub(r'(?<!\w)\*(.+?)\*(?!\w)', r'\1', t)  # *italic*
    t = _re.sub(r'(?<!\w)_(.+?)_(?!\w)', r'\1', t)    # _italic_
    t = _re.sub(r'^#{1,6}\s+', '', t, flags=_re.MULTILINE)  # ## headings
    t = _re.sub(r'`(.+?)`', r'\1', t)               # `code`
    t = _re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t) # [text](url) → text
    t = _re.sub(r'^[\s]*[-*]\s+', '• ', t, flags=_re.MULTILINE)  # - list → • list
    t = _re.sub(r'^---+$', '', t, flags=_re.MULTILINE)  # ---
    t = _re.sub(r'^>\s*', '', t, flags=_re.MULTILINE)   # > blockquote
    return t.strip()


# ── Project (top-level) endpoints ──


async def _verify_project_access(project_id: UUID, current_user: CurrentUser) -> tuple[dict, str]:
    """Check project access. For admins, uses collaborator table. Returns (project, role)."""
    from ..services import project_service as proj_svc
    if current_user.role == "admin":
        result = await proj_svc.get_project_as_collaborator(project_id, current_user.id)
        if result:
            return result
        raise HTTPException(status_code=404, detail="Project not found")
    company_id = await get_client_company_id(current_user)
    if not company_id:
        raise HTTPException(status_code=404, detail="Project not found")
    project = await proj_svc.get_project(project_id, company_id, user_id=current_user.id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not project.get("collaborator_role"):
        project["collaborator_role"] = "owner"
    return project, project["collaborator_role"]


@router.get("/projects")
async def list_projects_endpoint(
    status: Optional[str] = Query(None, pattern="^(active|archived)$"),
    hiring_client_id: Optional[UUID] = Query(None),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List all projects for the current user."""
    from ..services import project_service as proj_svc
    if current_user.role == "admin":
        return await proj_svc.list_projects(None, status, user_id=current_user.id)
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        return []
    return await proj_svc.list_projects(company_id, status, hiring_client_id=hiring_client_id)


@router.post("/projects", status_code=201)
async def create_project_endpoint(
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Create a new project with an auto-created first chat."""
    from ..services import project_service as proj_svc
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=400, detail="No company associated")
    title = body.get("title", "Untitled Project")
    project_type = body.get("project_type", "general")
    hiring_client_id_raw = body.get("hiring_client_id")
    hiring_client_id = UUID(hiring_client_id_raw) if hiring_client_id_raw else None

    extra_data: Optional[dict] = None
    if project_type == "consultation":
        extra_data = {
            "client": body.get("client") or {},
            "engagement": body.get("engagement") or {},
            "tags": body.get("tags") or [],
        }
        # Consultation title defaults to the client name when caller didn't set one.
        if body.get("title") in (None, "", "Untitled Project"):
            client_name = (extra_data["client"] or {}).get("name")
            if client_name:
                title = client_name
    elif project_type == "blog":
        blog = body.get("blog") or {}
        extra_data = {
            "title": title,
            "audience": blog.get("audience"),
            "tone": blog.get("tone"),
            "tags": blog.get("tags") or [],
            "author": blog.get("author") or {},
        }

    try:
        return await proj_svc.create_project(
            company_id, current_user.id, title, project_type,
            hiring_client_id=hiring_client_id, extra_data=extra_data,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


# ── Recruiting clients (hiring clients a recruiter works for) ──


@router.get("/recruiting-clients")
async def list_recruiting_clients_endpoint(
    include_archived: bool = Query(False),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import recruiting_client_service as rc_svc
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        return []
    return await rc_svc.list_clients(company_id, include_archived=include_archived)


@router.post("/recruiting-clients", status_code=201)
async def create_recruiting_client_endpoint(
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import recruiting_client_service as rc_svc
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=400, detail="No company associated")
    name = (body.get("name") or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="Name required")
    return await rc_svc.create_client(
        company_id,
        current_user.id,
        name=name,
        website=body.get("website"),
        logo_url=body.get("logo_url"),
        notes=body.get("notes"),
    )


@router.get("/recruiting-clients/{client_id}")
async def get_recruiting_client_endpoint(
    client_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import recruiting_client_service as rc_svc
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Not found")
    rc = await rc_svc.get_client(client_id, company_id)
    if not rc:
        raise HTTPException(status_code=404, detail="Not found")
    return rc


@router.patch("/recruiting-clients/{client_id}")
async def update_recruiting_client_endpoint(
    client_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import recruiting_client_service as rc_svc
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Not found")
    rc = await rc_svc.update_client(client_id, company_id, body)
    if not rc:
        raise HTTPException(status_code=404, detail="Not found")
    return rc


@router.post("/recruiting-clients/{client_id}/archive")
async def archive_recruiting_client_endpoint(
    client_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import recruiting_client_service as rc_svc
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Not found")
    ok = await rc_svc.archive_client(client_id, company_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Not found")
    return {"status": "archived"}


@router.post("/recruiting-clients/{client_id}/unarchive")
async def unarchive_recruiting_client_endpoint(
    client_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import recruiting_client_service as rc_svc
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Not found")
    ok = await rc_svc.unarchive_client(client_id, company_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Not found")
    return {"status": "active"}


@router.get("/projects/{project_id}")
async def get_project_endpoint(
    project_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Get a project with its chat list."""
    project, _role = await _verify_project_access(project_id, current_user)
    return project


@router.patch("/projects/{project_id}")
async def update_project_endpoint(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Update project title, pin, status, or hiring client."""
    from ..services import project_service as proj_svc
    from ..services import recruiting_client_service as rc_svc
    await _verify_project_access(project_id, current_user)
    if "hiring_client_id" in body and body["hiring_client_id"] is not None:
        company_id = await get_client_company_id(current_user)
        try:
            body["hiring_client_id"] = UUID(body["hiring_client_id"])
        except (ValueError, TypeError):
            raise HTTPException(status_code=400, detail="Invalid hiring_client_id")
        if company_id is None or not await rc_svc.get_client(body["hiring_client_id"], company_id):
            raise HTTPException(status_code=400, detail="Hiring client does not belong to this workspace")
    return await proj_svc.update_project(project_id, body)


@router.delete("/projects/{project_id}")
async def archive_project_endpoint(
    project_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Archive a project."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    await proj_svc.archive_project(project_id)
    return {"status": "archived"}


# ── Consultation endpoints (project_type == 'consultation') ──


@router.patch("/projects/{project_id}/consultation")
async def patch_consultation_endpoint(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Partial update of a consultation's client/engagement/stage/tags/custom_fields."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    try:
        return await proj_svc.patch_consultation(project_id, body)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/projects/{project_id}/sessions")
async def add_session_endpoint(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Append a real-world session entry (time + notes + billable)."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    try:
        return await proj_svc.append_session(project_id, body)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.patch("/projects/{project_id}/sessions/{session_id}")
async def patch_session_endpoint(
    project_id: UUID,
    session_id: str,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    try:
        return await proj_svc.update_session(project_id, session_id, body)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.delete("/projects/{project_id}/sessions/{session_id}")
async def delete_session_endpoint(
    project_id: UUID,
    session_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.delete_session(project_id, session_id)


@router.post("/projects/{project_id}/action-items")
async def add_action_item_endpoint(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    text = (body.get("text") or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="text is required")
    return await proj_svc.append_action_item(
        project_id,
        text=text,
        source_thread_id=body.get("source_thread_id"),
        pending_confirmation=bool(body.get("pending_confirmation", False)),
    )


@router.patch("/projects/{project_id}/action-items/{item_id}")
async def patch_action_item_endpoint(
    project_id: UUID,
    item_id: str,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    try:
        return await proj_svc.patch_action_item(project_id, item_id, body)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.delete("/projects/{project_id}/action-items/{item_id}")
async def delete_action_item_endpoint(
    project_id: UUID,
    item_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.delete_action_item(project_id, item_id)


# ── Blog endpoints (project_type == 'blog') ──


@router.patch("/projects/{project_id}/blog")
async def patch_blog_endpoint(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    try:
        return await proj_svc.patch_blog(project_id, body)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/projects/{project_id}/blog/status")
async def transition_blog_status_endpoint(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    to = (body or {}).get("to")
    if not to:
        raise HTTPException(status_code=400, detail="Missing 'to'")
    try:
        return await proj_svc.transition_blog_status(project_id, to)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/projects/{project_id}/sections")
async def add_project_section_endpoint(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Add a section to the project."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    company_id = await get_client_company_id(current_user)
    raw_content = body.get("content", "")

    # Convert inline SVGs to uploaded images so TipTap can render them
    raw_content, diagrams = await _convert_svgs_to_images(raw_content, company_id, project_id)

    section_data = {**body, "content": raw_content}
    if diagrams:
        section_data["diagram_data"] = diagrams
    return await proj_svc.add_section(project_id, section_data)


@router.put("/projects/{project_id}/sections/reorder")
async def reorder_project_sections_endpoint(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Reorder project sections."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.reorder_sections(project_id, body.get("section_ids", []))


@router.put("/projects/{project_id}/sections/{section_id}")
async def update_project_section_endpoint(
    project_id: UUID,
    section_id: str,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Update a project section."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.update_section(project_id, section_id, body)


@router.delete("/projects/{project_id}/sections/{section_id}")
async def delete_project_section_endpoint(
    project_id: UUID,
    section_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Delete a project section."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.delete_section(project_id, section_id)


# ── Project file attachment endpoints ──

ALLOWED_PROJECT_FILE_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".csv", ".xlsx", ".xls",
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg",
    ".pptx", ".md",
}
PROJECT_FILE_MAX_BYTES = 10 * 1024 * 1024  # 10 MB


@router.post("/projects/{project_id}/files")
async def upload_project_file(
    project_id: UUID,
    file: UploadFile = File(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Upload a file attachment to a project."""
    from ..services import project_file_service

    project, _role = await _verify_project_access(project_id, current_user)
    company_id = project.get("company_id")
    if not company_id:
        company_id = await get_client_company_id(current_user)

    fname = file.filename or "file"
    ext = os.path.splitext(fname)[1].lower()
    if ext not in ALLOWED_PROJECT_FILE_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    content = await file.read()
    if len(content) > PROJECT_FILE_MAX_BYTES:
        raise HTTPException(status_code=400, detail="File exceeds 10 MB limit")

    storage_url = await get_storage().upload_file(
        content, fname,
        prefix=f"matcha-work/{company_id}/{project_id}/files",
        content_type=file.content_type,
    )

    record = await project_file_service.add_project_file(
        project_id=project_id,
        uploaded_by=current_user.id,
        filename=fname,
        storage_url=storage_url,
        content_type=file.content_type,
        file_size=len(content),
    )
    return record


@router.get("/projects/{project_id}/files")
async def list_project_files_endpoint(
    project_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List file attachments for a project."""
    from ..services import project_file_service

    await _verify_project_access(project_id, current_user)
    return await project_file_service.list_project_files(project_id)


@router.delete("/projects/{project_id}/files/{file_id}")
async def delete_project_file_endpoint(
    project_id: UUID,
    file_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Delete a file attachment from a project."""
    from ..services import project_file_service

    await _verify_project_access(project_id, current_user)

    record = await project_file_service.get_project_file(file_id, project_id)
    if not record:
        raise HTTPException(status_code=404, detail="File not found")

    # Remove from storage
    try:
        await get_storage().delete_file(record["storage_url"])
    except Exception:
        pass

    await project_file_service.delete_project_file(file_id, project_id)
    return {"deleted": True}


# ── Research task endpoints ──


@router.post("/projects/{project_id}/research-tasks")
async def create_research_task(
    project_id: UUID,
    body: dict = Body(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Create a new research task in a project."""
    from ..services import project_service as proj_svc
    import uuid as _uuid

    await _verify_project_access(project_id, current_user)

    task = {
        "id": str(_uuid.uuid4()),
        "name": body.get("name", "Untitled Research"),
        "instructions": body.get("instructions", ""),
        "inputs": [],
        "results": [],
    }

    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}
            tasks = data.get("research_tasks", [])
            tasks.append(task)
            data["research_tasks"] = tasks
            await conn.execute(
                "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                json.dumps(data), project_id,
            )

    return task


@router.put("/projects/{project_id}/research-tasks/{task_id}")
async def update_research_task(
    project_id: UUID,
    task_id: str,
    body: dict = Body(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Update a research task definition."""
    await _verify_project_access(project_id, current_user)

    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}

            for task in data.get("research_tasks", []):
                if task["id"] == task_id:
                    if "name" in body:
                        task["name"] = body["name"]
                    if "instructions" in body:
                        task["instructions"] = body["instructions"]
                    await conn.execute(
                        "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                        json.dumps(data), project_id,
                    )
                    return task

    raise HTTPException(status_code=404, detail="Research task not found")


@router.delete("/projects/{project_id}/research-tasks/{task_id}")
async def delete_research_task(
    project_id: UUID,
    task_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Delete a research task and all its results."""
    await _verify_project_access(project_id, current_user)

    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}
            data["research_tasks"] = [t for t in data.get("research_tasks", []) if t["id"] != task_id]
            await conn.execute(
                "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                json.dumps(data), project_id,
            )

    return {"deleted": True}


@router.post("/projects/{project_id}/research-tasks/{task_id}/inputs")
async def add_research_inputs(
    project_id: UUID,
    task_id: str,
    body: dict = Body(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Add URLs to a research task."""
    import uuid as _uuid

    await _verify_project_access(project_id, current_user)
    urls = body.get("urls", [])
    if not urls:
        raise HTTPException(status_code=400, detail="No URLs provided")

    new_inputs = []
    for url in urls:
        url = url.strip()
        if not url or not (url.startswith("http://") or url.startswith("https://")):
            continue
        new_inputs.append({
            "id": str(_uuid.uuid4()),
            "url": url,
            "status": "pending",
            "queued_at": datetime.now(timezone.utc).isoformat(),
        })

    if not new_inputs:
        raise HTTPException(status_code=400, detail="No valid URLs provided")

    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}

            for task in data.get("research_tasks", []):
                if task["id"] == task_id:
                    task.setdefault("inputs", []).extend(new_inputs)
                    await conn.execute(
                        "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                        json.dumps(data), project_id,
                    )
                    return {"added": len(new_inputs), "inputs": new_inputs}

    raise HTTPException(status_code=404, detail="Research task not found")


@router.delete("/projects/{project_id}/research-tasks/{task_id}/inputs/{input_id}")
async def delete_research_input(
    project_id: UUID,
    task_id: str,
    input_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Remove a URL from a research task."""
    await _verify_project_access(project_id, current_user)

    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}

            for task in data.get("research_tasks", []):
                if task["id"] == task_id:
                    task["inputs"] = [i for i in task.get("inputs", []) if i["id"] != input_id]
                    task["results"] = [r for r in task.get("results", []) if r.get("input_id") != input_id]
                    await conn.execute(
                        "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                        json.dumps(data), project_id,
                    )
                    return {"deleted": True}

    raise HTTPException(status_code=404, detail="Research task not found")


@router.post("/projects/{project_id}/research-tasks/{task_id}/run")
async def run_research_task(
    project_id: UUID,
    task_id: str,
    capture_screenshot: bool = Query(False),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Run all pending research inputs sequentially with SSE status streaming."""
    from ..services.research_browse_service import run_research_for_input
    from starlette.responses import StreamingResponse

    project, _role = await _verify_project_access(project_id, current_user)
    company_id = str(project.get("company_id") or await get_client_company_id(current_user))

    # Collect pending inputs
    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}

            pending_inputs = []
            instructions = ""
            for task in data.get("research_tasks", []):
                if task["id"] == task_id:
                    instructions = task.get("instructions", "")
                    if not instructions:
                        raise HTTPException(status_code=400, detail="Task has no instructions")
                    for inp in task.get("inputs", []):
                        if inp["status"] in ("pending", "error"):
                            inp["status"] = "running"
                            inp.pop("error", None)
                            pending_inputs.append({"id": inp["id"], "url": inp["url"]})

                    await conn.execute(
                        "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                        json.dumps(data), project_id,
                    )
                    break

    if not pending_inputs:
        return {"queued": 0}

    async def event_stream():
        for i, inp in enumerate(pending_inputs):
            yield _sse_data({"type": "status", "input_id": inp["id"], "url": inp["url"],
                             "message": f"Starting research ({i + 1}/{len(pending_inputs)}): {inp['url']}"})

            async def on_status(msg):
                pass  # will be replaced per-input below

            status_queue: asyncio.Queue = asyncio.Queue()

            async def stream_status(msg: str):
                await status_queue.put(msg)

            # Run browse in background, stream statuses
            browse_task = asyncio.create_task(
                run_research_for_input(
                    project_id, task_id, inp["id"], inp["url"], instructions,
                    on_status=stream_status,
                    capture_screenshot=capture_screenshot, company_id=company_id,
                )
            )

            while not browse_task.done():
                try:
                    msg = await asyncio.wait_for(status_queue.get(), timeout=1.0)
                    yield _sse_data({"type": "status", "input_id": inp["id"], "message": msg})
                except asyncio.TimeoutError:
                    pass

            # Drain remaining messages
            while not status_queue.empty():
                msg = status_queue.get_nowait()
                yield _sse_data({"type": "status", "input_id": inp["id"], "message": msg})

            result = browse_task.result()
            yield _sse_data({
                "type": "complete" if not result.get("error") else "error",
                "input_id": inp["id"],
                "url": inp["url"],
                "findings": result.get("findings", {}),
                "summary": result.get("summary", ""),
                "error": result.get("error"),
            })

        yield _sse_data({"type": "done", "message": f"Finished researching {len(pending_inputs)} URL(s)"})

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/projects/{project_id}/research-tasks/{task_id}/inputs/{input_id}/follow-up")
async def follow_up_research_input(
    project_id: UUID,
    task_id: str,
    input_id: str,
    body: dict = Body(...),
    capture_screenshot: bool = Query(False),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Re-research a URL with additional instructions, building on previous findings."""
    from ..services.research_browse_service import run_research_for_input
    from starlette.responses import StreamingResponse

    project, _role = await _verify_project_access(project_id, current_user)
    company_id = str(project.get("company_id") or await get_client_company_id(current_user))

    follow_up = body.get("follow_up", "").strip()
    if not follow_up:
        raise HTTPException(status_code=400, detail="follow_up is required")

    follow_url = ""
    combined_instructions = ""

    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}

            for task in data.get("research_tasks", []):
                if task["id"] == task_id:
                    base_instructions = task.get("instructions", "")

                    # Find previous findings for context
                    prev_findings = {}
                    for r in task.get("results", []):
                        if r.get("input_id") == input_id:
                            prev_findings = r.get("findings", {})
                            break

                    # Build combined instructions with previous context
                    combined_instructions = base_instructions
                    if prev_findings:
                        combined_instructions += f"\n\nPREVIOUS FINDINGS (already gathered):\n{json.dumps(prev_findings, indent=2)}"
                    combined_instructions += f"\n\nADDITIONAL REQUEST:\n{follow_up}"

                    for inp in task.get("inputs", []):
                        if inp["id"] == input_id:
                            inp["status"] = "running"
                            inp.pop("error", None)
                            inp.pop("completed_at", None)
                            follow_url = inp["url"]
                            # Keep old results — new ones will merge
                            break

                    await conn.execute(
                        "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                        json.dumps(data), project_id,
                    )
                    break

    if not follow_url:
        raise HTTPException(status_code=404, detail="Input not found")

    async def event_stream():
        status_queue: asyncio.Queue = asyncio.Queue()

        async def stream_status(msg: str):
            await status_queue.put(msg)

        browse_task = asyncio.create_task(
            run_research_for_input(
                project_id, task_id, input_id, follow_url, combined_instructions,
                on_status=stream_status,
                capture_screenshot=capture_screenshot, company_id=company_id,
            )
        )

        while not browse_task.done():
            try:
                msg = await asyncio.wait_for(status_queue.get(), timeout=1.0)
                yield _sse_data({"type": "status", "input_id": input_id, "message": msg})
            except asyncio.TimeoutError:
                pass

        while not status_queue.empty():
            msg = status_queue.get_nowait()
            yield _sse_data({"type": "status", "input_id": input_id, "message": msg})

        result = browse_task.result()
        yield _sse_data({
            "type": "complete" if not result.get("error") else "error",
            "input_id": input_id, "url": follow_url,
            "findings": result.get("findings", {}),
            "summary": result.get("summary", ""),
            "error": result.get("error"),
        })
        yield _sse_data({"type": "done"})

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/projects/{project_id}/research-tasks/{task_id}/inputs/{input_id}/retry")
async def retry_research_input(
    project_id: UUID,
    task_id: str,
    input_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Retry a failed research input with SSE streaming."""
    from ..services.research_browse_service import run_research_for_input
    from starlette.responses import StreamingResponse

    await _verify_project_access(project_id, current_user)

    retry_url = ""
    retry_instructions = ""

    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}

            for task in data.get("research_tasks", []):
                if task["id"] == task_id:
                    retry_instructions = task.get("instructions", "")
                    for inp in task.get("inputs", []):
                        if inp["id"] == input_id:
                            inp["status"] = "running"
                            inp.pop("error", None)
                            inp.pop("completed_at", None)
                            retry_url = inp["url"]
                            task["results"] = [r for r in task.get("results", []) if r.get("input_id") != input_id]

                            await conn.execute(
                                "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                                json.dumps(data), project_id,
                            )
                            break
                    break

    if not retry_url:
        raise HTTPException(status_code=404, detail="Input not found")

    async def event_stream():
        status_queue: asyncio.Queue = asyncio.Queue()

        async def stream_status(msg: str):
            await status_queue.put(msg)

        browse_task = asyncio.create_task(
            run_research_for_input(
                project_id, task_id, input_id, retry_url, retry_instructions,
                on_status=stream_status,
            )
        )

        while not browse_task.done():
            try:
                msg = await asyncio.wait_for(status_queue.get(), timeout=1.0)
                yield _sse_data({"type": "status", "input_id": input_id, "message": msg})
            except asyncio.TimeoutError:
                pass

        while not status_queue.empty():
            msg = status_queue.get_nowait()
            yield _sse_data({"type": "status", "input_id": input_id, "message": msg})

        result = browse_task.result()
        yield _sse_data({
            "type": "complete" if not result.get("error") else "error",
            "input_id": input_id, "url": retry_url,
            "findings": result.get("findings", {}),
            "summary": result.get("summary", ""),
            "error": result.get("error"),
        })
        yield _sse_data({"type": "done"})

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/projects/{project_id}/research-tasks/{task_id}/stop")
async def stop_research_task(
    project_id: UUID,
    task_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Reset all running inputs back to pending (cancel in-flight research)."""
    await _verify_project_access(project_id, current_user)

    async with get_connection() as conn:
        async with conn.transaction():
            row = await conn.fetchrow(
                "SELECT project_data FROM mw_projects WHERE id = $1 FOR UPDATE", project_id,
            )
            data = row["project_data"] if row else {}
            if isinstance(data, str):
                data = json.loads(data)
            data = data or {}

            reset_count = 0
            for task in data.get("research_tasks", []):
                if task["id"] == task_id:
                    for inp in task.get("inputs", []):
                        if inp["status"] == "running":
                            inp["status"] = "pending"
                            reset_count += 1
                    break

            if reset_count:
                await conn.execute(
                    "UPDATE mw_projects SET project_data = $1::jsonb, updated_at = NOW() WHERE id = $2",
                    json.dumps(data), project_id,
                )

    return {"reset": reset_count}


# ── Diagram editing endpoints ──


@router.post("/projects/{project_id}/sections/{section_id}/edit-diagram")
async def edit_diagram_ai(
    project_id: UUID,
    section_id: str,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Use AI to modify a diagram based on a natural language instruction."""
    from ..services import project_service as proj_svc

    await _verify_project_access(project_id, current_user)
    company_id = await get_client_company_id(current_user)
    instruction = body.get("instruction", "").strip()
    if not instruction:
        raise HTTPException(status_code=400, detail="Instruction is required")

    project = await proj_svc.get_project(project_id, company_id)
    section = next((s for s in project.get("sections", []) if s.get("id") == section_id), None)
    if not section:
        raise HTTPException(status_code=404, detail="Section not found")

    diagram_data = section.get("diagram_data")
    if not diagram_data or not isinstance(diagram_data, list) or len(diagram_data) == 0:
        raise HTTPException(status_code=400, detail="No diagram data found in this section")

    svg_source = diagram_data[0].get("svg_source", "")
    if not svg_source:
        raise HTTPException(status_code=400, detail="No SVG source available for editing")

    # Optional region selection (percentages of image dimensions)
    region = body.get("region")  # { x, y, width, height } in %

    # Call Gemini to modify the SVG
    import google.genai as genai
    import re as _re_vb
    from ...config import get_settings
    settings = get_settings()

    client = genai.Client(api_key=settings.gemini_api_key)

    # Build prompt — region-constrained or full edit
    if region and isinstance(region, dict):
        # Parse viewBox to convert % region to absolute SVG coordinates
        vb_match = _re_vb.search(r'viewBox=["\']([^"\']+)["\']', svg_source)
        if vb_match:
            vb_parts = vb_match.group(1).split()
            vb_x, vb_y, vb_w, vb_h = float(vb_parts[0]), float(vb_parts[1]), float(vb_parts[2]), float(vb_parts[3])
        else:
            w_match = _re_vb.search(r'\bwidth=["\'](\d+)', svg_source)
            h_match = _re_vb.search(r'\bheight=["\'](\d+)', svg_source)
            vb_x, vb_y = 0.0, 0.0
            vb_w = float(w_match.group(1)) if w_match else 480.0
            vb_h = float(h_match.group(1)) if h_match else 300.0

        rx, ry, rw, rh = float(region["x"]), float(region["y"]), float(region["width"]), float(region["height"])
        abs_x1 = vb_x + (rx / 100) * vb_w
        abs_y1 = vb_y + (ry / 100) * vb_h
        abs_x2 = vb_x + ((rx + rw) / 100) * vb_w
        abs_y2 = vb_y + ((ry + rh) / 100) * vb_h

        prompt = f"""You are an SVG diagram editor. Here is an SVG diagram:

{svg_source}

IMPORTANT CONSTRAINT — REGION-ONLY EDIT:
The user has selected a specific region of the diagram for editing.
The selected region in SVG coordinates is: top-left ({abs_x1:.1f}, {abs_y1:.1f}) to bottom-right ({abs_x2:.1f}, {abs_y2:.1f}).

Rules:
1. ONLY modify SVG elements that are within or overlap this bounding box.
2. DO NOT move, resize, restyle, recolor, or delete ANY element outside this region.
3. DO NOT change the viewBox, overall SVG dimensions, or add new elements outside this region.
4. Elements partially inside the region may be modified, but preserve their parts outside the region as much as possible.
5. The rest of the SVG MUST remain EXACTLY as-is, character for character.

User's instruction (applies ONLY to the selected region):
{instruction}

Return ONLY the modified SVG code, nothing else. No markdown fences, no explanation. Just the raw <svg>...</svg> content."""
    else:
        prompt = f"""You are an SVG diagram editor. Here is an SVG diagram:

{svg_source}

Modify this SVG according to the following instruction:
{instruction}

Return ONLY the modified SVG code, nothing else. No markdown fences, no explanation. Just the raw <svg>...</svg> content."""

    try:
        response = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt,
        )
        new_svg = response.text.strip()
        # Clean up any markdown fences
        if new_svg.startswith("```"):
            new_svg = new_svg.split("\n", 1)[1] if "\n" in new_svg else new_svg
        if new_svg.endswith("```"):
            new_svg = new_svg.rsplit("```", 1)[0].strip()
        if not new_svg.startswith("<svg"):
            raise HTTPException(status_code=500, detail="AI did not return valid SVG")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI diagram edit failed: {e}")

    # Upload new SVG
    storage = get_storage()
    prefix = f"matcha-work/{company_id}/{project_id}/diagrams"
    svg_bytes = new_svg.encode("utf-8")
    new_url = await storage.upload_file(svg_bytes, f"diagram-edited-{section_id[:8]}.svg", prefix=prefix, content_type="image/svg+xml")

    # Update section content — replace old img with new
    import re as _re
    new_img = f'<img src="{new_url}" alt="Diagram" data-diagram-index="0" style="max-width:100%;margin:8px 0;" />'
    old_content = section.get("content", "")
    updated_content = _re.sub(r'<img[^>]*data-diagram-index[^>]*/>', new_img, old_content)
    if updated_content == old_content:
        # Fallback: replace first img with diagram alt
        updated_content = _re.sub(r'<img[^>]*alt="Diagram"[^>]*/>', new_img, old_content, count=1)
    if updated_content == old_content:
        updated_content = old_content + new_img

    edit_source = "ai_region_edit" if (region and isinstance(region, dict)) else "ai_edit"
    new_diagram_data = [{"svg_source": new_svg, "storage_url": new_url, "created_from": edit_source}]
    await proj_svc.update_section(project_id, section_id, {
        "content": updated_content,
        "diagram_data": new_diagram_data,
    })

    updated = await proj_svc.get_project(project_id, company_id)
    return updated


@router.post("/projects/{project_id}/sections/{section_id}/edit-diagram-text")
async def edit_diagram_text(
    project_id: UUID,
    section_id: str,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Edit text labels in a diagram by direct string replacement."""
    from ..services import project_service as proj_svc

    await _verify_project_access(project_id, current_user)
    company_id = await get_client_company_id(current_user)
    edits = body.get("edits", [])
    if not edits:
        raise HTTPException(status_code=400, detail="No edits provided")

    project = await proj_svc.get_project(project_id, company_id)
    section = next((s for s in project.get("sections", []) if s.get("id") == section_id), None)
    if not section or not section.get("diagram_data"):
        raise HTTPException(status_code=404, detail="Section or diagram not found")

    svg_source = section["diagram_data"][0].get("svg_source", "")
    if not svg_source:
        raise HTTPException(status_code=400, detail="No SVG source")

    new_svg = svg_source
    for edit in edits:
        old_text = edit.get("old_text", "")
        new_text = edit.get("new_text", "")
        if old_text:
            new_svg = new_svg.replace(f">{old_text}<", f">{new_text}<")

    storage = get_storage()
    prefix = f"matcha-work/{company_id}/{project_id}/diagrams"
    new_url = await storage.upload_file(new_svg.encode("utf-8"), f"diagram-textedit-{section_id[:8]}.svg", prefix=prefix, content_type="image/svg+xml")

    import re as _re
    new_img = f'<img src="{new_url}" alt="Diagram" data-diagram-index="0" style="max-width:100%;margin:8px 0;" />'
    old_content = section.get("content", "")
    updated_content = _re.sub(r'<img[^>]*alt="Diagram"[^>]*/>', new_img, old_content, count=1)

    await proj_svc.update_section(project_id, section_id, {
        "content": updated_content,
        "diagram_data": [{"svg_source": new_svg, "storage_url": new_url, "created_from": "text_edit"}],
    })
    return await proj_svc.get_project(project_id, company_id)


@router.post("/projects/{project_id}/sections/{section_id}/save-diagram")
async def save_diagram(
    project_id: UUID,
    section_id: str,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Save a diagram from the visual editor (Excalidraw SVG export)."""
    from ..services import project_service as proj_svc

    await _verify_project_access(project_id, current_user)
    company_id = await get_client_company_id(current_user)
    svg = body.get("svg", "").strip()
    if not svg or "<svg" not in svg.lower():
        raise HTTPException(status_code=400, detail="Invalid SVG")

    storage = get_storage()
    prefix = f"matcha-work/{company_id}/{project_id}/diagrams"
    new_url = await storage.upload_file(svg.encode("utf-8"), f"diagram-visual-{section_id[:8]}.svg", prefix=prefix, content_type="image/svg+xml")

    import re as _re
    new_img = f'<img src="{new_url}" alt="Diagram" data-diagram-index="0" style="max-width:100%;margin:8px 0;" />'

    project = await proj_svc.get_project(project_id, company_id)
    section = next((s for s in project.get("sections", []) if s.get("id") == section_id), None)
    if not section:
        raise HTTPException(status_code=404, detail="Section not found")

    old_content = section.get("content", "")
    updated_content = _re.sub(r'<img[^>]*alt="Diagram"[^>]*/>', new_img, old_content, count=1)

    await proj_svc.update_section(project_id, section_id, {
        "content": updated_content,
        "diagram_data": [{"svg_source": svg, "storage_url": new_url, "created_from": "visual_editor"}],
    })
    return await proj_svc.get_project(project_id, company_id)


# ── Project collaborator endpoints ──


@router.get("/projects/{project_id}/collaborators")
async def list_project_collaborators(
    project_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List collaborators on a project."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.list_collaborators(project_id)


@router.post("/projects/{project_id}/collaborators")
async def add_project_collaborator(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Add a user as a collaborator. Only the project owner can invite."""
    from ..services import project_service as proj_svc
    _project, role = await _verify_project_access(project_id, current_user)
    if role != "owner":
        raise HTTPException(status_code=403, detail="Only the project owner can add collaborators")
    user_id = body.get("user_id")
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")
    try:
        return await proj_svc.add_collaborator(project_id, UUID(user_id), current_user.id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.delete("/projects/{project_id}/collaborators/{user_id}")
async def remove_project_collaborator(
    project_id: UUID,
    user_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Remove a collaborator from a project. Only the owner can do this."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    try:
        return await proj_svc.remove_collaborator(project_id, user_id, current_user.id)
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.post("/projects/{project_id}/invite")
async def invite_to_project(
    project_id: UUID,
    body: dict = Body(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Invite a user to a project by email. Creates pending collaborator + inbox notification + email."""
    from ...core.services.email import get_email_service

    await _verify_project_access(project_id, current_user)

    email = (body.get("email") or "").strip().lower()
    if not email:
        raise HTTPException(status_code=400, detail="email is required")

    async with get_connection() as conn:
        # Look up user by email
        invitee = await conn.fetchrow("SELECT id, email FROM users WHERE email = $1 AND is_active = true", email)
        if not invitee:
            raise HTTPException(status_code=404, detail="User not found. They need to create an account first.")

        invitee_id = invitee["id"]

        if invitee_id == current_user.id:
            raise HTTPException(status_code=400, detail="You cannot invite yourself")

        # Check if already a collaborator
        existing = await conn.fetchrow(
            "SELECT status FROM mw_project_collaborators WHERE project_id = $1 AND user_id = $2",
            project_id, invitee_id,
        )
        if existing:
            if existing["status"] == "active":
                raise HTTPException(status_code=400, detail="User is already a collaborator")
            if existing["status"] == "pending":
                raise HTTPException(status_code=400, detail="Invitation already pending")
            # Was removed — re-invite
            await conn.execute(
                "UPDATE mw_project_collaborators SET status = 'pending', invited_by = $3, created_at = NOW() WHERE project_id = $1 AND user_id = $2",
                project_id, invitee_id, current_user.id,
            )
        else:
            await conn.execute(
                """INSERT INTO mw_project_collaborators (project_id, user_id, invited_by, role, status)
                   VALUES ($1, $2, $3, 'collaborator', 'pending')""",
                project_id, invitee_id, current_user.id,
            )

        # Get project title and inviter name for notifications
        project = await conn.fetchrow("SELECT title FROM mw_projects WHERE id = $1", project_id)
        inviter = await conn.fetchrow("SELECT email FROM users WHERE id = $1", current_user.id)
        inviter_client = await conn.fetchrow("SELECT name FROM clients WHERE user_id = $1", current_user.id)
        inviter_name = (inviter_client["name"] if inviter_client else None) or inviter["email"].split("@")[0]
        project_title = project["title"] if project else "a project"

        # Create inbox notification
        msg_content = f"**{inviter_name}** has invited you to join the project **{project_title}**. Go to your projects to accept or decline."
        conversation = await conn.fetchrow(
            """INSERT INTO inbox_conversations (title, is_group, created_by, last_message_at, last_message_preview)
               VALUES ($1, false, $2, NOW(), $3)
               RETURNING id""",
            f"Project Invite: {project_title}", current_user.id, msg_content[:100],
        )
        conv_id = conversation["id"]
        await conn.execute(
            "INSERT INTO inbox_participants (conversation_id, user_id) VALUES ($1, $2)", conv_id, current_user.id,
        )
        await conn.execute(
            "INSERT INTO inbox_participants (conversation_id, user_id) VALUES ($1, $2)", conv_id, invitee_id,
        )
        await conn.execute(
            """INSERT INTO inbox_messages (conversation_id, sender_id, content)
               VALUES ($1, $2, $3)""",
            conv_id, current_user.id, msg_content,
        )

    # Send email notification
    email_svc = get_email_service()
    if email_svc.is_configured():
        settings = get_settings()
        base_url = settings.app_base_url.rstrip("/")
        try:
            await email_svc.send_email(
                to_email=email,
                to_name=email.split("@")[0],
                subject=f"{inviter_name} invited you to a project on Matcha",
                html_content=f"""
                <div style="font-family: -apple-system, sans-serif; max-width: 480px; margin: 0 auto;">
                    <h2 style="color: #e4e4e7;">Project Invitation</h2>
                    <p style="color: #a1a1aa;"><strong>{inviter_name}</strong> invited you to join <strong>{project_title}</strong>.</p>
                    <a href="{base_url}/work"
                       style="display: inline-block; background: #10b981; color: white; padding: 12px 28px;
                              border-radius: 8px; text-decoration: none; font-size: 14px; font-weight: 600;">
                        View Projects
                    </a>
                </div>
                """,
            )
        except Exception as exc:
            logger.warning("Failed to send project invite email: %s", exc)

    # Create MW notification for the invitee
    try:
        from ..services import notification_service as notif_svc
        company_id = await get_client_company_id(current_user)
        await notif_svc.create_notification(
            user_id=invitee_id,
            company_id=company_id,
            type="project_invite",
            title=f"Project invite from {inviter_name}",
            body=f"You've been invited to join \"{project_title}\"",
            link=f"/work",
            metadata={"project_id": str(project_id), "invited_by": str(current_user.id)},
        )
    except Exception as e:
        logger.warning("Failed to create invite notification: %s", e)

    return {"invited": True, "email": email}


@router.post("/projects/{project_id}/invite/accept")
async def accept_project_invite(
    project_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Accept a pending project invitation."""
    async with get_connection() as conn:
        result = await conn.execute(
            """UPDATE mw_project_collaborators
               SET status = 'active'
               WHERE project_id = $1 AND user_id = $2 AND status = 'pending'""",
            project_id, current_user.id,
        )
        if result.endswith("0"):
            raise HTTPException(status_code=404, detail="No pending invitation found")
    return {"accepted": True}


@router.post("/projects/{project_id}/invite/decline")
async def decline_project_invite(
    project_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Decline a pending project invitation."""
    async with get_connection() as conn:
        result = await conn.execute(
            """UPDATE mw_project_collaborators
               SET status = 'removed'
               WHERE project_id = $1 AND user_id = $2 AND status = 'pending'""",
            project_id, current_user.id,
        )
        if result.endswith("0"):
            raise HTTPException(status_code=404, detail="No pending invitation found")
    return {"declined": True}


@router.get("/project-invites")
async def list_pending_invites(
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List all pending project invitations for the current user."""
    async with get_connection() as conn:
        rows = await conn.fetch(
            """SELECT c.project_id, p.title AS project_title,
                      u.email AS invited_by_email, cl.name AS invited_by_name,
                      c.created_at
               FROM mw_project_collaborators c
               JOIN mw_projects p ON p.id = c.project_id
               JOIN users u ON u.id = c.invited_by
               LEFT JOIN clients cl ON cl.user_id = c.invited_by
               WHERE c.user_id = $1 AND c.status = 'pending'
               ORDER BY c.created_at DESC""",
            current_user.id,
        )
    return [
        {
            "project_id": str(r["project_id"]),
            "project_title": r["project_title"],
            "invited_by": r["invited_by_name"] or r["invited_by_email"].split("@")[0],
            "invited_at": r["created_at"].isoformat() if r["created_at"] else None,
        }
        for r in rows
    ]


@router.get("/admin-users/search")
async def search_admin_users_endpoint(
    q: str = Query(..., min_length=2, max_length=100),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Search admin users for the collaborator invite picker."""
    from ..services import project_service as proj_svc
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only admins can search admin users")
    return await proj_svc.search_admin_users(q, current_user.id)


@router.post("/projects/{project_id}/chats")
async def create_project_chat_endpoint(
    project_id: UUID,
    body: dict = {},
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Create a new chat within a project."""
    from ..services import project_service as proj_svc
    project, _role = await _verify_project_access(project_id, current_user)
    company_id = project.get("company_id")
    if company_id is None:
        raise HTTPException(status_code=400, detail="No company associated")
    return await proj_svc.create_project_chat(project_id, company_id, current_user.id, body.get("title"))


@router.post("/projects/{project_id}/posting/from-chat")
async def populate_posting_from_chat(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Extract structured posting fields from a chat message using AI."""
    from ..services import project_service as proj_svc

    project, _role = await _verify_project_access(project_id, current_user)

    content = body.get("content", "")
    if not content:
        raise HTTPException(status_code=400, detail="No content provided")

    # Use AI to extract structured fields
    ai_provider = get_ai_provider()
    prompt = (
        "Extract job posting fields from this text. Return ONLY valid JSON with these fields "
        "(use null for missing fields):\n"
        '{"title":"...","description":"...","requirements":"...","compensation":"...",'
        '"location":"...","employment_type":"full-time|part-time|contract"}\n\n'
        f"Text:\n---\n{content[:5000]}\n---"
    )
    ai_resp = await ai_provider.generate(
        [{"role": "user", "content": prompt}], {}, company_context=""
    )

    # Parse the AI response
    raw = ai_resp.assistant_reply.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
    try:
        fields = json.loads(raw)
    except Exception:
        # Fallback: put everything in description
        fields = {"description": _strip_markdown(content)}

    # Merge with existing posting (don't overwrite non-null fields with null)
    existing = (project.get("project_data") or {}).get("posting") or {}
    merged = {**existing}
    for k, v in fields.items():
        if v is not None and str(v).strip():
            merged[k] = v

    result = await proj_svc.update_project_data(project_id, {"posting": merged})
    return result


@router.put("/projects/{project_id}/posting")
async def update_project_posting(
    project_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Update the job posting data for a recruiting project."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.update_project_data(project_id, {"posting": body})


@router.post("/projects/{project_id}/shortlist/{candidate_id}")
async def toggle_project_shortlist(
    project_id: UUID,
    candidate_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Toggle a candidate on/off the shortlist."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.toggle_shortlist(project_id, candidate_id)


@router.post("/projects/{project_id}/dismiss/{candidate_id}")
async def toggle_project_dismiss(
    project_id: UUID,
    candidate_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Toggle a candidate on/off the dismissed list."""
    from ..services import project_service as proj_svc
    await _verify_project_access(project_id, current_user)
    return await proj_svc.toggle_dismiss(project_id, candidate_id)


@router.post("/projects/{project_id}/reject/{candidate_id}")
async def reject_project_candidate(
    project_id: UUID,
    candidate_id: str,
    body: RejectCandidateRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Reject a candidate, optionally sending a polite rejection email.

    When `send_email=false` this is equivalent to dismissing the candidate
    with a `status='rejected'` marker — no email goes out.
    """
    from ..services import project_service as proj_svc
    from app.core.services.email import EmailService

    project, _role = await _verify_project_access(project_id, current_user)
    data = project.get("project_data") or {}
    candidates = list(data.get("candidates") or [])

    candidate_idx = next(
        (i for i, c in enumerate(candidates) if c.get("id") == candidate_id),
        None,
    )
    if candidate_idx is None:
        raise HTTPException(status_code=404, detail="Candidate not found in project")

    candidate = candidates[candidate_idx]
    name = candidate.get("name") or candidate.get("filename", "Candidate")
    email = candidate.get("email")

    # Idempotency guard: if the candidate was already rejected + hidden, don't
    # re-send the email. Return the current project so the caller's state stays
    # in sync but the `email_sent` flag is honest.
    dismissed_ids = list(data.get("dismissed_ids") or [])
    already_rejected = (
        candidate.get("status") == "rejected" and candidate_id in dismissed_ids
    )
    if already_rejected:
        return {"project": project, "email_sent": False, "already_rejected": True}

    email_sent = False
    if body.send_email and email:
        company_id = project.get("company_id")
        async with get_connection() as conn:
            company_row = await conn.fetchrow(
                "SELECT name FROM companies WHERE id = $1", company_id
            )
        company_name = company_row["name"] if company_row else "the company"
        position_title = project.get("title") or "Open Position"

        email_svc = EmailService()
        try:
            email_sent = await email_svc.send_candidate_rejection_email(
                to_email=email,
                to_name=name,
                company_name=company_name,
                position_title=position_title,
                custom_message=body.custom_message,
            )
        except Exception as e:
            logger.error("Rejection email failed for %s: %s", email, e, exc_info=True)
            email_sent = False

    # Mutate candidate: mark rejected, store internal reason
    candidates[candidate_idx] = {
        **candidate,
        "status": "rejected",
        "rejection_reason": body.rejection_reason,
    }

    # Add to dismissed_ids so existing filters hide the candidate by default
    if candidate_id not in dismissed_ids:
        dismissed_ids.append(candidate_id)

    updated_project = await proj_svc.update_project_data(
        project_id,
        {"candidates": candidates, "dismissed_ids": dismissed_ids},
    )

    return {
        "project": updated_project,
        "email_sent": email_sent,
    }


@router.post("/projects/{project_id}/resume/upload")
async def upload_project_resumes(
    project_id: UUID,
    files: list[UploadFile] = File(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Upload resumes to a recruiting project — extract candidates into project_data."""
    from ..services import project_service as proj_svc

    project, _role = await _verify_project_access(project_id, current_user)

    # Require finalized posting before accepting resumes
    data = project.get("project_data") or {}
    posting = data.get("posting") or {}
    if not posting.get("finalized"):
        raise HTTPException(status_code=400, detail="Finalize the job posting before uploading resumes")

    # Validate files
    parsed_files: list[tuple[str, bytes, str]] = []
    for f in files:
        fname = f.filename or "resume"
        ext = os.path.splitext(fname)[1].lower()
        if ext not in RESUME_UPLOAD_EXTENSIONS:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {fname}")
        raw = await f.read()
        if len(raw) > RESUME_UPLOAD_MAX_BYTES:
            raise HTTPException(status_code=400, detail=f"File exceeds 10 MB limit: {fname}")
        parsed_files.append((fname, raw, f.content_type or "application/octet-stream"))

    async def event_stream():
        try:
            from google import genai as _genai
            from google.genai import types as _types

            api_key = os.getenv("GEMINI_API_KEY") or get_settings().gemini_api_key
            client = _genai.Client(api_key=api_key)
            extract_model = "gemini-3.1-flash-lite-preview"
            parser = ERDocumentParser()
            new_candidates = []

            for idx, (fname, raw, ct) in enumerate(parsed_files, 1):
                yield _sse_data({"type": "status", "message": f"Extracting text from {fname} ({idx}/{len(parsed_files)})..."})
                try:
                    text, _ = parser.extract_text_from_bytes(raw, fname)
                except Exception:
                    continue
                if not text or len(text.strip()) < 50:
                    continue

                yield _sse_data({"type": "status", "message": f"Analyzing {fname} ({idx}/{len(parsed_files)})..."})
                capped = text[:RESUME_TEXT_CAP]
                try:
                    resp = await asyncio.wait_for(
                        asyncio.to_thread(
                            lambda t=capped: client.models.generate_content(
                                model=extract_model,
                                contents=[_types.Content(role="user", parts=[_types.Part.from_text(text=RESUME_EXTRACT_PROMPT % t)])],
                                config=_types.GenerateContentConfig(temperature=0.1),
                            )
                        ),
                        timeout=60,
                    )
                    raw_json = (resp.text or "").strip()
                    if raw_json.startswith("```"):
                        raw_json = raw_json.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
                    data = json.loads(raw_json)
                    new_candidates.append({
                        "id": os.urandom(8).hex(),
                        "filename": fname,
                        "name": data.get("name"),
                        "email": data.get("email"),
                        "phone": data.get("phone"),
                        "location": data.get("location"),
                        "current_title": data.get("current_title"),
                        "experience_years": data.get("experience_years"),
                        "skills": data.get("skills"),
                        "education": data.get("education"),
                        "certifications": data.get("certifications"),
                        "summary": data.get("summary"),
                        "strengths": data.get("strengths"),
                        "flags": data.get("flags"),
                        "status": "analyzed",
                    })
                except Exception as e:
                    logger.warning("Resume extraction failed for %s: %s", fname, e)

            if new_candidates:
                yield _sse_data({"type": "status", "message": f"Adding {len(new_candidates)} candidates to project..."})
                result = await proj_svc.add_candidates_to_project(project_id, new_candidates)
                yield _sse_data({"type": "complete", "data": {"candidates_added": len(new_candidates), "project": result}})
            else:
                yield _sse_data({"type": "complete", "data": {"candidates_added": 0}})
        except Exception as e:
            logger.error("Project resume upload failed: %s", e, exc_info=True)
            yield _sse_data({"type": "error", "message": "Failed to process resumes."})
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/projects/placeholder-questions")
async def generate_placeholder_questions(
    body: dict,
    _current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Generate human-friendly questions for each placeholder using AI."""
    placeholders = body.get("placeholders") or []  # [{placeholder, label}]
    if not placeholders:
        return {"questions": []}

    items = "\n".join(
        f"- Placeholder: {p['placeholder']}, Context: \"{p.get('label', p['placeholder'])}\""
        for p in placeholders
    )

    try:
        from google import genai as _genai
        from google.genai import types as _types

        api_key = os.getenv("GEMINI_API_KEY") or get_settings().gemini_api_key
        client = _genai.Client(api_key=api_key)
        resp = await asyncio.to_thread(
            lambda: client.models.generate_content(
                model="gemini-3.1-flash-lite-preview",
                contents=[_types.Content(role="user", parts=[_types.Part.from_text(
                    text=f"Generate a short, friendly question for each placeholder below. The question should help someone fill in the blank in a job posting. Return ONE question per line, in order, no numbering or bullets.\n\n{items}"
                )])],
                config=_types.GenerateContentConfig(temperature=0.3),
            )
        )
        lines = [l.strip() for l in (resp.text or "").strip().split("\n") if l.strip()]
        # Pair questions with placeholders
        questions = []
        for i, p in enumerate(placeholders):
            q = lines[i] if i < len(lines) else f"What's the {p['placeholder']}?"
            questions.append({"placeholder": p["placeholder"], "label": p.get("label", ""), "question": q})
        return {"questions": questions}
    except Exception as e:
        logger.warning("Placeholder question generation failed: %s", e)
        # Fallback to raw names
        return {"questions": [
            {"placeholder": p["placeholder"], "label": p.get("label", ""), "question": f"What's the {p['placeholder']}?"}
            for p in placeholders
        ]}


@router.post("/projects/extract-value")
async def extract_placeholder_value(
    body: dict,
    _current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Extract a clean replacement value from a user's natural-language answer."""
    user_input = (body.get("input") or "").strip()
    placeholder = body.get("placeholder") or ""
    context = body.get("context") or ""

    if not user_input:
        return {"value": user_input}

    # Simple case: short input (≤ 3 words) — use directly
    if len(user_input.split()) <= 3:
        return {"value": user_input}

    # Complex input: use Gemini flash lite to extract the actual value
    try:
        from google import genai as _genai
        from google.genai import types as _types

        api_key = os.getenv("GEMINI_API_KEY") or get_settings().gemini_api_key
        client = _genai.Client(api_key=api_key)
        resp = await asyncio.to_thread(
            lambda: client.models.generate_content(
                model="gemini-3.1-flash-lite-preview",
                contents=[_types.Content(role="user", parts=[_types.Part.from_text(
                    text=f"Extract the exact value to fill in the placeholder {placeholder} from this user answer: \"{user_input}\"\n"
                         f"Context from the document: \"{context}\"\n"
                         f"Return ONLY the extracted value — no quotes, no explanation, just the value itself."
                )])],
                config=_types.GenerateContentConfig(temperature=0.0),
            )
        )
        extracted = (resp.text or "").strip().strip('"').strip("'")
        return {"value": extracted if extracted else user_input}
    except Exception as e:
        logger.warning("Value extraction failed, using raw input: %s", e)
        return {"value": user_input}


@router.post("/projects/{project_id}/resume/analyze")
async def analyze_project_candidates(
    project_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Rank candidates against the job posting using AI."""
    from ..services import project_service as proj_svc

    project, _role = await _verify_project_access(project_id, current_user)

    # Build posting text from sections (strip HTML tags for clean AI input)
    sections = project.get("sections") or []
    def _strip_html(html: str) -> str:
        return re.sub(r'<[^>]+>', '', html).strip()

    posting_text = "\n\n".join(
        f"{s.get('title', 'Untitled')}:\n{_strip_html(s.get('content', ''))}"
        for s in sections
    )
    if not posting_text.strip():
        raise HTTPException(status_code=400, detail="No posting content to analyze against")

    data = project.get("project_data") or {}
    candidates = data.get("candidates") or []
    if not candidates:
        raise HTTPException(status_code=400, detail="No candidates to analyze")

    # Build candidate summaries for the prompt
    candidate_entries = []
    for c in candidates:
        if c.get("status") not in ("analyzed", "interview_sent", "interview_completed", "interview_in_progress"):
            continue
        entry = f"ID: {c['id']}\nName: {c.get('name', 'Unknown')}\n"
        if c.get("current_title"):
            entry += f"Title: {c['current_title']}\n"
        if c.get("experience_years") is not None:
            entry += f"Experience: {c['experience_years']} years\n"
        if c.get("skills"):
            entry += f"Skills: {', '.join(c['skills'][:15])}\n"
        if c.get("education"):
            entry += f"Education: {c['education']}\n"
        if c.get("certifications"):
            entry += f"Certifications: {', '.join(c['certifications'])}\n"
        if c.get("summary"):
            entry += f"Summary: {c['summary']}\n"
        if c.get("strengths"):
            entry += f"Strengths: {', '.join(c['strengths'])}\n"
        if c.get("flags"):
            entry += f"Flags: {', '.join(c['flags'])}\n"
        candidate_entries.append(entry)

    if not candidate_entries:
        raise HTTPException(status_code=400, detail="No analyzed candidates to rank")

    try:
        from google import genai as _genai
        from google.genai import types as _types

        api_key = os.getenv("GEMINI_API_KEY") or get_settings().gemini_api_key
        client = _genai.Client(api_key=api_key)

        prompt = (
            "You are a recruiting analyst. Given the job posting and candidate profiles below, "
            "score each candidate on how well they match the posting (0-100). "
            "Return ONLY valid JSON — an array of objects with these exact fields:\n"
            '  [{"id": "<candidate_id>", "score": <0-100>, "summary": "<1-2 sentence reason>"}]\n'
            "Order by score descending (best match first).\n\n"
            f"=== JOB POSTING ===\n{posting_text}\n\n"
            f"=== CANDIDATES ===\n" + "\n---\n".join(candidate_entries)
        )

        resp = await asyncio.wait_for(
            asyncio.to_thread(
                lambda: client.models.generate_content(
                    model="gemini-3.1-flash-lite-preview",
                    contents=[_types.Content(role="user", parts=[_types.Part.from_text(text=prompt)])],
                    config=_types.GenerateContentConfig(temperature=0.1),
                )
            ),
            timeout=60,
        )

        raw = (resp.text or "").strip()
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        rankings = json.loads(raw)

        # Merge scores into candidates
        score_map = {r["id"]: r for r in rankings}
        updated_candidates = []
        for c in candidates:
            match = score_map.get(c["id"])
            if match:
                updated_candidates.append({
                    **c,
                    "match_score": match.get("score"),
                    "match_summary": match.get("summary"),
                })
            else:
                updated_candidates.append(c)

        # Sort by match_score descending
        updated_candidates.sort(key=lambda x: x.get("match_score") or 0, reverse=True)

        await proj_svc.update_project_data(project_id, {"candidates": updated_candidates})
        return {"analyzed": len(rankings), "candidates": updated_candidates}

    except json.JSONDecodeError as e:
        logger.error("Failed to parse ranking JSON: %s", e)
        raise HTTPException(status_code=500, detail="Failed to parse AI ranking response")
    except Exception as e:
        logger.error("Candidate analysis failed: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/projects/{project_id}/resume/send-interviews")
async def send_project_interviews(
    project_id: UUID,
    body: SendInterviewsRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Create screening interviews for selected project candidates and send invite emails."""
    import secrets as _secrets
    from ..services import project_service as proj_svc
    from app.core.services.email import EmailService

    project, _role = await _verify_project_access(project_id, current_user)

    data = project.get("project_data") or {}
    candidates = data.get("candidates") or []
    if not candidates:
        raise HTTPException(status_code=400, detail="No candidates in this project")

    company_id = project.get("company_id")
    async with get_connection() as conn:
        company_row = await conn.fetchrow("SELECT name FROM companies WHERE id = $1", company_id)
    company_name = company_row["name"] if company_row else "the company"

    position_title = body.position_title or project.get("title") or "Open Position"
    email_svc = EmailService()
    settings = get_settings()

    sent = []
    failed = []
    updated_candidates = list(candidates)

    for cid in body.candidate_ids:
        candidate = None
        candidate_idx = None
        for idx, c in enumerate(updated_candidates):
            if c.get("id") == cid:
                candidate = c
                candidate_idx = idx
                break

        if candidate is None:
            failed.append({"id": cid, "error": "Candidate not found in project"})
            continue

        email = candidate.get("email")
        name = candidate.get("name") or candidate.get("filename", "Candidate")
        if not email:
            failed.append({"id": cid, "error": f"No email for {name}"})
            continue

        try:
            invite_token = _secrets.token_urlsafe(32)
            interview_data = json.dumps({
                "invite_token": invite_token,
                "candidate_name": name,
                "candidate_email": email,
                "position_title": position_title,
                "company_name": company_name,
                "resume_batch_project_id": str(project_id),
                "resume_batch_candidate_id": cid,
            })

            async with get_connection() as interview_conn:
                interview_row = await interview_conn.fetchrow(
                    """
                    INSERT INTO interviews (id, company_id, interviewer_name, interview_type, raw_culture_data, status, created_at)
                    VALUES (gen_random_uuid(), $1, $2, 'screening', $3, 'pending', NOW())
                    RETURNING id
                    """,
                    company_id,
                    name,
                    interview_data,
                )
            interview_id = interview_row["id"]

            invite_url = f"{settings.app_base_url}/candidate-interview/{invite_token}"
            email_sent = await email_svc.send_candidate_interview_invite_email(
                to_email=email,
                to_name=name,
                company_name=company_name,
                position_title=position_title,
                invite_url=invite_url,
                custom_message=body.custom_message,
            )

            updated_candidates[candidate_idx] = {
                **candidate,
                "status": "interview_sent",
                "interview_id": str(interview_id),
            }

            sent.append({
                "id": cid,
                "name": name,
                "email": email,
                "interview_id": str(interview_id),
                "email_sent": email_sent,
            })
        except Exception as e:
            logger.error("Failed to create interview for project candidate %s: %s", cid, e, exc_info=True)
            failed.append({"id": cid, "error": str(e)})

    if sent:
        await proj_svc.update_project_data(project_id, {"candidates": updated_candidates})

    return {"sent": sent, "failed": failed}


@router.post("/projects/{project_id}/resume/sync-interviews")
async def sync_project_interviews(
    project_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Sync interview statuses back into project candidates."""
    from ..services import project_service as proj_svc

    project, _role = await _verify_project_access(project_id, current_user)

    data = project.get("project_data") or {}
    candidates = data.get("candidates") or []

    interview_ids = [c.get("interview_id") for c in candidates if c.get("interview_id")]
    if not interview_ids:
        return {"updated": 0}

    async with get_connection() as conn:
        rows = await conn.fetch(
            """
            SELECT id, status, screening_analysis
            FROM interviews
            WHERE id = ANY($1::uuid[])
            """,
            interview_ids,
        )

    interview_map = {str(r["id"]): r for r in rows}
    updated = 0
    updated_candidates = list(candidates)

    for idx, c in enumerate(updated_candidates):
        iid = c.get("interview_id")
        if not iid or iid not in interview_map:
            continue
        row = interview_map[iid]
        new_status = c.get("status")
        interview_status = row["status"]

        # Treat analyzing as completed for UI purposes — interview is over,
        # only the AI analysis is still running. The review modal handles the
        # "no analysis yet" state gracefully.
        if interview_status in ("completed", "analyzed", "analyzing") and c.get("status") != "interview_completed":
            new_status = "interview_completed"
        elif interview_status == "in_progress" and c.get("status") == "interview_sent":
            new_status = "interview_in_progress"

        score = None
        summary = None
        analysis = row.get("screening_analysis")
        if analysis:
            if isinstance(analysis, str):
                analysis = json.loads(analysis)
            score = analysis.get("overall_score") or analysis.get("score")
            summary = analysis.get("summary") or analysis.get("overall_assessment")

        if new_status != c.get("status") or score or summary:
            updated_candidates[idx] = {
                **c,
                "status": new_status,
                "interview_status": interview_status,
                "interview_score": score,
                "interview_summary": summary,
            }
            updated += 1

    if updated > 0:
        await proj_svc.update_project_data(project_id, {"candidates": updated_candidates})

    return {"updated": updated}


@router.get("/projects/{project_id}/export/{fmt}")
async def export_project_endpoint(
    project_id: UUID,
    fmt: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Export project as PDF, DOCX, or Markdown."""
    from ..services import project_service as proj_svc
    project, _role = await _verify_project_access(project_id, current_user)
    company_id = await get_client_company_id(current_user)

    title = project["title"]
    sections = project["sections"]

    if fmt not in ("pdf", "md", "docx", "md_frontmatter"):
        raise HTTPException(status_code=400, detail="Supported formats: pdf, md, md_frontmatter, docx")

    if fmt in ("md", "md_frontmatter"):
        md_lines: list[str] = []
        if fmt == "md_frontmatter":
            pdata = project.get("project_data") or {}
            slug = pdata.get("slug") or ""
            excerpt = (pdata.get("excerpt") or "").replace('"', '\\"')
            tags = pdata.get("tags") or []
            status = pdata.get("status") or "draft"
            published_at = pdata.get("published_at") or ""
            safe_title = title.replace('"', '\\"')
            md_lines.append("---")
            md_lines.append(f'title: "{safe_title}"')
            if slug:
                md_lines.append(f"slug: {slug}")
            if excerpt:
                md_lines.append(f'excerpt: "{excerpt}"')
            if tags:
                md_lines.append("tags:")
                for t in tags:
                    md_lines.append(f"  - {t}")
            md_lines.append(f"status: {status}")
            if published_at:
                md_lines.append(f"published_at: {published_at}")
            md_lines.append("---\n")
        md_lines.append(f"# {title}\n")
        for s in sections:
            if s.get("title"):
                md_lines.append(f"## {s['title']}\n")
            md_lines.append(s.get("content", "") + "\n")
        suffix = ".md"
        return Response(
            content="\n".join(md_lines),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{title}{suffix}"'},
        )

    if fmt == "pdf":
        import html as _html

        # Inline remote images as base64 data URIs so WeasyPrint can render them
        async def _inline_images(html_str: str) -> str:
            import re as _re, base64
            storage = get_storage()
            img_pattern = _re.compile(r'<img\s+[^>]*src="([^"]+)"', _re.IGNORECASE)
            result = html_str
            for match in reversed(list(img_pattern.finditer(result))):
                src = match.group(1)
                if not storage.is_supported_storage_path(src):
                    continue
                try:
                    data = await storage.download_file(src)
                    ext = src.rsplit(".", 1)[-1].lower() if "." in src else "png"
                    mime = {"svg": "image/svg+xml", "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}.get(ext, "image/png")
                    b64 = base64.b64encode(data).decode()
                    data_uri = f"data:{mime};base64,{b64}"
                    result = result[:match.start(1)] + data_uri + result[match.end(1):]
                except Exception:
                    pass
            return result

        sections_html = []
        for idx, s in enumerate(sections):
            heading = ""
            if s.get("title"):
                heading = f'<h2><span class="section-num">{idx + 1}.</span> {_html.escape(s["title"])}</h2>'
            content = s.get("content", "")
            content = await _inline_images(content)
            if content.lstrip().startswith("<"):
                content_html = content
            else:
                try:
                    import markdown as _md
                    content_html = _md.markdown(content, extensions=["tables", "fenced_code", "nl2br"])
                except ImportError:
                    content_html = f"<p>{_html.escape(content)}</p>"
            sections_html.append(f"{heading}\n<div class='section-body'>{content_html}</div>")

        body_html = "\n".join(sections_html)
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>
  @page {{ size: A4; margin: 50px 60px; }}
  * {{ box-sizing: border-box; }}
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; font-size: 11pt; line-height: 1.6; color: #1a1a1a; margin: 0; }}
  h1 {{ font-size: 22pt; font-weight: 700; color: #0f172a; margin: 0 0 6px 0; }}
  .title-rule {{ border: none; border-top: 3px solid #22c55e; margin: 0 0 30px 0; }}
  h2 {{ font-size: 14pt; font-weight: 600; color: #0f172a; margin: 28px 0 10px 0; padding-bottom: 6px; border-bottom: 1px solid #e2e8f0; }}
  .section-num {{ color: #22c55e; font-weight: 700; }}
  img {{ max-width: 100%; height: auto; page-break-inside: avoid; margin: 12px 0; border-radius: 4px; }}
  .section-body {{ margin-bottom: 16px; }}
  .section-body p {{ margin: 6px 0; color: #334155; }}
  .section-body ul, .section-body ol {{ margin: 6px 0; padding-left: 22px; color: #334155; }}
  .section-body li {{ margin: 3px 0; }}
  .section-body strong {{ color: #0f172a; }}
  pre {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 4px; padding: 10px; font-size: 9pt; white-space: pre-wrap; overflow-wrap: break-word; }}
  code {{ background: #f1f5f9; padding: 1px 5px; border-radius: 3px; font-size: 9pt; color: #b45309; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 9.5pt; }}
  th, td {{ border: 1px solid #e2e8f0; padding: 6px 10px; text-align: left; }}
  th {{ background: #f8fafc; font-weight: 600; color: #0f172a; }}
  a {{ color: #2563eb; text-decoration: none; }}
  .footer {{ margin-top: 40px; padding-top: 12px; border-top: 1px solid #e2e8f0; font-size: 8pt; color: #94a3b8; text-align: center; }}
</style>
</head><body>
<h1>{_html.escape(title)}</h1>
<hr class="title-rule">
{body_html}
<div class="footer">Generated with Matcha Work</div>
</body></html>"""

        try:
            from weasyprint import HTML
            pdf_bytes = await asyncio.to_thread(lambda: HTML(string=full_html).write_pdf())
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF generation not available")

        prefix = doc_svc.build_matcha_work_thread_storage_prefix(company_id, project_id, "project-exports")
        pdf_url = await get_storage().upload_file(
            pdf_bytes, f"{title}.pdf", prefix=prefix, content_type="application/pdf"
        )
        return {"pdf_url": pdf_url}

    if fmt == "docx":
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX generation not available")

        def _build_docx():
            doc = DocxDocument()
            doc.add_heading(title, level=0)
            for s in sections:
                if s.get("title"):
                    doc.add_heading(s["title"], level=1)
                for para in (s.get("content") or "").split("\n"):
                    if para.strip():
                        doc.add_paragraph(para)
            import io
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        docx_bytes = await asyncio.to_thread(_build_docx)
        prefix = doc_svc.build_matcha_work_thread_storage_prefix(company_id, project_id, "project-exports")
        docx_url = await get_storage().upload_file(
            docx_bytes, f"{title}.docx", prefix=prefix,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return {"docx_url": docx_url}


# ── Thread-scoped project endpoints (legacy, kept for backward compat) ──


@router.post("/threads/{thread_id}/project/init")
async def init_project(
    thread_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Initialize a project document with a title."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    result = await doc_svc.apply_update(thread_id, {
        "project_title": body.get("title", "Untitled Project"),
        "project_sections": (thread.get("current_state") or {}).get("project_sections") or [],
        "project_status": "drafting",
    })
    return {"current_state": result["current_state"], "version": result["version"]}


@router.post("/threads/{thread_id}/project/sections")
async def add_project_section(
    thread_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Add a section to the project."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    current_state = thread.get("current_state") or {}
    sections = list(current_state.get("project_sections") or [])

    raw_content = body.get("content", "")
    # Strip markdown when content comes from chat AI responses
    if body.get("source_message_id"):
        raw_content = _strip_markdown(raw_content)

    new_section = {
        "id": os.urandom(8).hex(),
        "title": body.get("title"),
        "content": raw_content,
        "source_message_id": body.get("source_message_id"),
    }
    sections.append(new_section)

    # Auto-init project if not already
    updates = {"project_sections": sections}
    if not current_state.get("project_title"):
        updates["project_title"] = "Untitled Project"
        updates["project_status"] = "drafting"

    result = await doc_svc.apply_update(thread_id, updates)
    return {"section": new_section, "current_state": result["current_state"], "version": result["version"]}


@router.put("/threads/{thread_id}/project/sections/reorder")
async def reorder_project_sections(
    thread_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Reorder project sections by providing an ordered list of section IDs."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    sections = list((thread.get("current_state") or {}).get("project_sections") or [])
    section_map = {s["id"]: s for s in sections}
    ordered_ids = body.get("section_ids", [])

    reordered = [section_map[sid] for sid in ordered_ids if sid in section_map]
    # Append any sections not in the ordered list
    seen = set(ordered_ids)
    for s in sections:
        if s["id"] not in seen:
            reordered.append(s)

    result = await doc_svc.apply_update(thread_id, {"project_sections": reordered})
    return {"current_state": result["current_state"], "version": result["version"]}


@router.put("/threads/{thread_id}/project/sections/{section_id}")
async def update_project_section(
    thread_id: UUID,
    section_id: str,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Update a project section's title or content."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    sections = list((thread.get("current_state") or {}).get("project_sections") or [])
    found = False
    for i, s in enumerate(sections):
        if s.get("id") == section_id:
            if "title" in body:
                sections[i] = {**s, "title": body["title"]}
            if "content" in body:
                sections[i] = {**sections[i], "content": body["content"]}
            found = True
            break

    if not found:
        raise HTTPException(status_code=404, detail="Section not found")

    result = await doc_svc.apply_update(thread_id, {"project_sections": sections})
    return {"current_state": result["current_state"], "version": result["version"]}


@router.delete("/threads/{thread_id}/project/sections/{section_id}")
async def delete_project_section(
    thread_id: UUID,
    section_id: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Remove a section from the project."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    sections = list((thread.get("current_state") or {}).get("project_sections") or [])
    sections = [s for s in sections if s.get("id") != section_id]

    result = await doc_svc.apply_update(thread_id, {"project_sections": sections})
    return {"current_state": result["current_state"], "version": result["version"]}


@router.get("/threads/{thread_id}/project/export/{fmt}")
async def export_project(
    thread_id: UUID,
    fmt: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Export project as PDF, DOCX, or Markdown."""
    if fmt not in ("pdf", "md", "docx"):
        raise HTTPException(status_code=400, detail="Supported formats: pdf, md, docx")

    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    state = thread.get("current_state") or {}
    title = state.get("project_title") or "Project"
    sections = state.get("project_sections") or []

    if fmt == "md":
        md_lines = [f"# {title}\n"]
        for s in sections:
            if s.get("title"):
                md_lines.append(f"## {s['title']}\n")
            md_lines.append(s.get("content", "") + "\n")
        md_content = "\n".join(md_lines)
        return Response(
            content=md_content,
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{title}.md"'},
        )

    if fmt == "pdf":
        import html as _html

        sections_html = []
        for idx, s in enumerate(sections):
            heading = ""
            if s.get("title"):
                heading = f'<h2><span class="section-num">{idx + 1}.</span> {_html.escape(s["title"])}</h2>'
            content = s.get("content", "")
            # Content may be HTML (from TipTap editor) or legacy markdown
            if content.lstrip().startswith("<"):
                content_html = content  # already HTML
            else:
                try:
                    import markdown as _md
                    content_html = _md.markdown(content, extensions=["tables", "fenced_code", "nl2br"])
                except ImportError:
                    content_html = f"<p>{_html.escape(content)}</p>"
            sections_html.append(f"{heading}\n<div class='section-body'>{content_html}</div>")

        body_html = "\n".join(sections_html)
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>
  @page {{ size: A4; margin: 50px 60px; }}
  * {{ box-sizing: border-box; }}
  body {{
    font-family: 'Helvetica Neue', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #1a1a1a;
    margin: 0;
  }}
  h1 {{
    font-size: 22pt;
    font-weight: 700;
    color: #0f172a;
    margin: 0 0 6px 0;
    letter-spacing: -0.5px;
  }}
  .subtitle {{
    font-size: 9pt;
    color: #94a3b8;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    margin-bottom: 20px;
  }}
  .title-rule {{
    border: none;
    border-top: 3px solid #22c55e;
    margin: 0 0 30px 0;
  }}
  h2 {{
    font-size: 14pt;
    font-weight: 600;
    color: #0f172a;
    margin: 28px 0 10px 0;
    padding-bottom: 6px;
    border-bottom: 1px solid #e2e8f0;
  }}
  .section-num {{
    color: #22c55e;
    font-weight: 700;
  }}
  img {{
    max-width: 100%;
    height: auto;
    page-break-inside: avoid;
    margin: 12px 0;
    border-radius: 4px;
  }}
  .section-body {{
    margin-bottom: 16px;
  }}
  .section-body p {{
    margin: 6px 0;
    color: #334155;
  }}
  .section-body ul, .section-body ol {{
    margin: 6px 0;
    padding-left: 22px;
    color: #334155;
  }}
  .section-body li {{
    margin: 3px 0;
  }}
  .section-body strong {{
    color: #0f172a;
  }}
  .section-body em {{
    color: #475569;
  }}
  pre {{
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 4px;
    padding: 10px 14px;
    font-size: 9pt;
    font-family: 'SF Mono', 'Menlo', monospace;
    overflow-wrap: break-word;
    white-space: pre-wrap;
    color: #334155;
  }}
  code {{
    background: #f1f5f9;
    padding: 1px 5px;
    border-radius: 3px;
    font-size: 9pt;
    font-family: 'SF Mono', 'Menlo', monospace;
    color: #b45309;
  }}
  table {{
    border-collapse: collapse;
    width: 100%;
    margin: 12px 0;
    font-size: 9.5pt;
  }}
  th, td {{
    border: 1px solid #e2e8f0;
    padding: 6px 10px;
    text-align: left;
  }}
  th {{
    background: #f8fafc;
    font-weight: 600;
    color: #0f172a;
  }}
  blockquote {{
    border-left: 3px solid #22c55e;
    margin: 12px 0;
    padding: 8px 16px;
    background: #f0fdf4;
    color: #334155;
  }}
  a {{
    color: #2563eb;
    text-decoration: none;
  }}
  .footer {{
    margin-top: 40px;
    padding-top: 12px;
    border-top: 1px solid #e2e8f0;
    font-size: 8pt;
    color: #94a3b8;
    text-align: center;
  }}
</style>
</head><body>
<h1>{_html.escape(title)}</h1>
<hr class="title-rule">
{body_html}
<div class="footer">Generated with Matcha Work</div>
</body></html>"""

        try:
            from weasyprint import HTML
            pdf_bytes = await asyncio.to_thread(lambda: HTML(string=full_html).write_pdf())
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF generation not available (WeasyPrint not installed)")

        prefix = doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "project-exports")
        pdf_url = await get_storage().upload_file(
            pdf_bytes, f"{title}.pdf", prefix=prefix, content_type="application/pdf"
        )
        return {"pdf_url": pdf_url}

    if fmt == "docx":
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX generation not available (python-docx not installed)")

        def _build_docx():
            doc = DocxDocument()
            doc.add_heading(title, level=0)
            for s in sections:
                if s.get("title"):
                    doc.add_heading(s["title"], level=1)
                for para in (s.get("content") or "").split("\n"):
                    if para.strip():
                        doc.add_paragraph(para)
            import io
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        docx_bytes = await asyncio.to_thread(_build_docx)
        prefix = doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "project-exports")
        docx_url = await get_storage().upload_file(
            docx_bytes, f"{title}.docx", prefix=prefix,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return {"docx_url": docx_url}


@router.post("/threads/{thread_id}/project/images")
async def upload_project_image(
    thread_id: UUID,
    file: UploadFile = File(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Upload an image for use in a project section. Returns the URL to embed as markdown."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    if not file.content_type or not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="Only image files are allowed")
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Image exceeds 10 MB limit")

    filename = file.filename or "image.png"
    url = await get_storage().upload_file(
        content,
        filename,
        prefix=doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "project-images"),
        content_type=file.content_type,
    )
    return {"url": url, "filename": filename}


INVENTORY_EXTRACT_PROMPT = """Extract inventory line items from this document (vendor invoice, inventory count, or order sheet).
Return ONLY valid JSON — an array of items:
[{"product_name":"...","sku":"...","category":"protein|produce|dairy|dry_goods|beverages|supplies|equipment|other","quantity":0,"unit":"case|lb|each|gal|oz|bag|box|doz|pack","unit_cost":0.00,"total_cost":0.00,"vendor":"..."}]

If this is a vendor invoice, extract the vendor name from the header or line items.
If quantities or costs are missing, use null. Compute total_cost as quantity * unit_cost when possible.

Document text:
---
%s
---"""


@router.post("/threads/{thread_id}/inventory/upload")
async def upload_thread_inventory(
    thread_id: UUID,
    files: list[UploadFile] = File(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Upload invoices/spreadsheets, extract structured inventory items, and stream batch insights."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    if thread["status"] in ("finalized", "archived"):
        raise HTTPException(status_code=400, detail=f"Cannot upload to a {thread['status']} thread")

    if current_user.role != "admin":
        await token_budget_service.check_token_budget(company_id)

    parsed_files: list[tuple[str, bytes, str]] = []
    for f in files:
        fname = f.filename or "document"
        ext = os.path.splitext(fname)[1].lower()
        if ext not in INVENTORY_UPLOAD_EXTENSIONS:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {fname}")
        raw = await f.read()
        if len(raw) > INVENTORY_UPLOAD_MAX_BYTES:
            raise HTTPException(status_code=400, detail=f"File exceeds 15 MB limit: {fname}")
        parsed_files.append((fname, raw, f.content_type or "application/octet-stream"))

    file_count = len(parsed_files)
    filenames = [pf[0] for pf in parsed_files]
    user_content = f"[Inventory batch: {file_count} file{'s' if file_count != 1 else ''}]\n" + "\n".join(f"- {fn}" for fn in filenames)
    user_msg = await doc_svc.add_message(thread_id, "user", user_content)

    async def event_stream():
        try:
            from google import genai as _genai
            from google.genai import types as _types

            api_key = os.getenv("GEMINI_API_KEY") or get_settings().gemini_api_key
            client = _genai.Client(api_key=api_key)
            extract_model = "gemini-3.1-flash-lite-preview"

            parser = ERDocumentParser()
            existing_items = list((thread.get("current_state") or {}).get("inventory_items") or [])
            new_items = []

            for idx, (fname, raw, ct) in enumerate(parsed_files, 1):
                yield _sse_data({"type": "status", "message": f"Reading {fname} ({idx}/{file_count})..."})

                # Extract text — CSV/TXT read directly, PDF/DOCX via parser
                ext = os.path.splitext(fname)[1].lower()
                try:
                    if ext == ".csv":
                        text = raw.decode("utf-8", errors="replace")
                    elif ext in (".xlsx", ".xls"):
                        # Excel files are binary — extract cell values via openpyxl
                        import io as _io
                        try:
                            import openpyxl
                            wb = openpyxl.load_workbook(_io.BytesIO(raw), read_only=True, data_only=True)
                            rows = []
                            ws = wb.active
                            for row in ws.iter_rows(values_only=True):
                                rows.append(",".join(str(c) if c is not None else "" for c in row))
                            wb.close()
                            text = "\n".join(rows)
                        except ImportError:
                            logger.warning("openpyxl not installed — cannot read Excel files")
                            yield _sse_data({"type": "status", "message": f"Cannot read Excel files (openpyxl not installed), skipping {fname}..."})
                            continue
                    else:
                        text, _ = parser.extract_text_from_bytes(raw, fname)
                except Exception:
                    yield _sse_data({"type": "status", "message": f"Could not read {fname}, skipping..."})
                    continue

                if not text or len(text.strip()) < 10:
                    continue

                # Upload to S3
                try:
                    await get_storage().upload_file(
                        raw, fname,
                        prefix=doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "inventory"),
                        content_type=ct,
                    )
                except Exception:
                    pass

                # AI extraction
                yield _sse_data({"type": "status", "message": f"Extracting items from {fname} ({idx}/{file_count})..."})
                capped = text[:INVENTORY_TEXT_CAP]

                try:
                    resp = await asyncio.wait_for(
                        asyncio.to_thread(
                            lambda t=capped: client.models.generate_content(
                                model=extract_model,
                                contents=[_types.Content(role="user", parts=[_types.Part.from_text(text=INVENTORY_EXTRACT_PROMPT % t)])],
                                config=_types.GenerateContentConfig(temperature=0.1),
                            )
                        ),
                        timeout=60,
                    )
                    raw_json = (resp.text or "").strip()
                    if raw_json.startswith("```"):
                        raw_json = raw_json.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
                    items_data = json.loads(raw_json)
                    if not isinstance(items_data, list):
                        items_data = [items_data]

                    for item in items_data:
                        new_items.append({
                            "id": os.urandom(8).hex(),
                            "filename": fname,
                            "product_name": item.get("product_name"),
                            "sku": item.get("sku"),
                            "category": item.get("category"),
                            "quantity": item.get("quantity"),
                            "unit": item.get("unit"),
                            "unit_cost": item.get("unit_cost"),
                            "total_cost": item.get("total_cost"),
                            "vendor": item.get("vendor"),
                            "status": "extracted",
                        })
                    yield _sse_data({"type": "status", "message": f"Extracted {len(items_data)} items from {fname}"})
                except Exception as e:
                    logger.warning("Inventory AI extraction failed for %s: %s", fname, e)
                    yield _sse_data({"type": "status", "message": f"Could not extract items from {fname}"})

            # Accumulate
            all_items = existing_items + new_items
            total_cost = sum(i.get("total_cost") or 0 for i in all_items)
            vendors = sorted(set(i.get("vendor") for i in all_items if i.get("vendor")))

            result = await doc_svc.apply_update(thread_id, {
                "inventory_items": all_items,
                "inventory_status": "ready",
                "inventory_total_count": len(all_items),
                "inventory_total_cost": round(total_cost, 2),
                "inventory_vendors": vendors,
            })
            current_state = result["current_state"]
            current_version = result["version"]

            # Batch summary
            yield _sse_data({"type": "status", "message": "Generating inventory insights..."})
            if new_items:
                # Build vendor breakdown
                vendor_totals: dict[str, float] = {}
                cat_totals: dict[str, int] = {}
                for i in new_items:
                    v = i.get("vendor") or "Unknown"
                    vendor_totals[v] = vendor_totals.get(v, 0) + (i.get("total_cost") or 0)
                    c = i.get("category") or "other"
                    cat_totals[c] = cat_totals.get(c, 0) + 1

                vendor_lines = ", ".join(f"{v}: ${t:,.2f}" for v, t in sorted(vendor_totals.items(), key=lambda x: -x[1]))
                cat_lines = ", ".join(f"{c}: {n}" for c, n in sorted(cat_totals.items(), key=lambda x: -x[1]))
                new_cost = sum(i.get("total_cost") or 0 for i in new_items)

                batch_reply = (
                    f"**Processed {len(new_items)} items** from {file_count} file{'s' if file_count != 1 else ''}\n\n"
                    f"**Total cost:** ${new_cost:,.2f}\n\n"
                    f"**By vendor:** {vendor_lines}\n\n"
                    f"**By category:** {cat_lines}"
                )
                if len(all_items) > len(new_items):
                    batch_reply += f"\n\n*Running total: {len(all_items)} items, ${total_cost:,.2f} across all uploads.*"
            else:
                batch_reply = f"Processed {file_count} file{'s' if file_count != 1 else ''} but could not extract any line items."

            assistant_msg = await doc_svc.add_message(thread_id, "assistant", batch_reply)

            response = SendMessageResponse(
                user_message=_row_to_message(user_msg),
                assistant_message=_row_to_message(assistant_msg),
                current_state=current_state,
                version=current_version,
                task_type=_infer_skill_from_state(current_state),
                pdf_url=None,
                token_usage=None,
            )
            yield _sse_data({"type": "complete", "data": response.model_dump(mode="json")})
        except Exception as e:
            logger.error("Inventory batch failed for thread %s: %s", thread_id, e, exc_info=True)
            yield _sse_data({"type": "error", "message": "Failed to process inventory files. Please try again."})
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ── Agent endpoints (email) ──


@router.get("/agent/email/status")
async def agent_email_status(
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Check if the current user has Gmail connected."""
    from ..services.gmail_service import GmailService
    gmail = GmailService(current_user.id)
    return await gmail.get_status()


@router.post("/agent/email/connect")
async def agent_email_connect(
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Start Google OAuth flow. Returns an auth_url to open in a popup."""
    from ..services.gmail_service import get_oauth_credentials, GMAIL_SCOPES
    import urllib.parse

    creds = get_oauth_credentials()
    if not creds:
        raise HTTPException(status_code=500, detail="Google OAuth credentials not configured on the server")

    settings = get_settings()
    redirect_uri = f"{settings.app_base_url}/api/matcha-work/agent/email/callback"

    # Encode user ID in state so callback knows who to store the token for
    from ...core.services.secret_crypto import encrypt_secret
    state = encrypt_secret(str(current_user.id))

    params = {
        "client_id": creds["client_id"],
        "redirect_uri": redirect_uri,
        "response_type": "code",
        "scope": " ".join(GMAIL_SCOPES),
        "access_type": "offline",
        "prompt": "consent",
        "state": state,
    }
    auth_url = f"https://accounts.google.com/o/oauth2/v2/auth?{urllib.parse.urlencode(params)}"
    return {"auth_url": auth_url}


@router.get("/agent/email/callback")
async def agent_email_callback(
    code: str = Query(...),
    state: str = Query(...),
):
    """OAuth callback — exchange code for tokens, store encrypted in DB, close popup."""
    from ..services.gmail_service import GmailService, get_oauth_credentials, GMAIL_SCOPES
    from ...core.services.secret_crypto import decrypt_secret as _decrypt

    # Recover user ID from state
    try:
        user_id_str = _decrypt(state)
        from uuid import UUID as _UUID
        user_id = _UUID(user_id_str)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid OAuth state")

    creds = get_oauth_credentials()
    if not creds:
        raise HTTPException(status_code=500, detail="Google OAuth credentials not configured")

    settings = get_settings()
    redirect_uri = f"{settings.app_base_url}/api/matcha-work/agent/email/callback"

    # Exchange authorization code for tokens
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": creds["client_id"],
                "client_secret": creds["client_secret"],
                "redirect_uri": redirect_uri,
                "grant_type": "authorization_code",
            },
            timeout=15.0,
        )
        if resp.status_code != 200:
            logger.error("Gmail OAuth token exchange failed: %s", resp.text)
            raise HTTPException(status_code=400, detail="Failed to exchange authorization code")
        tokens = resp.json()

    # Store encrypted token in DB
    gmail = GmailService(user_id)
    await gmail.save_token({
        "token": tokens["access_token"],
        "refresh_token": tokens.get("refresh_token"),
        "client_id": creds["client_id"],
        "client_secret": creds["client_secret"],
        "scopes": GMAIL_SCOPES,
    })

    # Return HTML that closes the popup
    return Response(
        content="""<!DOCTYPE html><html><body>
<script>window.opener && window.opener.postMessage('gmail-connected', '*'); window.close();</script>
<p>Gmail connected. You can close this window.</p>
</body></html>""",
        media_type="text/html",
    )


@router.delete("/agent/email/disconnect")
async def agent_email_disconnect(
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Remove the current user's Gmail connection."""
    async with get_connection() as conn:
        await conn.execute("UPDATE users SET gmail_token=NULL WHERE id=$1", current_user.id)
    return {"status": "disconnected"}


@router.post("/agent/email/fetch")
async def agent_email_fetch(
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Fetch unread emails for the current user."""
    from ..services.gmail_service import GmailService
    gmail = GmailService(current_user.id)
    await gmail.load_token()
    if not gmail.is_configured:
        raise HTTPException(status_code=400, detail="Gmail not connected")
    emails = await gmail.fetch_unread(max_results=25)
    return {"emails": emails}


@router.post("/agent/email/draft")
async def agent_email_draft(
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Draft a reply to an email using AI."""
    from ..services.gmail_service import GmailService
    gmail = GmailService(current_user.id)
    await gmail.load_token()
    if not gmail.is_configured:
        raise HTTPException(status_code=400, detail="Gmail not connected")

    email_id = body.get("email_id")
    instructions = body.get("instructions", "Write a helpful, concise reply.")
    if not email_id:
        raise HTTPException(status_code=400, detail="email_id is required")

    email = await gmail.get_message(email_id)

    ai_provider = get_ai_provider()
    prompt = (
        f"Draft a professional reply to this email. Return ONLY the reply body text, no subject line.\n"
        f"Instructions: {instructions}\n\n"
        f"Original email:\nFrom: {email['from']}\nSubject: {email['subject']}\nBody:\n{email['body'][:3000]}"
    )
    ai_resp = await ai_provider.generate(
        [{"role": "user", "content": prompt}], {}, company_context=""
    )
    draft_body = ai_resp.assistant_reply

    result = await gmail.create_draft(
        to=email["from"],
        subject=f"Re: {email['subject']}" if not email["subject"].startswith("Re:") else email["subject"],
        body=draft_body,
        reply_to_id=email_id,
    )

    return {
        "draft_id": result.get("id"),
        "to": email["from"],
        "subject": email["subject"],
        "body": draft_body,
    }


@router.post("/agent/email/send")
async def agent_email_send(
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Send an email via Gmail."""
    from ..services.gmail_service import GmailService
    gmail = GmailService(current_user.id)
    await gmail.load_token()
    if not gmail.is_configured:
        raise HTTPException(status_code=400, detail="Gmail not connected")

    to = body.get("to")
    subject = body.get("subject")
    email_body = body.get("body")
    reply_to_id = body.get("reply_to_id")

    if not all([to, subject, email_body]):
        raise HTTPException(status_code=400, detail="to, subject, and body are required")

    result = await gmail.send_email(to=to, subject=subject, body=email_body, reply_to_id=reply_to_id)
    return {"message_id": result.get("id"), "to": to, "subject": subject}


@router.get("/threads", response_model=list[ThreadListItem])
async def list_threads(
    status: Optional[str] = Query(None, pattern="^(active|finalized|archived)$"),
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List threads for the current company (includes threads where user is a collaborator)."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        return []

    threads = await doc_svc.list_threads(company_id, status=status, limit=limit, offset=offset, user_id=current_user.id)
    return [ThreadListItem(**t) for t in threads]


@router.get("/elements", response_model=list[ElementListItem])
async def list_elements(
    status: Optional[str] = Query(None, pattern="^(active|finalized|archived)$"),
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List Matcha Work elements for the current company."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        return []

    elements = await doc_svc.list_elements(company_id, status=status, limit=limit, offset=offset)
    return [ElementListItem(**e) for e in elements]


@router.get("/threads/{thread_id}", response_model=ThreadDetailResponse)
async def get_thread(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Get thread detail with all messages."""
    company_id = await get_client_company_id(current_user)
    # Don't 404 on missing company_id — collaborators (individuals invited to a
    # thread) may have no company of their own. Let _build_thread_detail_response
    # resolve access via the collaborator check in get_thread().
    return await _build_thread_detail_response(thread_id, company_id, user_id=current_user.id)


@router.post("/threads/{thread_id}/messages", response_model=SendMessageResponse)
async def send_message(
    thread_id: UUID,
    body: SendMessageRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Send a user message → AI response → state update → PDF."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id, user_id=current_user.id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    if thread["status"] == "finalized":
        raise HTTPException(status_code=400, detail="Cannot send messages to a finalized thread")

    if thread["status"] == "archived":
        raise HTTPException(status_code=400, detail="Cannot send messages to an archived thread")

    if current_user.role != "admin":
        await token_budget_service.check_token_budget(company_id)

    # Save user message
    user_msg = await doc_svc.add_message(thread_id, "user", body.content)

    # Fetch message history + company profile + context summary in parallel
    messages, profile, (context_summary, summary_at_count) = await asyncio.gather(
        doc_svc.get_thread_messages(thread_id, limit=20),
        doc_svc.get_company_profile_for_ai(company_id),
        doc_svc.get_context_summary(thread_id),
    )
    msg_dicts = [{"role": m["role"], "content": m["content"]} for m in messages]

    # Inject selected slide content into the AI-facing message (not saved to DB)
    _inject_slide_context(msg_dicts, thread["current_state"], body.slide_index)

    # Call AI with company context
    ai_provider = get_ai_provider()
    ctx = _build_company_context(profile)

    # Inject project file attachments metadata
    if thread.get("project_id"):
        from ..services import project_file_service
        pfiles = await project_file_service.list_project_files(thread["project_id"])
        if pfiles:
            listing = "\n".join(f"- {f['filename']} ({f['content_type']}, {f['file_size']:,} bytes)" for f in pfiles)
            ctx += f"\n\n=== PROJECT ATTACHMENTS ===\nThe user has attached these files to the project. Reference them when relevant:\n{listing}\n"

    # Inject recruiting project context so AI generates posting sections in the right project
    ctx = await _inject_recruiting_project_context(ctx, thread, thread["current_state"])

    # No-project guard: prevent the AI from hallucinating a "project panel" for
    # plain chat threads. Without this, once current_state accumulates
    # project_title/project_sections from a prior reply, _infer_skill_from_state
    # locks skill="project" and the AI keeps claiming it updated a document the
    # user has no UI to see.
    if not thread.get("project_id"):
        ctx += (
            "\n\n=== PLAIN THREAD (NO PROJECT) ==="
            "\nThis is a plain chat thread. Per the Surface architecture section at the top of the system prompt: threads cannot contain projects, and no project is attached to this thread."
            "\n- Set mode=\"general\", skill=\"none\", operation=\"none\". Never emit project_title, project_sections, blog_outline, blog_section_draft, or blog_section_revision."
            "\n- Your reply text is the only artifact. Do not reference any panel, canvas, document, or draft surface — none exists for this chat."
            "\n- Short-form content (LinkedIn posts, social captions, emails, summaries, cover letters): write it directly in your reply."
            "\n- If the user wants a blog post: tell them to use + next to Projects in the sidebar and choose 'Blog Post'."
            "\n- If the user wants a multi-section long-form document (strategy report, multi-page brief, job posting): tell them to create a Project from the sidebar and chat inside it."
        )

    # Grounded web search pre-pass for time-sensitive questions
    # (markets today, news, weather, scores, etc.) — fetches current facts via
    # Gemini Google Search grounding and injects them into the context.
    if needs_live_web_context(body.content):
        from ...config import get_settings as _get_settings
        live_ctx = await fetch_live_web_context(body.content, _get_settings())
        if live_ctx:
            ctx += live_ctx

    compliance_result: ComplianceContextResult | None = None
    if thread.get("node_mode"):
        node_ctx = await build_node_context(company_id)
        ctx += "\n\n" + node_ctx
    if thread.get("compliance_mode"):
        compliance_result = await build_compliance_context(company_id)
        ctx += "\n\n" + compliance_result.context_text

        # RAG augmentation — find requirements most relevant to the user's question
        rag_ctx = await _get_rag_context(body.content, company_id)
        if rag_ctx:
            ctx += "\n\n=== RELEVANT REGULATIONS (semantic search) ===\n" + rag_ctx

    # Payer mode — build dedicated medical policy prompt (separate from HR copilot)
    payer_prompt = None
    payer_sources: list[dict] = []
    if thread.get("payer_mode"):
        try:
            import os as _os
            from ...core.services.embedding_service import EmbeddingService
            from ...core.services.payer_policy_rag import PayerPolicyRAGService
            from ...config import get_settings as _get_settings
            from ..services.matcha_work_ai import PAYER_MODE_SYSTEM_PROMPT
            from datetime import date as _date

            user_msg = body.content or ""
            _api_key = _os.getenv("GEMINI_API_KEY") or _get_settings().gemini_api_key
            if _api_key and user_msg:
                _emb = EmbeddingService(api_key=_api_key)
                _rag = PayerPolicyRAGService(_emb)
                async with get_connection() as _pconn:
                    payer_ctx, payer_sources = await _rag.get_context_for_query(
                        query=user_msg, conn=_pconn,
                        company_id=company_id, max_tokens=6000,
                    )
                company_name = profile.get("name", "your company")
                payer_prompt = PAYER_MODE_SYSTEM_PROMPT.format(
                    company_name=company_name,
                    today=_date.today().isoformat(),
                    payer_context=payer_ctx or "No matching payer policy data found. Answer based on general knowledge but clearly state this is not from verified policy data.",
                )
        except Exception as _e:
            logger.warning("Failed to build payer policy prompt: %s", _e)

    # If no project is attached, scrub any leftover project_* state so
    # _infer_skill_from_state doesn't keep locking the AI into "project" skill.
    ai_facing_state = thread["current_state"]
    if not thread.get("project_id") and isinstance(ai_facing_state, dict):
        if "project_title" in ai_facing_state or "project_sections" in ai_facing_state:
            ai_facing_state = {
                k: v for k, v in ai_facing_state.items()
                if k not in ("project_title", "project_sections", "project_status")
            }

    project_meta = await _fetch_project_meta(thread.get("project_id"))
    blog_mode_state = _blog_mode_state_from_meta(project_meta)

    ai_resp = await ai_provider.generate(
        msg_dicts, ai_facing_state, company_context=ctx,
        slide_index=body.slide_index, context_summary=context_summary,
        payer_mode_prompt=payer_prompt,
        compliance_mode=bool(thread.get("compliance_mode")),
        payer_mode=bool(thread.get("payer_mode")),
        node_mode=bool(thread.get("node_mode")),
        blog_mode_state=blog_mode_state,
    )
    _scope_slide_update(ai_resp, thread["current_state"], body.slide_index)
    final_usage = ai_resp.token_usage

    current_version = thread["version"]
    (
        current_state,
        current_version,
        pdf_url,
        changed,
        assistant_reply_text,
    ) = await _apply_ai_updates_and_operations(
        thread_id=thread_id,
        company_id=company_id,
        ai_resp=ai_resp,
        current_state=thread["current_state"],
        current_version=current_version,
        user_message=body.content,
        current_user_id=current_user.id,
        project_id=thread.get("project_id"),
        project_meta=project_meta,
    )

    # Build metadata from compliance reasoning chains + payer sources
    msg_metadata = _build_compliance_metadata(compliance_result, ai_resp)
    if payer_sources:
        if msg_metadata is None:
            msg_metadata = {}
        msg_metadata["payer_sources"] = payer_sources

    # Cross-reference affected employees + detect policy gaps when both node + compliance are on
    if thread.get("node_mode") and thread.get("compliance_mode") and msg_metadata:
        if msg_metadata.get("referenced_locations"):
            affected = await _get_affected_employees(company_id, msg_metadata)
            if affected:
                msg_metadata["affected_employees"] = affected
        gaps = await _detect_compliance_gaps(company_id, msg_metadata)
        if gaps:
            msg_metadata["compliance_gaps"] = gaps

    # Annotate reply with change summary for conversation continuity
    if changed and ai_resp.structured_update and isinstance(ai_resp.structured_update, dict):
        update_slides = ai_resp.structured_update.get("slides")
        if update_slides and body.slide_index is not None and 0 <= body.slide_index < len(update_slides):
            changed_slide = update_slides[body.slide_index]
            if isinstance(changed_slide, dict):
                n_bullets = len(changed_slide.get("bullets", []))
                change_note = f"\n\n[Applied changes to Slide {body.slide_index + 1}: title=\"{changed_slide.get('title', '')}\", {n_bullets} bullets]"
                assistant_reply_text += change_note

    # Consultation: parse ACTION ITEMS DETECTED block from the reply, strip it
    # from the visible text, and append each item to project_data.action_items
    # with pending_confirmation=true so the UI can show them as ✨ proposed.
    if thread.get("project_id"):
        try:
            async with get_connection() as _conn:
                _ptype = await _conn.fetchval(
                    "SELECT project_type FROM mw_projects WHERE id = $1",
                    thread["project_id"],
                )
            if _ptype == "consultation":
                cleaned, detected_items = _extract_action_items(assistant_reply_text)
                if detected_items:
                    assistant_reply_text = cleaned
                    from ..services import project_service as _proj_svc
                    for item_text in detected_items:
                        try:
                            await _proj_svc.append_action_item(
                                thread["project_id"],
                                text=item_text,
                                source_thread_id=str(thread_id),
                                pending_confirmation=True,
                            )
                        except Exception:
                            logger.warning(
                                "Failed to append detected action item for project %s",
                                thread["project_id"], exc_info=True,
                            )
        except Exception:
            logger.warning("Consultation action-item parser failed", exc_info=True)

    # Save assistant message
    assistant_msg = await doc_svc.add_message(
        thread_id,
        "assistant",
        assistant_reply_text,
        version_created=current_version if changed else None,
        metadata=msg_metadata,
    )

    # Escalate low-confidence queries for human review
    if should_escalate(ai_resp):
        try:
            await create_escalation(
                company_id=company_id,
                thread_id=thread_id,
                user_message_id=user_msg["id"],
                assistant_message_id=assistant_msg["id"],
                user_query=body.content,
                ai_resp=ai_resp,
            )
        except Exception:
            logger.exception("Failed to create escalation for thread %s", thread_id)

    cost = calculate_call_cost(
        model=str((final_usage or {}).get("model") or "unknown"),
        prompt_tokens=(final_usage or {}).get("prompt_tokens"),
        completion_tokens=(final_usage or {}).get("completion_tokens"),
    )
    if final_usage is not None:
        final_usage["cost_dollars"] = float(cost)

    try:
        await doc_svc.log_token_usage_event(
            company_id=company_id,
            user_id=current_user.id,
            thread_id=thread_id,
            token_usage=final_usage,
            operation="send_message",
            cost_dollars=float(cost),
        )
    except Exception as e:
        logger.warning("Failed to log Matcha Work token usage for thread %s: %s", thread_id, e)

    if current_user.role != "admin":
        total_tokens = (final_usage or {}).get("total_tokens") or 0
        if total_tokens > 0:
            try:
                async with get_connection() as conn:
                    async with conn.transaction():
                        await token_budget_service.deduct_tokens(conn, company_id, total_tokens)
            except HTTPException:
                logger.warning("Token budget exhausted during deduction for thread %s", thread_id)
            except Exception as exc:
                logger.warning("Failed to deduct tokens for thread %s: %s", thread_id, exc)

    # Trigger conversation compaction in the background if needed
    asyncio.create_task(_maybe_compact(thread_id, ai_provider, summary_at_count))

    return SendMessageResponse(
        user_message=_row_to_message(user_msg),
        assistant_message=_row_to_message(assistant_msg),
        current_state=current_state,
        version=current_version,
        task_type=_infer_skill_from_state(current_state),
        pdf_url=pdf_url,
        token_usage=final_usage,
    )


_compacting_threads: set[UUID] = set()  # simple guard against concurrent compaction
_COMPACTION_REFRESH_INTERVAL = 20  # re-compact after this many new messages


async def _maybe_compact(thread_id: UUID, ai_provider, summary_at_count: int | None) -> None:
    """Check message count and run compaction if threshold is exceeded or summary is stale."""
    if thread_id in _compacting_threads:
        return
    try:
        _compacting_threads.add(thread_id)
        msg_count = await doc_svc.get_thread_message_count(thread_id)
        if msg_count < 30:
            return
        # Skip if summary is recent enough
        if summary_at_count is not None and (msg_count - summary_at_count) < _COMPACTION_REFRESH_INTERVAL:
            return
        all_messages = await doc_svc.get_thread_messages(thread_id)
        msg_dicts = [{"role": m["role"], "content": m["content"]} for m in all_messages]
        summary = await compact_conversation(msg_dicts, ai_provider.client)
        if summary:
            await doc_svc.save_context_summary(thread_id, summary, msg_count)
            logger.info("Compacted conversation for thread %s (%d messages)", thread_id, msg_count)
    except Exception:
        logger.warning("Background compaction failed for thread %s", thread_id, exc_info=True)
    finally:
        _compacting_threads.discard(thread_id)


@router.post("/threads/{thread_id}/messages/stream")
async def send_message_stream(
    thread_id: UUID,
    body: SendMessageRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Send message with SSE progress + token usage events."""
    caller_company_id = await get_client_company_id(current_user)
    # Don't 404 on None — collaborators (individuals invited to another user's
    # thread) may have no company of their own.
    thread = await doc_svc.get_thread(thread_id, caller_company_id, user_id=current_user.id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    # Use the thread's actual company for all downstream operations (AI profile,
    # token budget, etc.) so collaborators don't accidentally scope ops to their
    # own (possibly absent) company.
    company_id = thread["company_id"]

    if thread["status"] == "finalized":
        raise HTTPException(status_code=400, detail="Cannot send messages to a finalized thread")

    if thread["status"] == "archived":
        raise HTTPException(status_code=400, detail="Cannot send messages to an archived thread")

    if current_user.role != "admin":
        await token_budget_service.check_token_budget(company_id)

    # Check token quota
    quota = await doc_svc.check_token_quota(current_user.id, company_id)
    if not quota["allowed"]:
        raise HTTPException(
            status_code=429,
            detail=f"Token limit reached ({quota['used']:,}/{quota['limit']:,} tokens used). Resets at {quota['resets_at']}.",
        )

    # Normalize & persist attachment URLs on the user message metadata. Client
    # uploads images separately (stored in currentState["images"]) and sends the
    # URLs here so they become part of the message itself — visible in the
    # bubble and passed to the AI as multimodal parts.
    attach_urls: list[str] = []
    if body.image_urls:
        attach_urls = [u for u in body.image_urls if isinstance(u, str) and u]
    user_meta = {"attachments": [{"url": u, "kind": "image"} for u in attach_urls]} if attach_urls else None

    # Save user message before streaming
    user_msg = await doc_svc.add_message(thread_id, "user", body.content, metadata=user_meta)

    # Once the attachments are persisted on the message itself, clear them from
    # thread state so they don't leak into the next send or get re-consumed by
    # the presentation skill.
    if attach_urls:
        try:
            await doc_svc.apply_update(thread_id, {"images": []}, diff_summary="Consumed inline chat attachments")
        except Exception:
            logger.warning("Failed to clear thread images after attaching to message %s", thread_id, exc_info=True)

    # Fetch message history + company profile + context summary in parallel
    messages, profile, (context_summary, summary_at_count) = await asyncio.gather(
        doc_svc.get_thread_messages(thread_id, limit=20),
        doc_svc.get_company_profile_for_ai(company_id),
        doc_svc.get_context_summary(thread_id),
    )
    msg_dicts = []
    for m in messages:
        entry = {"role": m["role"], "content": m["content"]}
        meta = m.get("metadata")
        if isinstance(meta, str):
            try:
                meta = json.loads(meta)
            except Exception:
                meta = None
        if isinstance(meta, dict):
            atts = meta.get("attachments") or []
            urls = [a.get("url") for a in atts if isinstance(a, dict) and a.get("url")]
            if urls:
                entry["image_urls"] = urls
        msg_dicts.append(entry)

    # Inject selected slide content into the AI-facing message (not saved to DB)
    _inject_slide_context(msg_dicts, thread["current_state"], body.slide_index)

    # Pre-fetch any image attachment bytes concurrently off the event loop so
    # the prompt builder (which runs in a thread pool) doesn't block on I/O.
    from ..services.matcha_work_ai import fetch_image_parts_for_messages
    await fetch_image_parts_for_messages(msg_dicts)

    ai_provider = get_ai_provider()
    ctx = _build_company_context(profile)

    # Inject project file attachments metadata
    if thread.get("project_id"):
        from ..services import project_file_service
        pfiles = await project_file_service.list_project_files(thread["project_id"])
        if pfiles:
            listing = "\n".join(f"- {f['filename']} ({f['content_type']}, {f['file_size']:,} bytes)" for f in pfiles)
            ctx += f"\n\n=== PROJECT ATTACHMENTS ===\nThe user has attached these files to the project. Reference them when relevant:\n{listing}\n"

    # Inject recruiting project context so AI generates posting sections in the right project
    ctx = await _inject_recruiting_project_context(ctx, thread, thread["current_state"])

    # Node/compliance context is built inside event_stream() so we can yield status events

    async def event_stream():
        nonlocal ctx
        compliance_result: ComplianceContextResult | None = None
        try:
            # Build mode-specific context with status updates
            if thread.get("node_mode"):
                yield _sse_data({"type": "status", "message": "Loading internal company data..."})
                node_ctx = await build_node_context(company_id)
                ctx += "\n\n" + node_ctx

            if thread.get("compliance_mode"):
                yield _sse_data({"type": "status", "message": "Loading compliance data for your locations..."})
                compliance_result = await build_compliance_context(company_id)
                compliance_ctx = compliance_result.context_text
                cat_count = compliance_ctx.count("Decision path:")
                trigger_count = compliance_ctx.count("[trigger:")
                loc_count = compliance_ctx.count("FACILITY PROFILE")
                if cat_count > 0:
                    parts = [f"{cat_count} regulatory categories across {loc_count} location{'s' if loc_count != 1 else ''}"]
                    if trigger_count > 0:
                        parts.append(f"{trigger_count} triggered requirement{'s' if trigger_count != 1 else ''}")
                    yield _sse_data({"type": "status", "message": f"Found {' with '.join(parts)} — building reasoning chains..."})
                elif compliance_ctx.count("legacy format") > 0:
                    yield _sse_data({"type": "status", "message": "Loaded compliance data (legacy format) — cross-referencing..."})
                else:
                    yield _sse_data({"type": "status", "message": "No compliance data found — will suggest running a check..."})
                ctx += "\n\n" + compliance_ctx

                # RAG augmentation — find requirements most relevant to the question
                yield _sse_data({"type": "status", "message": "Searching relevant regulations..."})
                rag_ctx = await _get_rag_context(body.content, company_id)
                if rag_ctx:
                    ctx += "\n\n=== RELEVANT REGULATIONS (semantic search) ===\n" + rag_ctx

            # Payer mode — build payer prompt inside stream for status events
            stream_payer_prompt = None
            stream_payer_sources: list[dict] = []
            if thread.get("payer_mode"):
                yield _sse_data({"type": "status", "message": "Searching payer coverage data..."})
                try:
                    import os as _os2
                    from ...core.services.embedding_service import EmbeddingService as _ES2
                    from ...core.services.payer_policy_rag import PayerPolicyRAGService as _PRAG2
                    from ...config import get_settings as _gs2
                    from ..services.matcha_work_ai import PAYER_MODE_SYSTEM_PROMPT as _PMSP
                    from datetime import date as _d2

                    _ak2 = _os2.getenv("GEMINI_API_KEY") or _gs2().gemini_api_key
                    if _ak2 and body.content:
                        _e2 = _ES2(api_key=_ak2)
                        _r2 = _PRAG2(_e2)
                        async with get_connection() as _pc2:
                            _pctx, stream_payer_sources = await _r2.get_context_for_query(
                                query=body.content, conn=_pc2,
                                company_id=company_id, max_tokens=6000,
                            )
                        cn2 = profile.get("name", "your company")
                        stream_payer_prompt = _PMSP.format(
                            company_name=cn2,
                            today=_d2.today().isoformat(),
                            payer_context=_pctx or "No matching payer policy data found.",
                        )
                        if stream_payer_sources:
                            yield _sse_data({"type": "status", "message": f"Found {len(stream_payer_sources)} relevant payer policies"})
                except Exception as _pe:
                    logger.warning("Stream payer context failed: %s", _pe)

            estimated_usage = await ai_provider.estimate_usage(msg_dicts, thread["current_state"], company_context=ctx, slide_index=body.slide_index)
            yield _sse_data(
                {
                    "type": "usage",
                    "data": {
                        **estimated_usage,
                        "stage": "estimate",
                    },
                }
            )

            yield _sse_data({"type": "status", "message": "Generating response..."})
            import time as _time
            _t0 = _time.monotonic()
            # Run generation as a background task and emit keepalives every 15 s
            # so proxies with short read-timeouts (e.g. nginx default 60 s) don't
            # close the SSE connection while we wait for the AI to finish.
            stream_project_meta = await _fetch_project_meta(thread.get("project_id"))
            stream_blog_mode_state = _blog_mode_state_from_meta(stream_project_meta)
            _ai_task = asyncio.create_task(ai_provider.generate(
                msg_dicts, thread["current_state"], company_context=ctx,
                slide_index=body.slide_index, context_summary=context_summary,
                payer_mode_prompt=stream_payer_prompt,
                model_override=body.model,
                company_id=str(company_id),
                compliance_mode=bool(thread.get("compliance_mode")),
                payer_mode=bool(thread.get("payer_mode")),
                node_mode=bool(thread.get("node_mode")),
                blog_mode_state=stream_blog_mode_state,
            ))
            while True:
                done, _ = await asyncio.wait({_ai_task}, timeout=15.0)
                if done:
                    break
                yield _sse_data({"type": "keepalive"})
            ai_resp = await _ai_task
            logger.info("[TIMING] AI generate took %.2fs for thread %s", _time.monotonic() - _t0, thread_id)
            _scope_slide_update(ai_resp, thread["current_state"], body.slide_index)

            current_version = thread["version"]
            (
                current_state,
                current_version,
                pdf_url,
                changed,
                assistant_reply_text,
            ) = await _apply_ai_updates_and_operations(
                thread_id=thread_id,
                company_id=company_id,
                ai_resp=ai_resp,
                current_state=thread["current_state"],
                current_version=current_version,
                user_message=body.content,
                current_user_id=current_user.id,
                project_id=thread.get("project_id"),
                project_meta=stream_project_meta,
            )

            # Build metadata from compliance reasoning chains + payer sources
            msg_metadata = _build_compliance_metadata(compliance_result, ai_resp)
            if stream_payer_sources:
                if msg_metadata is None:
                    msg_metadata = {}
                msg_metadata["payer_sources"] = stream_payer_sources

            # Cross-reference affected employees + detect policy gaps when both node + compliance are on
            if thread.get("node_mode") and thread.get("compliance_mode") and msg_metadata:
                if msg_metadata.get("referenced_locations"):
                    affected = await _get_affected_employees(company_id, msg_metadata)
                    if affected:
                        msg_metadata["affected_employees"] = affected
                gaps = await _detect_compliance_gaps(company_id, msg_metadata)
                if gaps:
                    msg_metadata["compliance_gaps"] = gaps

            # Annotate reply with change summary for conversation continuity
            if changed and ai_resp.structured_update and isinstance(ai_resp.structured_update, dict):
                update_slides = ai_resp.structured_update.get("slides")
                if update_slides and body.slide_index is not None and 0 <= body.slide_index < len(update_slides):
                    changed_slide = update_slides[body.slide_index]
                    if isinstance(changed_slide, dict):
                        n_bullets = len(changed_slide.get("bullets", []))
                        change_note = f"\n\n[Applied changes to Slide {body.slide_index + 1}: title=\"{changed_slide.get('title', '')}\", {n_bullets} bullets]"
                        assistant_reply_text += change_note

            # Save assistant message
            assistant_msg = await doc_svc.add_message(
                thread_id,
                "assistant",
                assistant_reply_text,
                version_created=current_version if changed else None,
                metadata=msg_metadata,
            )

            # Broadcast new messages to collaborators via WS — fire-and-forget so
            # a CancelledError inside the lock doesn't kill the SSE generator before
            # the complete event is sent.
            try:
                from .thread_ws import thread_manager
                user_msg_dict = _row_to_message(user_msg).model_dump(mode="json")
                assistant_msg_dict = _row_to_message(assistant_msg).model_dump(mode="json")
                asyncio.create_task(
                    thread_manager.broadcast_new_message(
                        str(thread_id), [user_msg_dict, assistant_msg_dict], exclude_user=current_user.id
                    )
                )
            except Exception:
                logger.warning("Thread WS broadcast failed for thread %s", thread_id)

            # Escalate low-confidence queries for human review
            if should_escalate(ai_resp):
                try:
                    await create_escalation(
                        company_id=company_id,
                        thread_id=thread_id,
                        user_message_id=user_msg["id"],
                        assistant_message_id=assistant_msg["id"],
                        user_query=body.content,
                        ai_resp=ai_resp,
                    )
                except Exception:
                    logger.exception("Failed to create escalation for thread %s", thread_id)

            final_usage = ai_resp.token_usage or estimated_usage
            stream_cost = calculate_call_cost(
                model=str((final_usage or {}).get("model") or "unknown"),
                prompt_tokens=(final_usage or {}).get("prompt_tokens"),
                completion_tokens=(final_usage or {}).get("completion_tokens"),
            )
            if final_usage is not None:
                final_usage["cost_dollars"] = float(stream_cost)

            try:
                await doc_svc.log_token_usage_event(
                    company_id=company_id,
                    user_id=current_user.id,
                    thread_id=thread_id,
                    token_usage=final_usage,
                    operation="send_message",
                    cost_dollars=float(stream_cost),
                )
            except Exception as e:
                logger.warning("Failed to log Matcha Work token usage for thread %s: %s", thread_id, e)

            if current_user.role != "admin":
                total_tokens = (final_usage or {}).get("total_tokens") or 0
                if total_tokens > 0:
                    try:
                        async with get_connection() as conn:
                            async with conn.transaction():
                                await token_budget_service.deduct_tokens(conn, company_id, total_tokens)
                    except HTTPException:
                        logger.warning("Token budget exhausted during stream deduction for thread %s", thread_id)
                    except Exception as exc:
                        logger.warning("Failed to deduct tokens for thread %s: %s", thread_id, exc)

            if final_usage:
                yield _sse_data(
                    {
                        "type": "usage",
                        "data": {
                            **final_usage,
                            "stage": "final",
                        },
                    }
                )

            response = SendMessageResponse(
                user_message=_row_to_message(user_msg),
                assistant_message=_row_to_message(assistant_msg),
                current_state=current_state,
                version=current_version,
                task_type=_infer_skill_from_state(current_state),
                pdf_url=pdf_url,
                token_usage=final_usage,
            )

            yield _sse_data({"type": "complete", "data": response.model_dump(mode="json")})

            # Trigger compaction in the background if needed
            asyncio.create_task(_maybe_compact(thread_id, ai_provider, summary_at_count))
        except BaseException as e:
            logger.error("Matcha Work stream failed for thread %s: %s (%s)", thread_id, e, type(e).__name__, exc_info=True)
            try:
                yield _sse_data(
                    {
                        "type": "error",
                        "message": "Failed to process message. Please try again.",
                    }
                )
            except Exception:
                pass
            if not isinstance(e, Exception):
                raise
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.get("/usage/summary", response_model=UsageSummaryResponse)
async def get_usage_summary(
    response: Response,
    period_days: int = Query(30, ge=1, le=365),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Get Matcha Work token usage totals for the current user, grouped by model."""
    response.headers["Cache-Control"] = "private, max-age=300"
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        return UsageSummaryResponse(
            period_days=period_days,
            generated_at=datetime.now(timezone.utc),
            totals={
                "prompt_tokens": 0,
                "completion_tokens": 0,
                "total_tokens": 0,
                "operation_count": 0,
                "estimated_operations": 0,
            },
            by_model=[],
        )

    summary = await doc_svc.get_token_usage_summary(company_id, current_user.id, period_days)
    return UsageSummaryResponse(**summary)


@router.get("/threads/{thread_id}/versions", response_model=list[DocumentVersionResponse])
async def list_versions(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List all document versions for a thread."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    versions = await doc_svc.list_versions(thread_id)
    return [DocumentVersionResponse(**v) for v in versions]


@router.post("/threads/{thread_id}/revert", response_model=SendMessageResponse)
async def revert_thread(
    thread_id: UUID,
    body: RevertRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Revert to a historical version (creates a new version with old state)."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    if thread["status"] == "finalized":
        raise HTTPException(status_code=400, detail="Cannot revert a finalized thread")
    if thread["status"] == "archived":
        raise HTTPException(status_code=400, detail="Cannot revert an archived thread")

    try:
        result = await doc_svc.revert_to_version(thread_id, body.version)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

    new_version = result["version"]
    new_state = result["current_state"]

    system_msg = await doc_svc.add_message(
        thread_id,
        "system",
        f"Document reverted to version {body.version}.",
        version_created=new_version,
    )

    assistant_msg = await doc_svc.add_message(
        thread_id,
        "assistant",
        f"I've reverted the document to version {body.version}. You can continue editing from there.",
        version_created=new_version,
    )

    pdf_url = None
    if _infer_skill_from_state(new_state) == "offer_letter":
        pdf_url = await doc_svc.generate_pdf(
            new_state,
            thread_id,
            new_version,
            is_draft=True,
            company_id=company_id,
        )

    return SendMessageResponse(
        user_message=_row_to_message(system_msg),
        assistant_message=_row_to_message(assistant_msg),
        current_state=new_state,
        version=new_version,
        pdf_url=pdf_url,
    )


@router.post("/threads/{thread_id}/finalize", response_model=FinalizeResponse)
async def finalize_thread(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Finalize the thread — lock status, generate final PDF without watermark."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    try:
        result = await doc_svc.finalize_thread(thread_id, company_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    return FinalizeResponse(**result)


@router.post("/threads/{thread_id}/save-draft", response_model=SaveDraftResponse)
async def save_draft(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Save the current Matcha Work state into offer_letters as a draft."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    if _infer_skill_from_state(thread.get("current_state") or {}) != "offer_letter":
        raise HTTPException(status_code=400, detail="Save draft is only available for offer letters")

    try:
        result = await doc_svc.save_offer_letter_draft(thread_id, company_id)
    except ValueError as e:
        detail = str(e)
        if detail == "Thread not found":
            raise HTTPException(status_code=404, detail=detail)
        raise HTTPException(status_code=400, detail=detail)

    return SaveDraftResponse(**result)


@router.get("/threads/{thread_id}/pdf")
async def get_pdf(
    thread_id: UUID,
    version: Optional[int] = Query(None),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Get or generate the PDF for a thread (optionally a specific version)."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    if _infer_skill_from_state(thread.get("current_state") or {}) != "offer_letter":
        raise HTTPException(status_code=400, detail="PDF preview is only available for offer letters")

    target_version = version if version is not None else thread["version"]
    state = thread["current_state"]

    # If requesting a historical version, load that state
    if version is not None and version != thread["version"]:
        versions = await doc_svc.list_versions(thread_id)
        ver_map = {v["version"]: v for v in versions}
        if target_version not in ver_map:
            raise HTTPException(status_code=404, detail=f"Version {target_version} not found")
        state = ver_map[target_version]["state_json"]

    is_draft = thread["status"] != "finalized"
    pdf_url = await doc_svc.generate_pdf(
        state,
        thread_id,
        target_version,
        is_draft=is_draft,
        company_id=company_id,
    )

    if pdf_url is None:
        raise HTTPException(status_code=503, detail="PDF generation unavailable")

    return {"pdf_url": pdf_url, "version": target_version}


@router.get("/threads/{thread_id}/pdf/proxy")
async def proxy_pdf(
    thread_id: UUID,
    version: Optional[int] = Query(None),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Stream PDF bytes for same-origin iframe embedding (avoids cross-origin iframe blocking)."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    if _infer_skill_from_state(thread.get("current_state") or {}) != "offer_letter":
        raise HTTPException(status_code=400, detail="PDF preview is only available for offer letters")

    target_version = version if version is not None else thread["version"]
    state = thread["current_state"]

    if version is not None and version != thread["version"]:
        versions = await doc_svc.list_versions(thread_id)
        ver_map = {v["version"]: v for v in versions}
        if target_version not in ver_map:
            raise HTTPException(status_code=404, detail=f"Version {target_version} not found")
        state = ver_map[target_version]["state_json"]

    is_draft = thread["status"] != "finalized"
    pdf_url = await doc_svc.generate_pdf(
        state,
        thread_id,
        target_version,
        is_draft=is_draft,
        company_id=company_id,
    )

    if pdf_url is None:
        raise HTTPException(status_code=503, detail="PDF generation unavailable")

    try:
        pdf_bytes = await get_storage().download_file(pdf_url)
    except Exception as e:
        logger.error("Failed to fetch PDF for proxy: %s", e)
        raise HTTPException(status_code=503, detail="Failed to load PDF")

    safe_title = re.sub(r"[^\w\s-]", "", thread.get("title") or "offer-letter").strip().replace(" ", "-") or "offer-letter"
    return StreamingResponse(
        iter([pdf_bytes]),
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{safe_title}.pdf"'},
    )


@router.delete("/threads/{thread_id}", status_code=204)
async def archive_thread(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Archive (soft-delete) a thread."""
    from ...database import get_connection

    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    async with get_connection() as conn:
        result = await conn.execute(
            """
            UPDATE mw_threads
            SET status='archived', updated_at=NOW()
            WHERE id=$1 AND company_id=$2 AND status != 'archived'
            """,
            thread_id,
            company_id,
        )
    if result == "UPDATE 0":
        raise HTTPException(status_code=404, detail="Thread not found")
    await doc_svc.sync_element_record(thread_id)


@router.get(
    "/threads/{thread_id}/review-requests",
    response_model=list[ReviewRequestStatus],
)
async def list_review_requests(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List review request status rows for a review thread."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    try:
        rows = await doc_svc.list_review_requests(thread_id, company_id)
    except ValueError as e:
        detail = str(e)
        if detail == "Thread not found":
            raise HTTPException(status_code=404, detail=detail)
        raise HTTPException(status_code=400, detail=detail)

    return [ReviewRequestStatus(**row) for row in rows]


@router.post(
    "/threads/{thread_id}/review-requests/send",
    response_model=SendReviewRequestsResponse,
)
async def send_review_requests(
    thread_id: UUID,
    body: SendReviewRequestsRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """
    Send anonymous review request links to recipient emails and track expected responses.
    """
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    try:
        result = await doc_svc.send_review_requests(
            thread_id=thread_id,
            company_id=company_id,
            recipient_emails=body.recipient_emails,
            custom_message=body.custom_message,
        )
    except ValueError as e:
        detail = str(e)
        if detail == "Thread not found":
            raise HTTPException(status_code=404, detail=detail)
        raise HTTPException(status_code=400, detail=detail)

    return SendReviewRequestsResponse(**result)


@router.post(
    "/threads/{thread_id}/handbook/send-signatures",
    response_model=SendHandbookSignaturesResponse,
)
async def send_handbook_signatures(
    thread_id: UUID,
    body: SendHandbookSignaturesRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Send handbook acknowledgement signatures from a workbook thread."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    if thread["status"] == "archived":
        raise HTTPException(status_code=400, detail="Cannot send handbook signatures for an archived thread")
    if _infer_skill_from_state(thread.get("current_state") or {}) != "workbook":
        raise HTTPException(status_code=400, detail="Handbook signatures are only available for workbook threads")

    employee_ids = [str(employee_id) for employee_id in body.employee_ids] if body.employee_ids else None

    try:
        distributed = await HandbookService.distribute_to_employees(
            handbook_id=str(body.handbook_id),
            company_id=str(company_id),
            distributed_by=str(current_user.id),
            employee_ids=employee_ids,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if distributed is None:
        raise HTTPException(status_code=404, detail="Handbook not found")

    payload = distributed.model_dump() if hasattr(distributed, "model_dump") else dict(distributed)
    return SendHandbookSignaturesResponse(**payload)


@router.post(
    "/threads/{thread_id}/presentation/generate",
    response_model=GeneratePresentationResponse,
)
async def generate_workbook_presentation(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Generate slide-ready presentation content from workbook sections."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    try:
        result = await doc_svc.generate_workbook_presentation(
            thread_id=thread_id,
            company_id=company_id,
        )
    except ValueError as e:
        detail = str(e)
        if detail == "Thread not found":
            raise HTTPException(status_code=404, detail=detail)
        raise HTTPException(status_code=400, detail=detail)

    return GeneratePresentationResponse(**result)


@router.get("/threads/{thread_id}/presentation/pdf")
async def get_presentation_pdf(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Return the presentation PDF URL, generating it on demand."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    state = await doc_svc.ensure_matcha_work_thread_storage_scope(
        thread_id,
        company_id,
        thread.get("current_state") or {},
    )
    if _infer_skill_from_state(state) not in ("presentation", "workbook"):
        raise HTTPException(status_code=400, detail="Presentation PDF is only available for presentation and workbook threads")

    version = int(thread.get("version") or 0)
    # For workbook threads, use the nested presentation; for standalone presentations, use top-level state
    if _infer_skill_from_state(state) == "workbook":
        pres = state.get("presentation")
        if not pres or not pres.get("slides"):
            raise HTTPException(status_code=400, detail="No presentation slides found. Generate a presentation first.")
        pdf_state = {
            "presentation_title": pres.get("title"),
            "subtitle": pres.get("subtitle"),
            "slides": pres.get("slides", []),
            "cover_image_url": pres.get("cover_image_url"),
        }
    else:
        pdf_state = state

    pdf_url = await doc_svc.generate_presentation_pdf(
        pdf_state,
        thread_id,
        version,
        company_id=company_id,
    )
    if not pdf_url:
        raise HTTPException(status_code=500, detail="Failed to generate presentation PDF")

    return {"pdf_url": pdf_url}


@router.patch("/threads/{thread_id}", response_model=ThreadListItem)
async def update_thread_title(
    thread_id: UUID,
    body: UpdateTitleRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Update thread title."""
    from ...database import get_connection

    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    async with get_connection() as conn:
        row = await conn.fetchrow(
            """
            UPDATE mw_threads
            SET title=$1, updated_at=NOW()
            WHERE id=$2 AND company_id=$3
            RETURNING id, title, status, version, is_pinned, node_mode, compliance_mode, created_at, updated_at, current_state
            """,
            body.title,
            thread_id,
            company_id,
        )
    if row is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    await doc_svc.sync_element_record(thread_id)

    row_dict = dict(row)
    current_state = _json_object(row_dict.pop("current_state", {}))
    return ThreadListItem(
        **{
            **row_dict,
            "task_type": _infer_skill_from_state(current_state),
        }
    )


@router.post("/threads/{thread_id}/pin", response_model=ThreadListItem)
async def set_thread_pin(
    thread_id: UUID,
    body: PinThreadRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Pin or unpin a thread for faster access in chat history."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    row = await doc_svc.set_thread_pinned(thread_id, company_id, body.is_pinned)
    if row is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    return ThreadListItem(**row)


@router.post("/threads/{thread_id}/node-mode", response_model=ThreadListItem)
async def set_thread_node_mode(
    thread_id: UUID,
    body: NodeModeRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Enable or disable Node mode (internal data search) for a thread."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    row = await doc_svc.set_thread_node_mode(thread_id, company_id, body.node_mode)
    if row is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    row_dict = dict(row)
    current_state = _json_object(row_dict.pop("current_state", {}))
    return ThreadListItem(
        **{
            **row_dict,
            "task_type": _infer_skill_from_state(current_state),
        }
    )


@router.post("/threads/{thread_id}/compliance-mode", response_model=ThreadListItem)
async def set_thread_compliance_mode(
    thread_id: UUID,
    body: ComplianceModeRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Enable or disable Compliance mode (jurisdiction requirements context) for a thread."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    row = await doc_svc.set_thread_compliance_mode(thread_id, company_id, body.compliance_mode)
    if row is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    row_dict = dict(row)
    current_state = _json_object(row_dict.pop("current_state", {}))
    return ThreadListItem(
        **{
            **row_dict,
            "task_type": _infer_skill_from_state(current_state),
        }
    )


@router.post("/threads/{thread_id}/payer-mode", response_model=ThreadListItem)
async def set_thread_payer_mode(
    thread_id: UUID,
    body: PayerModeRequest,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Enable or disable Payer mode (medical policy coverage data) for a thread."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    row = await doc_svc.set_thread_payer_mode(thread_id, company_id, body.payer_mode)
    if row is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    row_dict = dict(row)
    current_state = _json_object(row_dict.pop("current_state", {}))
    return ThreadListItem(
        **{
            **row_dict,
            "task_type": _infer_skill_from_state(current_state),
        }
    )


@public_router.get("/review-requests/{token}", response_model=PublicReviewRequestResponse)
async def get_public_review_request(token: str):
    """Get a public review-request payload by one-time token."""
    row = await doc_svc.get_public_review_request(token)
    if row is None:
        raise HTTPException(status_code=404, detail="Review request not found")
    return PublicReviewRequestResponse(**row)


@public_router.post(
    "/review-requests/{token}/submit",
    response_model=PublicReviewSubmitResponse,
)
async def submit_public_review_request(
    token: str,
    body: PublicReviewSubmitRequest,
):
    """Submit an anonymous review response via public token link."""
    try:
        result = await doc_svc.submit_public_review_request(
            token=token,
            feedback=body.feedback,
            rating=body.rating,
        )
    except ValueError as e:
        detail = str(e)
        if detail == "Review request not found":
            raise HTTPException(status_code=404, detail=detail)
        if detail == "Review response already submitted":
            raise HTTPException(status_code=409, detail=detail)
        raise HTTPException(status_code=400, detail=detail)

    return PublicReviewSubmitResponse(**result)


# ── Task Board endpoints ──


@router.get("/tasks")
async def list_tasks(
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Combined task board: auto-populated items + manual tasks + dismissals."""
    from .dashboard import _UPCOMING_SOURCES, _apply_company_filter, _severity_from_days, UpcomingItem
    from datetime import date as _date, timedelta as _td

    company_id = await get_client_company_id(current_user)
    today = _date.today()
    lookahead = today + _td(days=90)

    # 1. Auto-populated items from upcoming sources
    auto_items = []
    async with get_connection() as conn:
        import asyncpg as _asyncpg
        for source in _UPCOMING_SOURCES:
            try:
                sql = _apply_company_filter(source["sql"], company_id)
                uses_p1 = "$1" in sql
                uses_p2 = "$2" in sql
                if uses_p1 and uses_p2:
                    rows = await conn.fetch(sql, company_id, lookahead)
                elif uses_p1:
                    rows = await conn.fetch(sql, company_id)
                elif uses_p2:
                    rows = await conn.fetch(sql, lookahead)
                else:
                    rows = await conn.fetch(sql)
            except (_asyncpg.UndefinedTableError, _asyncpg.UndefinedColumnError):
                continue
            except Exception:
                continue
            for row in rows:
                deadline = row["deadline"]
                if deadline is None:
                    continue
                days_until = (deadline - today).days
                link = source["link"]
                row_id = row.get("id")
                if row_id and "{id}" in link:
                    link = link.replace("{id}", row_id)
                auto_items.append({
                    "category": source["category"],
                    "source_id": row.get("id") or "",
                    "title": row["title"] or source["category"].title(),
                    "subtitle": row.get("subtitle"),
                    "date": str(deadline),
                    "days_until": days_until,
                    "severity": _severity_from_days(days_until),
                    "link": link,
                })

    auto_items.sort(key=lambda x: x["days_until"])

    # 2. Manual tasks
    manual_items = []
    async with get_connection() as conn:
        try:
            rows = await conn.fetch(
                """
                SELECT id, title, description, due_date, horizon, priority, status,
                       completed_at, link, category, created_at, updated_at
                FROM mw_tasks
                WHERE company_id = $1 AND status != 'cancelled'
                ORDER BY
                    CASE WHEN status = 'completed' THEN 1 ELSE 0 END,
                    due_date ASC NULLS LAST,
                    created_at DESC
                """,
                company_id,
            )
            for r in rows:
                d = dict(r)
                d["id"] = str(d["id"])
                d["source"] = "manual"
                if d["due_date"]:
                    d["days_until"] = (d["due_date"] - today).days
                    d["date"] = str(d["due_date"])
                else:
                    d["days_until"] = None
                    d["date"] = None
                manual_items.append(d)
        except Exception:
            pass  # table may not exist yet

    # 3. Dismissed IDs
    dismissed_ids = []
    async with get_connection() as conn:
        try:
            rows = await conn.fetch(
                "SELECT source_category, source_id FROM mw_task_dismissals WHERE user_id = $1",
                current_user.id,
            )
            dismissed_ids = [f"{r['source_category']}:{r['source_id']}" for r in rows]
        except Exception:
            pass

    return {
        "auto_items": auto_items,
        "manual_items": manual_items,
        "dismissed_ids": dismissed_ids,
        "total": len(auto_items) + len(manual_items),
    }


@router.post("/tasks", status_code=201)
async def create_task(
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Create a manual task."""
    from datetime import date as _date

    company_id = await get_client_company_id(current_user)
    if not company_id:
        raise HTTPException(status_code=400, detail="No company associated")

    title = body.get("title", "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title is required")

    due_date = body.get("due_date")
    if due_date and isinstance(due_date, str):
        due_date = _date.fromisoformat(due_date)

    horizon = body.get("horizon")
    priority = body.get("priority", "medium")
    if priority not in ("critical", "high", "medium", "low"):
        priority = "medium"

    async with get_connection() as conn:
        row = await conn.fetchrow(
            """
            INSERT INTO mw_tasks (company_id, created_by, title, description, due_date, horizon, priority, link)
            VALUES ($1, $2, $3, $4, $5, $6, $7, $8)
            RETURNING *
            """,
            company_id, current_user.id, title,
            body.get("description"), due_date, horizon, priority, body.get("link"),
        )
    d = dict(row)
    d["id"] = str(d["id"])
    d["company_id"] = str(d["company_id"])
    d["created_by"] = str(d["created_by"])
    d["source"] = "manual"
    if d.get("due_date"):
        d["days_until"] = (d["due_date"] - _date.today()).days
        d["date"] = str(d["due_date"])
    else:
        d["days_until"] = None
        d["date"] = None
    return d


@router.patch("/tasks/{task_id}")
async def update_task(
    task_id: UUID,
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Update a manual task."""
    company_id = await get_client_company_id(current_user)

    allowed = {"title", "description", "due_date", "horizon", "priority", "status", "link"}
    sets = []
    vals = []
    idx = 1
    for k, v in body.items():
        if k in allowed:
            if k == "due_date" and isinstance(v, str):
                from datetime import date as _date
                v = _date.fromisoformat(v)
            sets.append(f"{k} = ${idx}")
            vals.append(v)
            idx += 1

    if not sets:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    # Auto-fill completed_at
    if body.get("status") == "completed":
        sets.append(f"completed_at = ${idx}")
        vals.append(datetime.now(timezone.utc))
        idx += 1
    elif body.get("status") == "pending":
        sets.append(f"completed_at = ${idx}")
        vals.append(None)
        idx += 1

    sets.append(f"updated_at = NOW()")
    vals.extend([task_id, company_id])

    async with get_connection() as conn:
        row = await conn.fetchrow(
            f"UPDATE mw_tasks SET {', '.join(sets)} WHERE id = ${idx} AND company_id = ${idx + 1} RETURNING *",
            *vals,
        )
    if not row:
        raise HTTPException(status_code=404, detail="Task not found")
    d = dict(row)
    d["id"] = str(d["id"])
    d["source"] = "manual"
    return d


@router.delete("/tasks/{task_id}")
async def delete_task(
    task_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Cancel a manual task."""
    company_id = await get_client_company_id(current_user)
    async with get_connection() as conn:
        await conn.execute(
            "UPDATE mw_tasks SET status = 'cancelled', updated_at = NOW() WHERE id = $1 AND company_id = $2",
            task_id, company_id,
        )
    return {"status": "cancelled"}


@router.post("/tasks/dismiss")
async def dismiss_auto_task(
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Dismiss an auto-populated task item."""
    cat = body.get("source_category", "")
    sid = body.get("source_id", "")
    if not cat or not sid:
        raise HTTPException(status_code=400, detail="source_category and source_id required")
    async with get_connection() as conn:
        await conn.execute(
            """
            INSERT INTO mw_task_dismissals (user_id, source_category, source_id)
            VALUES ($1, $2, $3)
            ON CONFLICT (user_id, source_category, source_id) DO NOTHING
            """,
            current_user.id, cat, sid,
        )
    return {"status": "dismissed"}


@router.delete("/tasks/dismiss")
async def undismiss_auto_task(
    body: dict,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Un-dismiss a previously dismissed auto-populated task."""
    cat = body.get("source_category", "")
    sid = body.get("source_id", "")
    if not cat or not sid:
        raise HTTPException(status_code=400, detail="source_category and source_id required")
    async with get_connection() as conn:
        await conn.execute(
            "DELETE FROM mw_task_dismissals WHERE user_id = $1 AND source_category = $2 AND source_id = $3",
            current_user.id, cat, sid,
        )
    return {"status": "restored"}


# ── Language Tutor Voice Sessions ──────────────────────────────────────────


@router.post("/threads/{thread_id}/tutor/start")
async def start_tutor_voice_session(
    thread_id: UUID,
    body: dict = Body(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Start a Gemini Live language tutor voice session linked to a matcha-work thread."""
    from ...core.services.auth import create_interview_ws_token

    language = body.get("language", "en")
    if language not in ("en", "es-mx", "fr"):
        raise HTTPException(status_code=400, detail="Language must be 'en', 'es-mx', or 'fr'")
    duration_minutes = body.get("duration_minutes", 5)
    if duration_minutes not in (0.33, 2, 5, 8):
        raise HTTPException(status_code=400, detail="Duration must be 0.33 (20s test), 2, 5, or 8 minutes")

    company_id = await get_client_company_id(current_user)

    async with get_connection() as conn:
        # Verify thread exists and belongs to user
        thread = await conn.fetchrow(
            "SELECT id, current_state FROM mw_threads WHERE id = $1 AND company_id IS NOT DISTINCT FROM $2",
            thread_id, company_id,
        )
        if not thread:
            raise HTTPException(status_code=404, detail="Thread not found")

        # Create interview record (same as POST /tutor/sessions with mode=language_test)
        row = await conn.fetchrow(
            """
            INSERT INTO interviews (company_id, interviewer_name, interviewer_role, interview_type, status)
            VALUES (NULL, $1, $2, 'tutor_language', 'pending')
            RETURNING id
            """,
            current_user.email,
            language,
        )
        interview_id = row["id"]

        # Store interview_id in thread current_state
        raw_state = thread["current_state"]
        if isinstance(raw_state, str):
            current_state = json.loads(raw_state) if raw_state else {}
        elif isinstance(raw_state, dict):
            current_state = dict(raw_state)
        else:
            current_state = {}
        current_state["language_tutor"] = {
            "interview_id": str(interview_id),
            "language": language,
            "status": "active",
        }
        await conn.execute(
            "UPDATE mw_threads SET current_state = $1, updated_at = NOW() WHERE id = $2",
            json.dumps(current_state), thread_id,
        )

    duration_seconds = int(duration_minutes * 60)
    return {
        "interview_id": str(interview_id),
        "websocket_url": f"/api/ws/interview/{interview_id}",
        "ws_auth_token": create_interview_ws_token(interview_id),
        "max_session_duration_seconds": duration_seconds,
    }


@router.get("/threads/{thread_id}/tutor/status")
async def get_tutor_voice_status(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Poll language tutor session status and analysis results.

    Runs analysis inline (no Celery) on first poll after session ends.
    """
    company_id = await get_client_company_id(current_user)

    async with get_connection() as conn:
        thread = await conn.fetchrow(
            "SELECT id, current_state FROM mw_threads WHERE id = $1 AND company_id IS NOT DISTINCT FROM $2",
            thread_id, company_id,
        )
        if not thread:
            raise HTTPException(status_code=404, detail="Thread not found")

        raw_state = thread["current_state"]
        if isinstance(raw_state, str):
            current_state = json.loads(raw_state) if raw_state else {}
        elif isinstance(raw_state, dict):
            current_state = dict(raw_state)
        else:
            current_state = {}
        tutor_state = current_state.get("language_tutor")
        if not tutor_state or not tutor_state.get("interview_id"):
            raise HTTPException(status_code=404, detail="No tutor session found for this thread")

        interview_id = tutor_state["interview_id"]
        interview = await conn.fetchrow(
            "SELECT status, tutor_analysis, transcript FROM interviews WHERE id = $1",
            UUID(interview_id),
        )
        if not interview:
            raise HTTPException(status_code=404, detail="Interview record not found")

        tutor_analysis = interview["tutor_analysis"]
        interview_status = interview["status"]

        # Run analysis inline if session ended but analysis hasn't run yet.
        # Atomically claim the analysis slot to prevent duplicate runs from concurrent polls.
        if interview_status in ("analyzing", "completed") and not tutor_analysis and interview["transcript"]:
            claimed = await conn.fetchval(
                "UPDATE interviews SET status = 'analyzing_inline' WHERE id = $1 AND status IN ('analyzing', 'completed') AND tutor_analysis IS NULL RETURNING id",
                UUID(interview_id),
            )
            if not claimed:
                # Another request is already running analysis
                return {"status": "analyzing", "tutor_analysis": None}
            try:
                from ..services.conversation_analyzer import ConversationAnalyzer
                settings = get_settings()
                # Use gemini-2.5-flash for Vertex (gemini-3-flash-preview not accessible on all projects)
                analysis_model = "gemini-2.5-flash" if settings.use_vertex else settings.analysis_model
                analyzer = ConversationAnalyzer(
                    api_key=settings.gemini_api_key,
                    vertex_project=settings.vertex_project if settings.use_vertex else None,
                    vertex_location=settings.vertex_location,
                    model=analysis_model,
                )
                language = tutor_state.get("language", "en")
                tutor_analysis = await analyzer.analyze_tutor_language(
                    transcript=interview["transcript"],
                    language=language,
                )
                # Save analysis and mark completed
                await conn.execute(
                    "UPDATE interviews SET tutor_analysis = $1, status = 'completed' WHERE id = $2",
                    json.dumps(tutor_analysis), UUID(interview_id),
                )
                interview_status = "completed"
            except Exception as e:
                logger.error("Inline tutor analysis failed: %s", e)
                # Reset status so next poll can retry
                await conn.execute(
                    "UPDATE interviews SET status = 'analyzing' WHERE id = $1",
                    UUID(interview_id),
                )
                return {"status": "analyzing", "tutor_analysis": None}

        result = {
            "status": interview_status,
            "tutor_analysis": tutor_analysis if isinstance(tutor_analysis, dict) else (json.loads(tutor_analysis) if tutor_analysis else None),
        }

        # When analysis is complete, save summary as assistant message (idempotent)
        if interview_status == "completed" and tutor_analysis and not tutor_state.get("message_saved"):
            analysis = tutor_analysis if isinstance(tutor_analysis, dict) else json.loads(tutor_analysis)
            proficiency = analysis.get("overall_proficiency", {})
            level = proficiency.get("level", "N/A")
            level_desc = proficiency.get("level_description", "")
            summary_text = f"**Language Practice Complete** — CEFR Level: **{level}** ({level_desc})\n\n"
            strengths = proficiency.get("strengths", [])
            if strengths:
                summary_text += "**Strengths:** " + ", ".join(strengths) + "\n\n"
            areas = proficiency.get("areas_to_improve", [])
            if areas:
                summary_text += "**Areas to Improve:** " + ", ".join(areas) + "\n\n"
            grammar_data = analysis.get("grammar", {})
            errors = grammar_data.get("common_errors", [])
            if errors:
                summary_text += "**Grammar Notes:**\n"
                for err in errors[:5]:
                    if isinstance(err, dict):
                        summary_text += f"- {err.get('error', '')}: {err.get('correction', '')}\n"
                    else:
                        summary_text += f"- {err}\n"

            await doc_svc.add_message(thread_id, "assistant", summary_text.strip())

            # Mark message as saved so we don't duplicate
            current_state_updated = dict(current_state)
            current_state_updated["language_tutor"]["message_saved"] = True
            current_state_updated["language_tutor"]["status"] = "completed"
            await conn.execute(
                "UPDATE mw_threads SET current_state = $1, updated_at = NOW() WHERE id = $2",
                json.dumps(current_state_updated), thread_id,
            )

        return result


# ── Real-time utterance error checking ─────────────────────────────────


UTTERANCE_CHECK_PROMPT_EN = """You are a language tutor analyzing a student's English utterance for errors.

Utterance: "{utterance}"

Return a JSON array of errors found. Each error object has:
- "error": the incorrect word/phrase exactly as spoken
- "correction": the correct form
- "type": one of "grammar", "vocabulary", "pronunciation"
- "brief": a 5-word max explanation

If no errors, return an empty array [].
Only flag clear mistakes, not stylistic preferences. Be concise.
Return ONLY valid JSON, no markdown."""

UTTERANCE_CHECK_PROMPT_ES = """Eres un tutor de idiomas analizando una frase en español de un estudiante.

Frase: "{utterance}"

Devuelve un array JSON de errores encontrados. Cada objeto tiene:
- "error": la palabra/frase incorrecta exactamente como fue dicha
- "correction": la forma correcta
- "type": uno de "grammar", "vocabulary", "pronunciation"
- "brief": explicación de máximo 5 palabras

Si no hay errores, devuelve un array vacío [].
Solo marca errores claros, no preferencias de estilo. Sé conciso.
Devuelve SOLO JSON válido, sin markdown."""


UTTERANCE_CHECK_PROMPT_FR = """Tu es un tuteur de langues analysant une phrase en français d'un étudiant.

Phrase: "{utterance}"

Renvoie un tableau JSON des erreurs trouvées. Chaque objet contient:
- "error": le mot/la phrase incorrecte exactement comme prononcé
- "correction": la forme correcte
- "type": l'un de "grammar", "vocabulary", "pronunciation"
- "brief": explication de 5 mots maximum

S'il n'y a pas d'erreurs, renvoie un tableau vide [].
Ne signale que les erreurs claires, pas les préférences stylistiques. Sois concis.
Renvoie UNIQUEMENT du JSON valide, pas de markdown."""


@router.post("/threads/{thread_id}/tutor/check")
async def check_tutor_utterance(
    thread_id: UUID,
    body: dict = Body(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Real-time error check on a single user utterance during a voice session."""
    utterance = (body.get("utterance") or "").strip()
    language = body.get("language", "en")

    if not utterance or len(utterance) < 3:
        return {"errors": []}

    settings = get_settings()
    try:
        from google import genai
        if settings.use_vertex:
            client = genai.Client(
                vertexai=True,
                project=settings.vertex_project,
                location=settings.vertex_location,
            )
        else:
            client = genai.Client(api_key=settings.gemini_api_key)

        if language in ("es", "es-mx"):
            prompt = UTTERANCE_CHECK_PROMPT_ES.format(utterance=utterance)
        elif language == "fr":
            prompt = UTTERANCE_CHECK_PROMPT_FR.format(utterance=utterance)
        else:
            prompt = UTTERANCE_CHECK_PROMPT_EN.format(utterance=utterance)
        response = await client.aio.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        text = response.text.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1] if "\n" in text else text[3:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()
        errors = json.loads(text)
        return {"errors": errors if isinstance(errors, list) else []}
    except Exception as e:
        logger.warning("Utterance check failed: %s", e)
        return {"errors": []}


# ---------------------------------------------------------------------------
# Thread Collaborator endpoints
# ---------------------------------------------------------------------------

@router.get("/threads/{thread_id}/collaborators")
async def list_thread_collaborators(
    thread_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """List collaborators on a thread with user info.

    Always includes the thread creator as a synthetic 'owner' row even if
    they don't have an explicit mw_thread_collaborators entry. Without this,
    the moment the first invitee is added the owner disappears from the
    collaborator list and their client-side `isOwner` flag flips to false.
    """
    company_id = await get_client_company_id(current_user)
    thread = await doc_svc.get_thread(thread_id, company_id, user_id=current_user.id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    owner_id = thread.get("created_by")

    async with get_connection() as conn:
        rows = await conn.fetch(
            """
            SELECT tc.user_id, tc.role, tc.created_at,
                   u.email, u.avatar_url,
                   COALESCE(cl.name, CONCAT(e.first_name, ' ', e.last_name), a.name, u.email) AS name
            FROM mw_thread_collaborators tc
            JOIN users u ON u.id = tc.user_id
            LEFT JOIN clients cl ON cl.user_id = tc.user_id
            LEFT JOIN employees e ON e.user_id = tc.user_id
            LEFT JOIN admins a ON a.user_id = tc.user_id
            WHERE tc.thread_id = $1
            ORDER BY tc.created_at
            """,
            thread_id,
        )

        result = [
            {
                "user_id": str(r["user_id"]),
                "name": r["name"],
                "email": r["email"],
                "role": r["role"],
                "avatar_url": r["avatar_url"],
                "added_at": r["created_at"].isoformat() if r["created_at"] else None,
            }
            for r in rows
        ]

        # Include the thread creator as a synthetic owner row if not already present
        if owner_id is not None and not any(c["user_id"] == str(owner_id) for c in result):
            owner_row = await conn.fetchrow(
                """
                SELECT u.id, u.email, u.avatar_url,
                       COALESCE(cl.name, CONCAT(e.first_name, ' ', e.last_name), a.name, u.email) AS name
                FROM users u
                LEFT JOIN clients cl ON cl.user_id = u.id
                LEFT JOIN employees e ON e.user_id = u.id
                LEFT JOIN admins a ON a.user_id = u.id
                WHERE u.id = $1
                """,
                owner_id,
            )
            if owner_row:
                result.insert(0, {
                    "user_id": str(owner_row["id"]),
                    "name": owner_row["name"],
                    "email": owner_row["email"],
                    "role": "owner",
                    "avatar_url": owner_row["avatar_url"],
                    "added_at": None,
                })

    return result


@router.post("/threads/{thread_id}/collaborators")
async def add_thread_collaborator(
    thread_id: UUID,
    body: dict = Body(...),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Add a collaborator to a thread. Only the thread owner or admin can invite."""
    company_id = await get_client_company_id(current_user)
    thread = await doc_svc.get_thread(thread_id, company_id, user_id=current_user.id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    # Only the owner or an admin can add collaborators
    if current_user.role != "admin" and thread["created_by"] != current_user.id:
        raise HTTPException(status_code=403, detail="Only the thread owner can add collaborators")

    user_id = body.get("user_id")
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    try:
        collab_user_id = UUID(user_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=400, detail="Invalid user_id")

    if collab_user_id == current_user.id:
        raise HTTPException(status_code=400, detail="You cannot add yourself as a collaborator")

    async with get_connection() as conn:
        # Verify user exists and is active
        target_user = await conn.fetchrow(
            "SELECT id, email FROM users WHERE id = $1 AND is_active = true",
            collab_user_id,
        )
        if not target_user:
            raise HTTPException(status_code=404, detail="User not found")

        # Check if already a collaborator
        existing = await conn.fetchval(
            "SELECT id FROM mw_thread_collaborators WHERE thread_id = $1 AND user_id = $2",
            thread_id, collab_user_id,
        )
        if existing:
            raise HTTPException(status_code=400, detail="User is already a collaborator")

        await conn.execute(
            """INSERT INTO mw_thread_collaborators (thread_id, user_id, invited_by, role)
               VALUES ($1, $2, $3, 'collaborator')""",
            thread_id, collab_user_id, current_user.id,
        )

        # Send inbox notification
        try:
            inviter = await conn.fetchrow("SELECT email FROM users WHERE id = $1", current_user.id)
            inviter_client = await conn.fetchrow("SELECT name FROM clients WHERE user_id = $1", current_user.id)
            inviter_name = (inviter_client["name"] if inviter_client else None) or inviter["email"].split("@")[0]
            thread_title = thread.get("title") or "a thread"

            msg_content = f"**{inviter_name}** has invited you to collaborate on the thread **{thread_title}**."
            conversation = await conn.fetchrow(
                """INSERT INTO inbox_conversations (title, is_group, created_by, last_message_at, last_message_preview)
                   VALUES ($1, false, $2, NOW(), $3)
                   RETURNING id""",
                f"Thread Invite: {thread_title}", current_user.id, msg_content[:100],
            )
            conv_id = conversation["id"]
            await conn.execute(
                "INSERT INTO inbox_participants (conversation_id, user_id) VALUES ($1, $2)", conv_id, current_user.id,
            )
            await conn.execute(
                "INSERT INTO inbox_participants (conversation_id, user_id) VALUES ($1, $2)", conv_id, collab_user_id,
            )
            await conn.execute(
                """INSERT INTO inbox_messages (conversation_id, sender_id, content)
                   VALUES ($1, $2, $3)""",
                conv_id, current_user.id, msg_content,
            )
        except Exception as e:
            logger.warning("Failed to send thread collaborator inbox notification: %s", e)

        # Create MW notification
        try:
            from ..services import notification_service as notif_svc
            inviter_client = await conn.fetchrow("SELECT name FROM clients WHERE user_id = $1", current_user.id)
            inviter_name = (inviter_client["name"] if inviter_client else None) or "Someone"
            await notif_svc.create_notification(
                user_id=collab_user_id,
                company_id=company_id,
                type="thread_collaborator_added",
                title=f"{inviter_name} added you to a thread",
                body=f"You've been added as a collaborator on \"{thread.get('title', 'a thread')}\"",
                link="/work",
                metadata={"thread_id": str(thread_id), "invited_by": str(current_user.id)},
            )
        except Exception as e:
            logger.warning("Failed to create thread collaborator notification: %s", e)

    return {"added": True, "user_id": str(collab_user_id)}


@router.delete("/threads/{thread_id}/collaborators/{user_id}")
async def remove_thread_collaborator(
    thread_id: UUID,
    user_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Remove a collaborator from a thread. Owner, admin, or the collaborator themselves can do this."""
    company_id = await get_client_company_id(current_user)
    thread = await doc_svc.get_thread(thread_id, company_id, user_id=current_user.id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    # Allow removal if: admin, thread owner, or self-removal
    is_owner = thread["created_by"] == current_user.id
    is_self = user_id == current_user.id
    if current_user.role != "admin" and not is_owner and not is_self:
        raise HTTPException(status_code=403, detail="Only the thread owner can remove collaborators")

    async with get_connection() as conn:
        result = await conn.execute(
            "DELETE FROM mw_thread_collaborators WHERE thread_id = $1 AND user_id = $2",
            thread_id, user_id,
        )
        if result.endswith("0"):
            raise HTTPException(status_code=404, detail="Collaborator not found")

    return {"removed": True, "user_id": str(user_id)}


@router.get("/threads/{thread_id}/collaborators/search")
async def search_thread_collaborator_candidates(
    thread_id: UUID,
    q: str = Query(..., min_length=2, max_length=100),
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Search users to invite as thread collaborators."""
    company_id = await get_client_company_id(current_user)
    thread = await doc_svc.get_thread(thread_id, company_id, user_id=current_user.id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    pattern = f"%{q}%"
    async with get_connection() as conn:
        # Search eligible invitees: same-company users + admins + accepted cross-tenant connections
        rows = await conn.fetch(
            """
            SELECT DISTINCT u.id AS user_id, u.email, u.avatar_url,
                   COALESCE(cl.name, CONCAT(e.first_name, ' ', e.last_name), a.name, u.email) AS name
            FROM users u
            LEFT JOIN clients cl ON cl.user_id = u.id
            LEFT JOIN employees e ON e.user_id = u.id
            LEFT JOIN admins a ON a.user_id = u.id
            WHERE u.id != $1
              AND u.is_active = true
              AND (
                  COALESCE(cl.name, CONCAT(e.first_name, ' ', e.last_name), a.name, '') ILIKE $2
                  OR u.email ILIKE $2
              )
              AND (
                  (cl.company_id = $3)
                  OR (e.org_id = $3)
                  OR (a.user_id IS NOT NULL)
                  OR EXISTS (
                      SELECT 1 FROM user_connections uc
                      WHERE uc.status = 'accepted'
                        AND (
                          (uc.user_id = $1 AND uc.connected_user_id = u.id)
                          OR (uc.connected_user_id = $1 AND uc.user_id = u.id)
                        )
                  )
              )
              AND NOT EXISTS(
                  SELECT 1 FROM mw_thread_collaborators tc
                  WHERE tc.thread_id = $4 AND tc.user_id = u.id
              )
            ORDER BY u.email
            LIMIT 10
            """,
            current_user.id, pattern, company_id, thread_id,
        )
    return [
        {
            "user_id": str(r["user_id"]),
            "name": r["name"],
            "email": r["email"],
            "avatar_url": r["avatar_url"],
        }
        for r in rows
    ]
