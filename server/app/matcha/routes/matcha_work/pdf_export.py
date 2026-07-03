"""Project/thread PDF and export rendering: markdown-to-PDF via WeasyPrint,
inline-markdown conversion, and the export endpoints (project, message,
thread-project) that use them.

Extracted from the original flat matcha_work.py during the package split
(2026-07-03). See matcha_work/CLAUDE.md.
"""
import asyncio
import logging
import re
from typing import Optional
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Response
from fastapi.responses import StreamingResponse

from app.core.models.auth import CurrentUser
from app.core.services.storage import get_storage
from app.database import get_connection
from app.matcha.dependencies import require_admin_or_client, get_client_company_id
from app.matcha.routes.matcha_work._shared import _verify_project_access
from app.matcha.services import matcha_work_document as doc_svc

logger = logging.getLogger(__name__)
router = APIRouter()

_IMG_LINE_RE = re.compile(r'^!\[([^\]]*)\]\(([^)]+)\)$')

def _render_inline_md(text: str) -> str:
    """Inline-only markdown matching SwiftUI inlineOnlyPreservingWhitespace.

    Converts **bold**, *italic*, `code` but leaves block structure (lists,
    headings, checkboxes) as literal text so the PDF matches the desktop preview.

    Standalone image lines `![alt](url)` are emitted as <img> tags so the
    downstream storage.inline_storage_images() pass can convert storage URLs to
    base64 data URIs. Mirrors the desktop preview's parseImage() in
    MarkdownPreviewView.swift.
    """
    import re as _re_i, html as _html_i
    out = []
    for line in text.split('\n'):
        # Standalone image line: ![alt](url) — emit <img>, bypass escape so
        # storage.inline_storage_images() can pick up the src.
        img_match = _IMG_LINE_RE.match(line.strip())
        if img_match:
            alt = _html_i.escape(img_match.group(1))
            src = _html_i.escape(img_match.group(2), quote=True)
            out.append(f'<img src="{src}" alt="{alt}" />')
            continue
        s = _html_i.escape(line)
        s = _re_i.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', s)
        s = _re_i.sub(r'__(.+?)__', r'<strong>\1</strong>', s)
        s = _re_i.sub(r'\*([^*\n]+?)\*', r'<em>\1</em>', s)
        s = _re_i.sub(r'(?<!\w)_([^_\n]+?)_(?!\w)', r'<em>\1</em>', s)
        s = _re_i.sub(r'`([^`\n]+?)`', r'<code>\1</code>', s)
        out.append(s)
    return '<br>'.join(out)

async def _render_project_pdf(project: dict) -> bytes:
    """Render a project's title + sections to PDF bytes.

    Mirrors the inline render in `/projects/{id}/export/pdf`. Used by
    the discipline signature-request flow to hand a freshly-rendered
    PDF to the SignatureProvider without re-issuing the export
    endpoint.
    """
    import html as _html

    title = project.get("title") or "Document"
    sections = project.get("sections") or []

    storage = get_storage()

    sections_html = []
    for idx, s in enumerate(sections):
        heading = ""
        if s.get("title"):
            heading = f'<h2><span class="section-num">{idx + 1}.</span> {_html.escape(s["title"])}</h2>'
        content = s.get("content", "") or ""
        # Render markdown FIRST so standalone ![alt](url) lines become <img> tags,
        # THEN inline images so those <img src="..."> URLs get base64'd. The
        # opposite order leaves the just-emitted <img src=https://s3...> tags
        # unprocessed and WeasyPrint can't fetch private storage URLs.
        # External (non-storage) image URLs stay un-inlined and are blocked by
        # the SSRF-safe fetcher at render time (intended).
        if content.lstrip().startswith("<"):
            content_html = await storage.inline_storage_images(content)
        else:
            content_html = f"<div>{_render_inline_md(content)}</div>"
            content_html = await storage.inline_storage_images(content_html)
        sections_html.append(f"{heading}\n<div class='section-body'>{content_html}</div>")

    body_html = "\n".join(sections_html)
    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>
  @page {{ size: A4; margin: 50px 60px; }}
  * {{ box-sizing: border-box; }}
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; font-size: 11pt; line-height: 1.6; color: #1a1a1a; margin: 0; }}
  h1 {{ font-size: 22pt; font-weight: 700; color: #0f172a; margin: 0 0 6px 0; }}
  .title-rule {{ border: none; border-top: 3px solid #22c55e; margin: 0 0 30px 0; }}
  h2 {{ font-size: 14pt; font-weight: 600; color: #0f172a; margin: 28px 0 10px 0; padding-bottom: 6px; border-bottom: 1px solid #e2e8f0; }}
  .section-num {{ color: #22c55e; font-weight: 700; }}
  img {{ max-width: 100%; max-height: 22cm; height: auto; margin: 12px 0; border-radius: 4px; }}
  .section-body {{ margin-bottom: 16px; }}
  .section-body div {{ color: #334155; line-height: 1.7; }}
  .section-body strong {{ color: #0f172a; }}
  pre {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 4px; padding: 10px; font-size: 9pt; white-space: pre-wrap; overflow-wrap: break-word; }}
  code {{ background: #f1f5f9; padding: 1px 5px; border-radius: 3px; font-size: 9pt; color: #b45309; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 9.5pt; }}
  th, td {{ border: 1px solid #e2e8f0; padding: 6px 10px; text-align: left; }}
  th {{ background: #f8fafc; font-weight: 600; color: #0f172a; }}
  a {{ color: #2563eb; text-decoration: none; }}
  .footer {{ margin-top: 40px; padding-top: 12px; border-top: 1px solid #e2e8f0; font-size: 8pt; color: #94a3b8; text-align: center; }}
</style>
</head><body>
<h1>{_html.escape(title)}</h1>
<hr class="title-rule">
{body_html}
<div class="footer">Generated with Matcha Work</div>
</body></html>"""

    try:
        from weasyprint import HTML
        from app.core.services.pdf import safe_url_fetcher
    except ImportError as ie:
        # Surface the real installation hint instead of an opaque 500 so
        # the desktop alert tells the operator what to do.
        logger.error("weasyprint import failed: %s", ie)
        raise HTTPException(
            status_code=501,
            detail="PDF generation not available — install weasyprint on the server (`pip install weasyprint`).",
        )

    try:
        return await asyncio.wait_for(
            asyncio.to_thread(lambda: HTML(string=full_html, url_fetcher=safe_url_fetcher).write_pdf()),
            timeout=60.0,
        )
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail="PDF generation timed out. Try a smaller document or fewer images.")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("PDF render failed for project %s", project.get("id"))
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {type(e).__name__}")

@router.get("/projects/{project_id}/export/{fmt}")
async def export_project_endpoint(
    project_id: UUID,
    fmt: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Export project as PDF, DOCX, or Markdown."""
    from app.matcha.services import project_service as proj_svc
    project, _role = await _verify_project_access(project_id, current_user)
    company_id = await get_client_company_id(current_user)

    title = project["title"]
    sections = project["sections"]

    if fmt not in ("pdf", "md", "docx", "md_frontmatter"):
        raise HTTPException(status_code=400, detail="Supported formats: pdf, md, md_frontmatter, docx")

    if fmt in ("md", "md_frontmatter"):
        md_lines: list[str] = []
        if fmt == "md_frontmatter":
            pdata = project.get("project_data") or {}
            slug = pdata.get("slug") or ""
            excerpt = (pdata.get("excerpt") or "").replace('"', '\\"')
            tags = pdata.get("tags") or []
            status = pdata.get("status") or "draft"
            published_at = pdata.get("published_at") or ""
            safe_title = title.replace('"', '\\"')
            md_lines.append("---")
            md_lines.append(f'title: "{safe_title}"')
            if slug:
                md_lines.append(f"slug: {slug}")
            if excerpt:
                md_lines.append(f'excerpt: "{excerpt}"')
            if tags:
                md_lines.append("tags:")
                for t in tags:
                    md_lines.append(f"  - {t}")
            md_lines.append(f"status: {status}")
            if published_at:
                md_lines.append(f"published_at: {published_at}")
            md_lines.append("---\n")
        md_lines.append(f"# {title}\n")
        for s in sections:
            if s.get("title"):
                md_lines.append(f"## {s['title']}\n")
            md_lines.append(s.get("content", "") + "\n")
        suffix = ".md"
        return Response(
            content="\n".join(md_lines),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{title}{suffix}"'},
        )

    if fmt == "pdf":
        import html as _html

        # Inline storage-owned images as base64 data URIs so WeasyPrint can
        # render them; external (non-storage) URLs stay un-inlined and are
        # blocked by the SSRF-safe fetcher at render time (intended).
        storage = get_storage()

        sections_html = []
        for idx, s in enumerate(sections):
            heading = ""
            if s.get("title"):
                heading = f'<h2><span class="section-num">{idx + 1}.</span> {_html.escape(s["title"])}</h2>'
            content = s.get("content", "")
            # Render markdown FIRST so standalone ![alt](url) lines become <img>
            # tags, THEN inline so those URLs get base64'd. The opposite order
            # left freshly-emitted <img src=https://s3...> tags unprocessed and
            # WeasyPrint couldn't fetch them — image rendered as the URL string.
            if content.lstrip().startswith("<"):
                content_html = await storage.inline_storage_images(content)
            else:
                content_html = f"<div>{_render_inline_md(content)}</div>"
                content_html = await storage.inline_storage_images(content_html)
            sections_html.append(f"{heading}\n<div class='section-body'>{content_html}</div>")

        body_html = "\n".join(sections_html)
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>
  @page {{ size: A4; margin: 50px 60px; }}
  * {{ box-sizing: border-box; }}
  body {{ font-family: 'Helvetica Neue', Arial, sans-serif; font-size: 11pt; line-height: 1.6; color: #1a1a1a; margin: 0; }}
  h1 {{ font-size: 22pt; font-weight: 700; color: #0f172a; margin: 0 0 6px 0; }}
  .title-rule {{ border: none; border-top: 3px solid #22c55e; margin: 0 0 30px 0; }}
  h2 {{ font-size: 14pt; font-weight: 600; color: #0f172a; margin: 28px 0 10px 0; padding-bottom: 6px; border-bottom: 1px solid #e2e8f0; }}
  .section-num {{ color: #22c55e; font-weight: 700; }}
  img {{ max-width: 100%; max-height: 22cm; height: auto; margin: 12px 0; border-radius: 4px; }}
  .section-body {{ margin-bottom: 16px; }}
  .section-body div {{ color: #334155; line-height: 1.7; }}
  .section-body strong {{ color: #0f172a; }}
  pre {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 4px; padding: 10px; font-size: 9pt; white-space: pre-wrap; overflow-wrap: break-word; }}
  code {{ background: #f1f5f9; padding: 1px 5px; border-radius: 3px; font-size: 9pt; color: #b45309; }}
  table {{ border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 9.5pt; }}
  th, td {{ border: 1px solid #e2e8f0; padding: 6px 10px; text-align: left; }}
  th {{ background: #f8fafc; font-weight: 600; color: #0f172a; }}
  a {{ color: #2563eb; text-decoration: none; }}
  .footer {{ margin-top: 40px; padding-top: 12px; border-top: 1px solid #e2e8f0; font-size: 8pt; color: #94a3b8; text-align: center; }}
</style>
</head><body>
<h1>{_html.escape(title)}</h1>
<hr class="title-rule">
{body_html}
<div class="footer">Generated with Matcha Work</div>
</body></html>"""

        try:
            from weasyprint import HTML
            from app.core.services.pdf import safe_url_fetcher
        except ImportError:
            raise HTTPException(status_code=500, detail="PDF generation not available")

        try:
            pdf_bytes = await asyncio.wait_for(
                asyncio.to_thread(lambda: HTML(string=full_html, url_fetcher=safe_url_fetcher).write_pdf()),
                timeout=60.0,
            )
        except asyncio.TimeoutError:
            logger.warning("PDF export timed out for project %s (size=%d bytes)", project_id, len(full_html))
            raise HTTPException(status_code=504, detail="PDF generation timed out. Try a smaller document or fewer images.")
        except HTTPException:
            raise
        except Exception as e:
            logger.exception("PDF export failed for project %s", project_id)
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {type(e).__name__}")

        # Return raw bytes so the desktop client can write them directly to
        # the save-panel URL. A previous implementation uploaded to S3 and
        # returned {"pdf_url": ...} JSON, which the client wrote as the .pdf
        # file — producing an unopenable JSON-in-PDF-extension on disk.
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{title}.pdf"'},
        )

    if fmt == "docx":
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX generation not available")

        def _build_docx():
            doc = DocxDocument()
            doc.add_heading(title, level=0)
            for s in sections:
                if s.get("title"):
                    doc.add_heading(s["title"], level=1)
                for para in (s.get("content") or "").split("\n"):
                    if para.strip():
                        doc.add_paragraph(para)
            import io
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        docx_bytes = await asyncio.to_thread(_build_docx)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{title}.docx"'},
        )

def _pdf_document_html(title: str, body_html: str) -> str:
    """Shared Matcha Work PDF document shell. The single source of PDF styling
    for BOTH the project export and per-message export, so every PDF renders
    identically. `body_html` is already-rendered HTML (project sections, or a
    single message's markdown wrapped in .section-body)."""
    import html as _html
    title = _html.escape(title)
    return f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>
  @page {{ size: A4; margin: 50px 60px; }}
  * {{ box-sizing: border-box; }}
  body {{
    font-family: 'Helvetica Neue', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #1a1a1a;
    margin: 0;
  }}
  h1 {{
    font-size: 22pt;
    font-weight: 700;
    color: #0f172a;
    margin: 0 0 6px 0;
    letter-spacing: -0.5px;
  }}
  .subtitle {{
    font-size: 9pt;
    color: #94a3b8;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    margin-bottom: 20px;
  }}
  .title-rule {{
    border: none;
    border-top: 3px solid #22c55e;
    margin: 0 0 30px 0;
  }}
  h2 {{
    font-size: 14pt;
    font-weight: 600;
    color: #0f172a;
    margin: 28px 0 10px 0;
    padding-bottom: 6px;
    border-bottom: 1px solid #e2e8f0;
  }}
  .section-num {{
    color: #22c55e;
    font-weight: 700;
  }}
  img {{
    max-width: 100%;
    max-height: 22cm;
    height: auto;
    margin: 12px 0;
    border-radius: 4px;
  }}
  .section-body {{
    margin-bottom: 16px;
  }}
  .section-body p {{
    margin: 6px 0;
    color: #334155;
  }}
  .section-body ul, .section-body ol {{
    margin: 6px 0;
    padding-left: 22px;
    color: #334155;
  }}
  .section-body li {{
    margin: 3px 0;
  }}
  .section-body strong {{
    color: #0f172a;
  }}
  .section-body em {{
    color: #475569;
  }}
  pre {{
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 4px;
    padding: 10px 14px;
    font-size: 9pt;
    font-family: 'SF Mono', 'Menlo', monospace;
    overflow-wrap: break-word;
    white-space: pre-wrap;
    color: #334155;
  }}
  code {{
    background: #f1f5f9;
    padding: 1px 5px;
    border-radius: 3px;
    font-size: 9pt;
    font-family: 'SF Mono', 'Menlo', monospace;
    color: #b45309;
  }}
  table {{
    border-collapse: collapse;
    width: 100%;
    margin: 12px 0;
    font-size: 9.5pt;
  }}
  th, td {{
    border: 1px solid #e2e8f0;
    padding: 6px 10px;
    text-align: left;
  }}
  th {{
    background: #f8fafc;
    font-weight: 600;
    color: #0f172a;
  }}
  blockquote {{
    border-left: 3px solid #22c55e;
    margin: 12px 0;
    padding: 8px 16px;
    background: #f0fdf4;
    color: #334155;
  }}
  a {{
    color: #2563eb;
    text-decoration: none;
  }}
  .footer {{
    margin-top: 40px;
    padding-top: 12px;
    border-top: 1px solid #e2e8f0;
    font-size: 8pt;
    color: #94a3b8;
    text-align: center;
  }}
</style>
</head><body>
<h1>{title}</h1>
<hr class="title-rule">
{body_html}
<div class="footer">Generated with Matcha Work</div>
</body></html>"""

async def _html_to_pdf_bytes(full_html: str) -> bytes:
    """Render an HTML document string to PDF bytes via WeasyPrint, off-thread.

    Storage-owned `<img src=...>` references are base64-inlined to `data:` URIs
    first, since the SSRF-safe fetcher blocks raw storage URLs. External/
    non-storage image URLs are intentionally left for the fetcher to block.
    """
    try:
        from weasyprint import HTML
        from app.core.services.pdf import safe_url_fetcher
    except ImportError:
        raise HTTPException(status_code=501, detail="PDF generation not available — install weasyprint")
    full_html = await get_storage().inline_storage_images(full_html)
    return await asyncio.to_thread(lambda: HTML(string=full_html, url_fetcher=safe_url_fetcher).write_pdf())

def _pdf_title_from_markdown(content: str, fallback: str = "Deal Memo") -> str:
    """First markdown heading (# / ## / ###) becomes the document title."""
    for line in (content or "").splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            if title:
                return title
    return fallback

def _strip_leading_title_heading(content: str) -> str:
    """Drop the first heading line (the one promoted to the document title via
    _pdf_title_from_markdown) so it isn't rendered a second time in the body."""
    lines = (content or "").splitlines()
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        if line.strip().startswith("#"):
            return "\n".join(lines[i + 1:]).lstrip("\n")
        break
    return content or ""

async def _render_markdown_pdf(title: str, content_md: str) -> bytes:
    """Render a thread reply's markdown to PDF through the SAME document shell
    as the project export, so a single message exports with the project look.
    The leading title heading is stripped so it isn't duplicated under the
    shell's own <h1> title."""
    import html as _html
    content = _strip_leading_title_heading(content_md or "")
    if content.lstrip().startswith("<"):
        inner = content  # already HTML
    else:
        try:
            import markdown as _md
            inner = _md.markdown(content, extensions=["tables", "fenced_code", "nl2br", "sane_lists"])
        except ImportError:
            inner = f"<p>{_html.escape(content)}</p>"
    body_html = f"<div class='section-body'>{inner}</div>"
    return await _html_to_pdf_bytes(_pdf_document_html(title, body_html))

@router.post("/threads/{thread_id}/messages/{message_id}/export/pdf")
async def export_message_pdf(
    thread_id: UUID,
    message_id: UUID,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Render a single thread message's markdown to a downloadable PDF. Lets a
    plain thread produce deal memos / briefs without a Project — the AI writes
    the memo as markdown in its reply and the user exports that reply."""
    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id, user_id=current_user.id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    async with get_connection() as conn:
        row = await conn.fetchrow(
            "SELECT id, content FROM mw_messages WHERE id = $1 AND thread_id = $2",
            message_id, thread_id,
        )
    if row is None:
        raise HTTPException(status_code=404, detail="Message not found")
    content = row["content"] or ""
    if not content.strip():
        raise HTTPException(status_code=400, detail="Message has no content to export")

    title = _pdf_title_from_markdown(content)
    pdf_bytes = await _render_markdown_pdf(title, content)
    safe_title = "".join(c for c in title if c.isalnum() or c in (" ", "-", "_")).strip() or "Deal Memo"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
    )

@router.get("/threads/{thread_id}/project/export/{fmt}")
async def export_project(
    thread_id: UUID,
    fmt: str,
    current_user: CurrentUser = Depends(require_admin_or_client),
):
    """Export project as PDF, DOCX, or Markdown."""
    if fmt not in ("pdf", "md", "docx"):
        raise HTTPException(status_code=400, detail="Supported formats: pdf, md, docx")

    company_id = await get_client_company_id(current_user)
    if company_id is None:
        raise HTTPException(status_code=404, detail="Thread not found")
    thread = await doc_svc.get_thread(thread_id, company_id)
    if thread is None:
        raise HTTPException(status_code=404, detail="Thread not found")

    state = thread.get("current_state") or {}
    title = state.get("project_title") or "Project"
    sections = state.get("project_sections") or []

    if fmt == "md":
        md_lines = [f"# {title}\n"]
        for s in sections:
            if s.get("title"):
                md_lines.append(f"## {s['title']}\n")
            md_lines.append(s.get("content", "") + "\n")
        md_content = "\n".join(md_lines)
        return Response(
            content=md_content,
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{title}.md"'},
        )

    if fmt == "pdf":
        import html as _html

        sections_html = []
        for idx, s in enumerate(sections):
            heading = ""
            if s.get("title"):
                heading = f'<h2><span class="section-num">{idx + 1}.</span> {_html.escape(s["title"])}</h2>'
            content = s.get("content", "")
            # Content may be HTML (from TipTap editor) or legacy markdown
            if content.lstrip().startswith("<"):
                content_html = content  # already HTML
            else:
                try:
                    import markdown as _md
                    content_html = _md.markdown(content, extensions=["tables", "fenced_code", "nl2br"])
                except ImportError:
                    content_html = f"<p>{_html.escape(content)}</p>"
            sections_html.append(f"{heading}\n<div class='section-body'>{content_html}</div>")

        body_html = "\n".join(sections_html)
        pdf_bytes = await _html_to_pdf_bytes(_pdf_document_html(title, body_html))

        prefix = doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "project-exports")
        pdf_url = await get_storage().upload_file(
            pdf_bytes, f"{title}.pdf", prefix=prefix, content_type="application/pdf"
        )
        return {"pdf_url": pdf_url}

    if fmt == "docx":
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
        except ImportError:
            raise HTTPException(status_code=500, detail="DOCX generation not available (python-docx not installed)")

        def _build_docx():
            doc = DocxDocument()
            doc.add_heading(title, level=0)
            for s in sections:
                if s.get("title"):
                    doc.add_heading(s["title"], level=1)
                for para in (s.get("content") or "").split("\n"):
                    if para.strip():
                        doc.add_paragraph(para)
            import io
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        docx_bytes = await asyncio.to_thread(_build_docx)
        prefix = doc_svc.build_matcha_work_thread_storage_prefix(company_id, thread_id, "project-exports")
        docx_url = await get_storage().upload_file(
            docx_bytes, f"{title}.docx", prefix=prefix,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return {"docx_url": docx_url}
