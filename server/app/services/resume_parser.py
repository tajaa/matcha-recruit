import os
import re
import tempfile
from typing import Optional, Any

import fitz  # PyMuPDF
from docx import Document


# Common skills to look for
TECH_SKILLS = {
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "ruby", "php",
    "sql", "nosql", "mongodb", "postgresql", "mysql", "redis", "elasticsearch",
    "react", "angular", "vue", "node", "express", "django", "flask", "fastapi", "spring",
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins", "ci/cd",
    "git", "linux", "agile", "scrum", "jira", "confluence",
    "machine learning", "ml", "ai", "deep learning", "nlp", "computer vision",
    "data science", "data analysis", "pandas", "numpy", "tensorflow", "pytorch",
    "html", "css", "sass", "tailwind", "bootstrap",
    "rest", "graphql", "api", "microservices", "serverless",
    "figma", "sketch", "adobe", "photoshop", "illustrator",
    "excel", "powerpoint", "word", "salesforce", "hubspot",
    "project management", "leadership", "communication", "problem solving",
}

# Education keywords
DEGREE_PATTERNS = [
    r"(?i)\b(ph\.?d|doctorate|doctoral)\b",
    r"(?i)\b(master'?s?|m\.?s\.?|m\.?a\.?|mba|m\.?eng)\b",
    r"(?i)\b(bachelor'?s?|b\.?s\.?|b\.?a\.?|b\.?eng)\b",
    r"(?i)\b(associate'?s?|a\.?s\.?|a\.?a\.?)\b",
]


class ResumeParser:
    """Fast algorithmic resume parser - extracts basic info without AI."""

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        text_parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text())
        return "\n".join(text_parts)

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return "\n".join(text_parts)

    def extract_text(self, file_path: str) -> str:
        """Extract text from a resume file (PDF or DOCX)."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def extract_text_from_bytes(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from file bytes."""
        ext = os.path.splitext(filename)[1].lower()

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            return self.extract_text(tmp_path)
        finally:
            os.unlink(tmp_path)

    def extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text."""
        pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
        match = re.search(pattern, text)
        return match.group(0) if match else None

    def extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text."""
        patterns = [
            r"\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
            r"\d{3}[-.\s]\d{3}[-.\s]\d{4}",
            r"\(\d{3}\)\s*\d{3}[-.\s]\d{4}",
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return None

    def extract_name(self, text: str) -> Optional[str]:
        """Extract name - typically the first non-empty line."""
        lines = text.strip().split("\n")
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            # Skip if it looks like contact info or header
            if not line or "@" in line or re.search(r"\d{3}.*\d{4}", line):
                continue
            # Skip common header words
            if line.lower() in ["resume", "curriculum vitae", "cv"]:
                continue
            # Likely a name if it's 2-4 words, all capitalized or title case
            words = line.split()
            if 1 <= len(words) <= 4:
                # Check if words look like names (capitalized, no numbers)
                if all(w[0].isupper() and not any(c.isdigit() for c in w) for w in words if w):
                    return line
        return None

    def extract_skills(self, text: str) -> list[str]:
        """Extract skills from text using keyword matching."""
        text_lower = text.lower()
        found_skills = []

        for skill in TECH_SKILLS:
            # Match whole words only
            pattern = r"\b" + re.escape(skill) + r"\b"
            if re.search(pattern, text_lower):
                found_skills.append(skill.title() if len(skill) > 3 else skill.upper())

        return list(set(found_skills))

    def extract_education(self, text: str) -> list[dict[str, Any]]:
        """Extract education information."""
        education = []

        for pattern in DEGREE_PATTERNS:
            matches = re.finditer(pattern, text)
            for match in matches:
                # Get surrounding context (100 chars before and after)
                start = max(0, match.start() - 100)
                end = min(len(text), match.end() + 100)
                context = text[start:end]

                # Try to find a year
                year_match = re.search(r"20\d{2}|19\d{2}", context)

                education.append({
                    "degree": match.group(0),
                    "field": None,
                    "institution": None,
                    "year": int(year_match.group(0)) if year_match else None,
                })

        return education

    def estimate_experience_years(self, text: str) -> Optional[int]:
        """Estimate years of experience from text."""
        # Look for explicit mentions like "10+ years" or "5 years of experience"
        patterns = [
            r"(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)",
            r"(\d+)\+?\s*years?\s*in\s*(?:the\s*)?(?:industry|field)",
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return int(match.group(1))

        # Count date ranges to estimate
        year_ranges = re.findall(r"(20\d{2}|19\d{2})\s*[-–]\s*(20\d{2}|19\d{2}|present|current)", text, re.IGNORECASE)
        if year_ranges:
            total_years = 0
            current_year = 2024
            for start, end in year_ranges:
                start_year = int(start)
                end_year = current_year if end.lower() in ["present", "current"] else int(end)
                total_years += max(0, end_year - start_year)
            if total_years > 0:
                return total_years

        return None

    def parse_resume(self, resume_text: str) -> dict[str, Any]:
        """Parse resume text into structured data algorithmically."""
        return {
            "name": self.extract_name(resume_text),
            "email": self.extract_email(resume_text),
            "phone": self.extract_phone(resume_text),
            "skills": self.extract_skills(resume_text),
            "experience_years": self.estimate_experience_years(resume_text),
            "education": self.extract_education(resume_text),
            "work_history": [],  # Would need more complex parsing
            "summary": None,  # Leave for AI analysis later
            "_resume_text": resume_text,
        }

    def parse_resume_file(self, file_bytes: bytes, filename: str) -> dict[str, Any]:
        """Extract text and parse a resume file."""
        resume_text = self.extract_text_from_bytes(file_bytes, filename)
        return self.parse_resume(resume_text)
