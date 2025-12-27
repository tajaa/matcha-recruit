"""Document Parser Service for ER Copilot.

Parses investigation documents: PDF, DOCX, TXT, CSV, JSON.
Extracts text, detects speaker turns, and identifies temporal references.
"""

import csv
import io
import json
import os
import re
import tempfile
from dataclasses import dataclass
from typing import Optional

import fitz  # PyMuPDF
from docx import Document


@dataclass
class SpeakerTurn:
    """A single speaker turn in a transcript."""
    speaker: str
    text: str
    line_start: int
    line_end: int


@dataclass
class TemporalReference:
    """A temporal reference found in text."""
    text: str
    type: str  # date, time, duration, relative
    position: int
    context: str


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    text: str
    speaker_turns: list[SpeakerTurn]
    temporal_refs: list[TemporalReference]
    page_count: Optional[int] = None
    metadata: Optional[dict] = None


class ERDocumentParser:
    """Parse investigation documents for ER Copilot."""

    # Patterns for detecting speaker turns in transcripts
    SPEAKER_PATTERNS = [
        # "Speaker Name:" or "SPEAKER NAME:"
        r"^([A-Z][A-Za-z\s\.]+):\s*(.+)$",
        # "Q:" and "A:" format
        r"^([QA]):\s*(.+)$",
        # "[Speaker Name]" or "(Speaker Name)"
        r"^\[([^\]]+)\]\s*(.+)$",
        r"^\(([^)]+)\)\s*(.+)$",
        # "Interviewer/Witness/Manager/HR:" etc
        r"^(Interviewer|Witness|Manager|Employee|HR|Complainant|Respondent|Attorney):\s*(.+)$",
    ]

    # Patterns for temporal references
    TEMPORAL_PATTERNS = {
        "date_mdy": r"\b(0?[1-9]|1[0-2])[-/](0?[1-9]|[12]\d|3[01])[-/](19|20)?\d{2}\b",
        "date_ymd": r"\b(19|20)\d{2}[-/](0?[1-9]|1[0-2])[-/](0?[1-9]|[12]\d|3[01])\b",
        "date_written": r"\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s*(19|20)?\d{2}\b",
        "date_short": r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+\d{1,2},?\s*(19|20)?\d{2}\b",
        "time_12h": r"\b(0?[1-9]|1[0-2]):([0-5]\d)\s*(am|pm|AM|PM)\b",
        "time_24h": r"\b([01]?\d|2[0-3]):([0-5]\d)\b",
        "duration": r"\b(\d+)\s*(minute|hour|day|week|month|year)s?\b",
        "relative": r"\b(yesterday|today|tomorrow|last week|next week|last month|this morning|that afternoon|the following day)\b",
        "ago": r"\b(\d+)\s*(minute|hour|day|week|month|year)s?\s+ago\b",
    }

    def extract_text_from_pdf(self, file_path: str) -> tuple[str, int]:
        """Extract text from a PDF file."""
        text_parts = []
        with fitz.open(file_path) as doc:
            page_count = len(doc)
            for page in doc:
                text_parts.append(page.get_text())
        return "\n".join(text_parts), page_count

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return "\n".join(text_parts)

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from a plain text file."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def parse_csv(self, file_path: str) -> tuple[str, list[dict]]:
        """Parse CSV file, return as text and structured data."""
        rows = []
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.DictReader(f)
            for row in reader:
                rows.append(dict(row))

        # Convert to readable text
        text_parts = []
        for i, row in enumerate(rows, 1):
            text_parts.append(f"Row {i}:")
            for key, value in row.items():
                text_parts.append(f"  {key}: {value}")
            text_parts.append("")

        return "\n".join(text_parts), rows

    def parse_json(self, file_path: str) -> tuple[str, dict]:
        """Parse JSON file, return as text and structured data."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Convert to readable text
        text = json.dumps(data, indent=2)
        return text, data

    def extract_text(self, file_path: str) -> tuple[str, Optional[int]]:
        """Extract text from any supported file type."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            text, page_count = self.extract_text_from_pdf(file_path)
            return text, page_count
        elif ext in [".docx", ".doc"]:
            return self.extract_text_from_docx(file_path), None
        elif ext == ".txt":
            return self.extract_text_from_txt(file_path), None
        elif ext == ".csv":
            text, _ = self.parse_csv(file_path)
            return text, None
        elif ext == ".json":
            text, _ = self.parse_json(file_path)
            return text, None
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def extract_text_from_bytes(self, file_bytes: bytes, filename: str) -> tuple[str, Optional[int]]:
        """Extract text from file bytes."""
        ext = os.path.splitext(filename)[1].lower()

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            return self.extract_text(tmp_path)
        finally:
            os.unlink(tmp_path)

    def detect_speaker_turns(self, text: str) -> list[SpeakerTurn]:
        """Detect speaker turns in interview transcript text."""
        turns = []
        lines = text.split("\n")
        current_speaker = None
        current_text_parts = []
        current_line_start = 0

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Try to match speaker patterns
            matched = False
            for pattern in self.SPEAKER_PATTERNS:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    # Save previous speaker's turn
                    if current_speaker and current_text_parts:
                        turns.append(SpeakerTurn(
                            speaker=current_speaker,
                            text=" ".join(current_text_parts),
                            line_start=current_line_start,
                            line_end=i - 1,
                        ))

                    # Start new speaker turn
                    current_speaker = match.group(1).strip()
                    current_text_parts = [match.group(2).strip()]
                    current_line_start = i
                    matched = True
                    break

            # If no speaker match, append to current speaker's text
            if not matched and current_speaker:
                current_text_parts.append(line)

        # Don't forget the last speaker's turn
        if current_speaker and current_text_parts:
            turns.append(SpeakerTurn(
                speaker=current_speaker,
                text=" ".join(current_text_parts),
                line_start=current_line_start,
                line_end=len(lines) - 1,
            ))

        return turns

    def extract_temporal_references(self, text: str) -> list[TemporalReference]:
        """Extract all temporal references from text."""
        refs = []

        for ref_type, pattern in self.TEMPORAL_PATTERNS.items():
            for match in re.finditer(pattern, text, re.IGNORECASE):
                # Get surrounding context (50 chars before and after)
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].strip()

                refs.append(TemporalReference(
                    text=match.group(0),
                    type=ref_type,
                    position=match.start(),
                    context=context,
                ))

        # Sort by position
        refs.sort(key=lambda r: r.position)
        return refs

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 500,
        overlap: int = 50,
    ) -> list[dict]:
        """
        Split text into overlapping chunks for embedding.

        Args:
            text: The text to chunk.
            chunk_size: Target size of each chunk in characters.
            overlap: Number of overlapping characters between chunks.

        Returns:
            List of chunks with metadata.
        """
        chunks = []

        # Split by paragraphs first
        paragraphs = text.split("\n\n")
        current_chunk = ""
        current_line_start = 0
        line_count = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_lines = para.count("\n") + 1

            # If adding this paragraph would exceed chunk size, save current chunk
            if len(current_chunk) + len(para) > chunk_size and current_chunk:
                chunks.append({
                    "content": current_chunk.strip(),
                    "line_start": current_line_start,
                    "line_end": line_count,
                    "char_start": len("".join(c["content"] for c in chunks)),
                })

                # Start new chunk with overlap from previous
                if overlap > 0 and len(current_chunk) > overlap:
                    current_chunk = current_chunk[-overlap:] + "\n\n" + para
                else:
                    current_chunk = para
                current_line_start = line_count
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para

            line_count += para_lines

        # Don't forget the last chunk
        if current_chunk:
            chunks.append({
                "content": current_chunk.strip(),
                "line_start": current_line_start,
                "line_end": line_count,
                "char_start": len("".join(c["content"] for c in chunks)),
            })

        # Add chunk indices
        for i, chunk in enumerate(chunks):
            chunk["chunk_index"] = i

        return chunks

    def parse_document(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """
        Parse a document and extract all relevant information.

        Args:
            file_bytes: The document content as bytes.
            filename: Original filename (for extension detection).

        Returns:
            ParsedDocument with text, speaker turns, and temporal references.
        """
        text, page_count = self.extract_text_from_bytes(file_bytes, filename)

        return ParsedDocument(
            text=text,
            speaker_turns=self.detect_speaker_turns(text),
            temporal_refs=self.extract_temporal_references(text),
            page_count=page_count,
            metadata={
                "filename": filename,
                "extension": os.path.splitext(filename)[1].lower(),
                "text_length": len(text),
            },
        )
