"""Generate job-description DOCX files, upload to S3, print URL mapping.

Run from `server/`:
    venv/bin/python scripts/generate_job_descriptions.py

Prints a Python dict you can paste into jobDescriptionsData.ts as downloadUrl values.
"""

import asyncio
import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from app.config import load_settings
from app.core.services.storage import get_storage

from scripts.data.job_description_content import JD_CONTENT, EEO_STATEMENT

DISCLAIMER = (
    "DISCLAIMER: This template is provided by Matcha for informational purposes only "
    "and does not constitute legal advice. Employment laws vary by jurisdiction and "
    "change frequently. Review with qualified employment counsel before use."
)

JD_TITLES: dict[str, str] = {
    "front-desk-agent": "Front Desk Agent",
    "housekeeper": "Housekeeper",
    "housekeeping-supervisor": "Housekeeping Supervisor",
    "concierge": "Concierge",
    "event-coordinator": "Event Coordinator",
    "registered-nurse": "Registered Nurse",
    "lvn-lpn": "LVN / LPN",
    "medical-assistant": "Medical Assistant",
    "cna": "Certified Nursing Assistant (CNA)",
    "phlebotomist": "Phlebotomist",
    "medical-receptionist": "Medical Receptionist",
    "behavioral-health-technician": "Behavioral Health Technician",
    "home-health-aide": "Home Health Aide",
    "retail-sales-associate": "Retail Sales Associate",
    "cashier": "Cashier",
    "store-manager": "Store Manager",
    "assistant-store-manager": "Assistant Store Manager",
    "visual-merchandiser": "Visual Merchandiser",
    "stock-associate": "Stock Associate",
    "line-cook": "Line Cook",
    "prep-cook": "Prep Cook",
    "server": "Server",
    "bartender": "Bartender",
    "host": "Host / Hostess",
    "dishwasher": "Dishwasher",
    "shift-leader": "Shift Leader",
    "general-manager-restaurant": "Restaurant General Manager",
    "delivery-driver": "Delivery Driver",
    "electrician": "Electrician",
    "plumber": "Plumber",
    "hvac-technician": "HVAC Technician",
    "carpenter": "Carpenter",
    "project-superintendent": "Project Superintendent",
    "safety-officer": "Safety Officer",
    "production-operator": "Production Operator",
    "forklift-operator": "Forklift Operator",
    "warehouse-associate": "Warehouse Associate",
    "shipping-receiving-clerk": "Shipping & Receiving Clerk",
    "maintenance-technician": "Maintenance Technician",
    "quality-inspector": "Quality Inspector",
    "hr-generalist": "HR Generalist",
    "hr-business-partner": "HR Business Partner",
    "recruiter": "Recruiter",
    "office-manager": "Office Manager",
    "executive-assistant": "Executive Assistant",
    "accountant": "Accountant",
    "bookkeeper": "Bookkeeper",
    "payroll-specialist": "Payroll Specialist",
    "paralegal": "Paralegal",
    "software-engineer": "Software Engineer",
    "senior-software-engineer": "Senior Software Engineer",
    "product-manager": "Product Manager",
    "designer": "Product Designer",
    "devops-engineer": "DevOps / SRE Engineer",
    "data-analyst": "Data Analyst",
    "it-support-specialist": "IT Support Specialist",
    "account-executive": "Account Executive",
    "sdr": "Sales Development Representative",
    "customer-success-manager": "Customer Success Manager",
    "marketing-manager": "Marketing Manager",
    "content-marketer": "Content Marketer",
    "social-media-manager": "Social Media Manager",
}


def _build_jd_doc(slug: str, title: str) -> bytes:
    content = JD_CONTENT[slug]
    doc = Document()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = t.add_run(f"Job Description: {title}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    # Brand subtitle
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Provided by Matcha — hey-matcha.com")
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Disclaimer
    disc = doc.add_paragraph()
    dr = disc.add_run(DISCLAIMER)
    dr.italic = True
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Job Summary
    _add_section_heading(doc, "Job Summary")
    doc.add_paragraph(content["summary"])
    doc.add_paragraph()

    # Key Responsibilities
    _add_section_heading(doc, "Key Responsibilities")
    for item in content["responsibilities"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    # Requirements
    _add_section_heading(doc, "Requirements")
    for item in content["requirements"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    # Preferred Qualifications
    _add_section_heading(doc, "Preferred Qualifications")
    for item in content["preferred"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()

    # Compensation & Benefits placeholder
    _add_section_heading(doc, "Compensation & Benefits")
    doc.add_paragraph(
        "Compensation is competitive and commensurate with experience. "
        "We offer a comprehensive benefits package including health, dental, vision, "
        "retirement savings, paid time off, and other benefits as described in the "
        "Employee Handbook. [Insert specific compensation range and benefits details.]"
    )
    doc.add_paragraph()

    # EEO
    _add_section_heading(doc, "Equal Opportunity Employer")
    doc.add_paragraph(EEO_STATEMENT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_section_heading(doc: Document, text: str) -> None:
    h = doc.add_paragraph()
    hr = h.add_run(text)
    hr.bold = True
    hr.font.size = Pt(13)
    hr.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


async def main():
    load_settings()
    storage = get_storage()
    if not storage.s3_client or not storage.bucket:
        print("ERROR: S3 not configured. Set AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, S3_BUCKET_NAME.")
        sys.exit(1)

    slugs_to_run = list(JD_TITLES.keys())
    print(f"Generating {len(slugs_to_run)} job descriptions and uploading to S3 bucket: {storage.bucket}")
    print()

    results: dict[str, str] = {}
    for slug in slugs_to_run:
        title = JD_TITLES[slug]
        filename = f"{slug}.docx"
        print(f"  - {filename}…", end=" ", flush=True)
        try:
            data = _build_jd_doc(slug, title)
            url = await storage.upload_file(
                file_bytes=data,
                filename=filename,
                prefix="resources/job-descriptions",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            results[slug] = url
            print(f"OK ({len(data):,} bytes)")
        except Exception as e:
            print(f"FAIL: {e}")

    print()
    print("=" * 80)
    print("downloadUrl MAPPING — paste into jobDescriptionsData.ts:")
    print("=" * 80)
    for slug, url in results.items():
        print(f'  // {JD_TITLES[slug]}')
        print(f'  "{slug}": "{url}",')
    print()
    print(f"Done. {len(results)}/{len(slugs_to_run)} uploaded.")


if __name__ == "__main__":
    asyncio.run(main())
