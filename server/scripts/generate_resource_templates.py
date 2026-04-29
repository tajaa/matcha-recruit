"""Generate HR resource templates as DOCX, upload to S3, print URL mapping.

Run from `server/`:
    venv/bin/python scripts/generate_resource_templates.py

Prints a Python dict block you can paste into
`app/core/routes/resources.py` ASSETS to flip them live.
"""

import asyncio
import io
import sys
from pathlib import Path

# Ensure local imports work when run from server/
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from app.config import load_settings
from app.core.services.storage import get_storage


DISCLAIMER = (
    "DISCLAIMER: This template is provided by Matcha for informational purposes only "
    "and does not constitute legal advice. Employment laws vary by jurisdiction and "
    "change frequently. Review with qualified employment counsel before use."
)


def _build_doc(title: str, sections: list[tuple[str, list[str]]]) -> bytes:
    """Build a DOCX with a title, disclaimer, and a list of (heading, paragraphs) sections."""
    doc = Document()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = t.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    # Subtitle / brand line
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Provided by Matcha — hey-matcha.com")
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Disclaimer
    disc = doc.add_paragraph()
    dr = disc.add_run(DISCLAIMER)
    dr.italic = True
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    for heading, paragraphs in sections:
        h = doc.add_paragraph()
        hr = h.add_run(heading)
        hr.bold = True
        hr.font.size = Pt(13)
        for p in paragraphs:
            if p.startswith("- "):
                doc.add_paragraph(p[2:], style="List Bullet")
            elif p.startswith("# "):
                # sub-heading
                sh = doc.add_paragraph()
                shr = sh.add_run(p[2:])
                shr.bold = True
                shr.font.size = Pt(11)
            else:
                doc.add_paragraph(p)
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Template content
# ---------------------------------------------------------------------------

TEMPLATES: dict[str, dict] = {
    "offer-letter": {
        "title": "Offer Letter",
        "sections": [
            ("[Date]", ["[Candidate Name]", "[Candidate Address]"]),
            ("Dear [Candidate Name],",
             ["[Company Name] is pleased to offer you the position of [Job Title], "
              "reporting to [Manager Name]. We're excited about the value you'll bring to our team."]),
            ("Position", [
                "Title: [Job Title]",
                "Department: [Department]",
                "Reports to: [Manager Name and Title]",
                "Start date: [Start Date]",
                "Work location: [Office address / Remote / Hybrid]",
                "Employment type: [Full-time / Part-time]",
                "Classification: [Exempt / Non-exempt under FLSA]",
            ]),
            ("Compensation", [
                "Base salary: $[Amount] per [year/hour], paid [bi-weekly / semi-monthly / monthly] in accordance with the Company's standard payroll practices.",
                "Bonus: [Eligible for X% target annual bonus / Sign-on bonus of $X / Not applicable]",
                "Equity: [Subject to a separate equity grant agreement / Not applicable]",
            ]),
            ("Benefits", [
                "You will be eligible for benefits offered to similarly situated employees, including health, dental, vision, retirement, and paid time off, subject to the terms of each plan.",
            ]),
            ("Contingencies", [
                "This offer is contingent on:",
                "- Satisfactory completion of a background check",
                "- Verification of your eligibility to work in the United States (Form I-9)",
                "- Execution of the Company's standard Confidentiality and Invention Assignment Agreement",
                "- [Reference checks / Drug screening / Other]",
            ]),
            ("At-Will Employment", [
                "Your employment with [Company Name] is at-will. Either you or the Company may terminate the employment relationship at any time, with or without cause or notice. Nothing in this letter, or any other Company document, creates an employment contract for any specific duration. (Note: Montana is not an at-will state — modify accordingly.)",
            ]),
            ("Acceptance", [
                "Please sign and return this letter by [Date] to indicate your acceptance.",
                "",
                "Sincerely,",
                "",
                "_______________________________",
                "[Hiring Manager Name]",
                "[Title]",
                "",
                "Accepted by:",
                "",
                "_______________________________   Date: __________",
                "[Candidate Name]",
            ]),
        ],
    },

    "pip": {
        "title": "Performance Improvement Plan (PIP)",
        "sections": [
            ("Employee Information", [
                "Employee name: ____________________________",
                "Position: ____________________________",
                "Manager: ____________________________",
                "Department: ____________________________",
                "PIP start date: ____________________________",
                "PIP review dates: 30 days [____], 60 days [____], 90 days [____]",
                "PIP end date: ____________________________",
            ]),
            ("Reason for PIP", [
                "Specific performance gaps requiring improvement (with examples and dates):",
                "",
                "1. _______________________________________________________________",
                "2. _______________________________________________________________",
                "3. _______________________________________________________________",
                "",
                "Prior coaching, feedback, or warnings provided (with dates):",
                "_____________________________________________________________________",
            ]),
            ("Performance Expectations", [
                "Specific, measurable standards the employee must meet during the PIP period:",
                "",
                "Goal 1: ___________________________________________________________",
                "  Measure: ________________________________________________________",
                "  Target: _________________________________________________________",
                "",
                "Goal 2: ___________________________________________________________",
                "  Measure: ________________________________________________________",
                "  Target: _________________________________________________________",
                "",
                "Goal 3: ___________________________________________________________",
                "  Measure: ________________________________________________________",
                "  Target: _________________________________________________________",
            ]),
            ("Support Provided", [
                "Resources, training, mentoring, or accommodations the company will provide:",
                "- ____________________________________________________________________",
                "- ____________________________________________________________________",
                "- ____________________________________________________________________",
            ]),
            ("Review Cadence", [
                "Manager and employee will meet every [weekly / bi-weekly] on [day] at [time] "
                "to discuss progress against the goals above. Written progress notes will be "
                "documented at each 30-day milestone.",
            ]),
            ("Consequences", [
                "Failure to meet the expectations outlined in this PIP by the end of the "
                "PIP period may result in further disciplinary action, up to and including "
                "termination of employment. The company reserves the right to take immediate "
                "action if performance worsens or new issues arise during the PIP period.",
            ]),
            ("Acknowledgment", [
                "I have read and understand this Performance Improvement Plan. My signature "
                "acknowledges receipt; it does not necessarily mean I agree with the contents.",
                "",
                "Employee: _______________________  Date: __________",
                "Manager:  _______________________  Date: __________",
                "HR:       _______________________  Date: __________",
            ]),
        ],
    },

    "termination-checklist": {
        "title": "Termination Checklist",
        "sections": [
            ("Employee Information", [
                "Name: ____________________________",
                "Position: ____________________________",
                "Manager: ____________________________",
                "Last day of work: ____________________________",
                "Type: [Voluntary resignation / Involuntary termination / Layoff / Retirement]",
                "Reason: ____________________________",
            ]),
            ("Pre-Termination (HR + Manager)", [
                "- Review documentation supporting the termination decision",
                "- Consult HR/legal on protected-class considerations and timing",
                "- Confirm final paycheck timing per state law (CA: same day; MA: same day; varies elsewhere)",
                "- Calculate accrued PTO payout per state law (required in CA, CO, MA, IL, NE, ND, others)",
                "- Prepare separation agreement / severance offer if applicable (use OWBPA-compliant language for employees 40+)",
                "- Schedule termination meeting (private location, witness if needed)",
                "- Disable system access at end of meeting (not before)",
            ]),
            ("Day of Termination", [
                "- Hold meeting in private; have HR or witness present",
                "- Deliver final paycheck (or confirm timing) and accrued PTO payout",
                "- Provide written termination notice if state requires (e.g., CA Notice to Employee as to Change in Relationship)",
                "- Provide COBRA election notice (within 14 days) and required state mini-COBRA notices",
                "- Deliver final benefits summary and 401(k) rollover information",
                "- Distribute unemployment insurance information (state-specific brochure)",
                "- Collect company property: laptop, phone, badges, keys, credit cards, vehicle, uniforms",
                "- Disable system access: SSO, email, VPN, MDM, slack, ticketing, code repos",
                "- Forward email and calendar to manager for [X days]",
                "- Discuss return of company files / IP per separation agreement",
            ]),
            ("Post-Termination (Within 30 Days)", [
                "- Process final paycheck and PTO payout per state-mandated timeline",
                "- File state-required separation notices (e.g., NY Form IA 12.3)",
                "- Update HRIS, payroll, benefits, and org chart",
                "- Reassign work, accounts, customers, and projects",
                "- Notify clients/vendors as needed",
                "- Conduct exit interview (if employee agrees)",
                "- Archive personnel file per state retention requirements",
                "- Respond to unemployment claims promptly",
            ]),
            ("Documentation", [
                "Date of last paycheck issued: __________  Amount: __________",
                "Date of PTO payout: __________  Amount: __________",
                "COBRA notice sent: __________",
                "Separation agreement signed: __________",
                "Property returned (Y/N): __________",
                "System access disabled: __________",
                "",
                "Completed by: _______________________  Date: __________",
            ]),
        ],
    },

    "interview-scorecard": {
        "title": "Interview Scorecard",
        "sections": [
            ("Candidate", [
                "Name: ____________________________",
                "Role: ____________________________",
                "Interview date: ____________________________",
                "Interviewer: ____________________________",
                "Interview type: [Phone screen / Technical / Behavioral / Panel / Final]",
            ]),
            ("Scoring Rubric", [
                "1 — Strong No: significant concerns, do not advance",
                "2 — No: more concerns than strengths",
                "3 — Mixed: balanced; need more data",
                "4 — Yes: more strengths than concerns",
                "5 — Strong Yes: exceptional, would lose talent if not hired",
            ]),
            ("Competencies", [
                "Rate each competency 1-5 with specific evidence (quote what the candidate said).",
                "",
                "# Required for the role",
                "1. [Competency 1, e.g., Technical proficiency]   Score: __  Evidence: __________________________________",
                "2. [Competency 2, e.g., Problem solving]          Score: __  Evidence: __________________________________",
                "3. [Competency 3, e.g., Communication]            Score: __  Evidence: __________________________________",
                "4. [Competency 4, e.g., Collaboration]            Score: __  Evidence: __________________________________",
                "5. [Competency 5, e.g., Ownership]                Score: __  Evidence: __________________________________",
            ]),
            ("Interview Questions Asked", [
                "List the questions you asked. Use behavioral questions (\"Tell me about a time…\") to elicit evidence.",
                "1. _______________________________________________________________",
                "2. _______________________________________________________________",
                "3. _______________________________________________________________",
                "4. _______________________________________________________________",
            ]),
            ("Recommendation", [
                "Overall recommendation: [Strong No / No / Mixed / Yes / Strong Yes]",
                "",
                "Top 2 strengths:",
                "1. _______________________________________________________________",
                "2. _______________________________________________________________",
                "",
                "Top 2 concerns:",
                "1. _______________________________________________________________",
                "2. _______________________________________________________________",
                "",
                "Risks if hired: _______________________________________________________________",
                "",
                "Decision: [Advance to next round / Hire / Pass / Hold]",
            ]),
            ("Compliance Note", [
                "Do NOT score on or document discussions about: age, family/childcare, marital status, "
                "pregnancy, religion, national origin, disability, arrest record (in ban-the-box "
                "jurisdictions), salary history (where prohibited). If the candidate volunteers protected "
                "information, do not document or consider it.",
            ]),
        ],
    },

    "interview-guide": {
        "title": "Interview Guide — What You Can and Can't Ask",
        "sections": [
            ("Purpose", [
                "Federal and state law prohibit pre-employment inquiries that could be used "
                "to discriminate. This guide lists the most common interview questions that "
                "create legal risk, what to ask instead, and the narrow Bona Fide Occupational "
                "Qualification (BFOQ) exceptions.",
            ]),
            ("Universal Rule", [
                "Ask only about the candidate's ability to perform the essential functions "
                "of the job. If a question doesn't directly relate to job performance, don't ask it.",
            ]),
            ("Age", [
                "AVOID: \"How old are you?\" / \"When did you graduate from high school?\" / \"When do you plan to retire?\"",
                "ASK INSTEAD: \"Are you legally authorized to work in the United States?\" (only at the I-9 stage). "
                "If a minimum age is legally required (e.g., serving alcohol — typically 18 or 21 by state), "
                "you may confirm the candidate meets that legal minimum.",
            ]),
            ("Childcare, Family Status, Pregnancy", [
                "AVOID: \"Do you have children?\" / \"Are you planning to have a family?\" / \"Who watches your kids?\" / \"Are you married?\"",
                "ASK INSTEAD: \"Are you available to work the schedule we've described?\" / \"Are you able to travel as required by this role?\" "
                "Ask the same questions of all candidates regardless of perceived family status.",
            ]),
            ("Transportation", [
                "AVOID: \"Do you have a car?\" / \"How will you get to work?\"",
                "ASK INSTEAD: If driving is an essential function (e.g., delivery driver, field service), "
                "you MAY ask: \"This role requires a valid driver's license and personal vehicle. Do you have both?\" "
                "If driving is NOT a job requirement, do not ask. Public transit is a protected commuting choice.",
            ]),
            ("Citizenship and National Origin", [
                "AVOID: \"What country are you from?\" / \"Are you a U.S. citizen?\" / \"What's your native language?\"",
                "ASK INSTEAD: \"Are you legally authorized to work in the United States?\" / "
                "\"Will you now or in the future require sponsorship for employment visa status?\" "
                "Verify work authorization only via Form I-9 after offer is extended.",
            ]),
            ("Disability and Health", [
                "AVOID: \"Do you have any disabilities?\" / \"Have you ever filed a workers' comp claim?\" / \"How many sick days did you take last year?\"",
                "ASK INSTEAD: \"Can you perform the essential functions of this job, with or without reasonable accommodation?\" "
                "Pre-offer medical exams are prohibited under the ADA. Post-offer medical exams are allowed only if required of all candidates in the same job category.",
            ]),
            ("Religion", [
                "AVOID: \"What religion are you?\" / \"Do you go to church on Sundays?\" / \"Will you need any holidays off?\"",
                "ASK INSTEAD: \"Are you available to work [specific schedule]?\" Religious accommodation can be discussed after hire.",
            ]),
            ("Arrest and Conviction Record", [
                "AVOID: Asking about arrest records — only convictions are relevant, and 37+ states/cities prohibit asking on the application (\"Ban the Box\").",
                "ASK INSTEAD: Delay any criminal-history inquiry until after a conditional offer, then conduct a job-related individualized assessment per EEOC 2012 guidance.",
            ]),
            ("Salary History", [
                "AVOID in: AL, CA, CO, CT, DE, HI, IL, KY (some), MA, ME, MD, NJ, NY, NC, OR, PR, UT (gov't), VT, WA, WV, DC, and many cities.",
                "ASK INSTEAD: \"What are your salary expectations for this role?\" Provide your pay range proactively where pay-transparency laws apply.",
            ]),
            ("Sex, Sexual Orientation, Gender Identity", [
                "AVOID any inquiry about sex, gender identity, sexual orientation, or pronouns at the recruiting stage.",
                "ASK INSTEAD: Use the candidate's name throughout the process. Ask about pronouns only if and when the candidate volunteers them or after hire.",
            ]),
            ("Bona Fide Occupational Qualification (BFOQ)", [
                "A BFOQ permits limited inquiry where the protected characteristic is reasonably necessary "
                "to the job. Examples:",
                "- Age: minimum age for selling alcohol or operating a commercial vehicle",
                "- Sex: actor playing a specific gender role; locker-room attendant for the same sex",
                "- Religion: clergy or religious-school teacher of that faith",
                "BFOQ does NOT exist for race or color under Title VII. Construe BFOQs narrowly. When in doubt, don't ask.",
            ]),
            ("If a Candidate Volunteers Protected Information", [
                "Politely redirect: \"Thank you for sharing that. Let's get back to your experience with [topic relevant to the role].\" "
                "Do not document the volunteered information. Do not consider it in your decision.",
            ]),
        ],
    },

    "pto-policy": {
        "title": "Paid Time Off (PTO) Policy",
        "sections": [
            ("Purpose", [
                "[Company Name] provides paid time off to support employee well-being, work-life balance, "
                "and personal needs. This policy outlines accrual, use, and payout of PTO.",
            ]),
            ("Eligibility", [
                "All [full-time / part-time at X+ hours per week] regular employees are eligible for PTO. "
                "Temporary, seasonal, and contract workers are not eligible unless required by state law.",
            ]),
            ("Accrual", [
                "PTO accrues per pay period based on tenure:",
                "- 0–2 years: [X] hours per pay period ([Y] days/year)",
                "- 2–5 years: [X] hours per pay period ([Y] days/year)",
                "- 5+ years: [X] hours per pay period ([Y] days/year)",
                "",
                "Accrual begins on the first day of employment. PTO is available for use as it accrues, "
                "subject to manager approval and operational needs.",
            ]),
            ("Cap and Carryover", [
                "Employees may carry over unused PTO up to a maximum balance of [X] hours. "
                "Once the cap is reached, additional accrual stops until the balance falls below the cap. "
                "(Note: \"use it or lose it\" forfeiture is illegal in CA, CO, MA, IL, NE, ND, MT and "
                "select other jurisdictions. Caps must be reasonable. Adjust this section accordingly for "
                "states where the company operates.)",
            ]),
            ("Requesting PTO", [
                "Requests should be submitted at least [2 weeks] in advance, except in emergencies. "
                "Approval is at the manager's discretion based on operational needs. "
                "Pre-approved PTO will not be unreasonably revoked.",
            ]),
            ("Sick Leave", [
                "Employees in jurisdictions with mandated paid sick leave (CA, CO, MA, MD, NJ, NM, NV, NY, OR, RI, VT, WA, "
                "and 30+ cities) accrue paid sick leave separately from PTO at the rate required by local law. "
                "Sick leave can be used for the employee's own illness, illness of a family member, "
                "preventive care, or related qualifying purposes.",
            ]),
            ("Payout at Separation", [
                "In states where accrued PTO is treated as wages (CA, CO, MA, IL, NE, ND, MT, others), "
                "all accrued and unused PTO will be paid out at the employee's regular rate of pay at "
                "separation, with the final paycheck. In all other states, payout is per company "
                "discretion as outlined in this policy.",
            ]),
            ("PTO Abuse", [
                "Employees who misuse PTO (e.g., consistent unscheduled absences, falsifying reasons) "
                "may be subject to disciplinary action, up to and including termination. "
                "Sick leave used per state law cannot be the basis for adverse action.",
            ]),
        ],
    },

    "workplace-investigation-report": {
        "title": "Workplace Investigation Report",
        "sections": [
            ("Case Information", [
                "Investigation ID: ____________________________",
                "Date of complaint: ____________________________",
                "Date of report: ____________________________",
                "Investigator: ____________________________",
                "Type of complaint: [Harassment / Discrimination / Retaliation / Policy violation / Safety / Other]",
                "Confidentiality level: ____________________________",
            ]),
            ("Parties", [
                "Complainant: ____________________________",
                "Respondent(s): ____________________________",
                "Witnesses interviewed: ____________________________",
            ]),
            ("Allegations", [
                "Specific allegations as reported:",
                "",
                "1. _______________________________________________________________",
                "2. _______________________________________________________________",
                "3. _______________________________________________________________",
                "",
                "Policies / laws potentially implicated: _______________________________________________________________",
            ]),
            ("Investigation Steps Taken", [
                "Document each step taken, in chronological order:",
                "- [Date] Initial intake meeting with complainant",
                "- [Date] Document review (emails, Slack, security footage, HRIS, etc.)",
                "- [Date] Witness interviews (list)",
                "- [Date] Interview with respondent",
                "- [Date] Follow-up interviews",
                "- [Date] Closing meeting with complainant",
            ]),
            ("Evidence Reviewed", [
                "Documents: _______________________________________________________________",
                "Electronic communications: _______________________________________________________________",
                "Physical evidence / video: _______________________________________________________________",
                "Witness statements: _______________________________________________________________",
            ]),
            ("Findings of Fact", [
                "For each allegation, state whether substantiated, partially substantiated, "
                "unsubstantiated, or inconclusive, and the evidence supporting that finding. "
                "Use the preponderance-of-evidence standard.",
                "",
                "Allegation 1 finding: _______________________________________________________________",
                "Supporting evidence: _______________________________________________________________",
                "",
                "Allegation 2 finding: _______________________________________________________________",
                "Supporting evidence: _______________________________________________________________",
            ]),
            ("Conclusion and Recommended Action", [
                "Based on the findings above, the investigator recommends:",
                "- [No action]",
                "- [Coaching / training]",
                "- [Written warning]",
                "- [Performance improvement plan]",
                "- [Suspension]",
                "- [Termination]",
                "- [Other: ____________]",
                "",
                "Anti-retaliation reminder: All parties should be reminded that retaliation against the complainant or witnesses is prohibited.",
            ]),
            ("Sign-off", [
                "Investigator: _______________________  Date: __________",
                "HR Reviewer: _______________________  Date: __________",
                "Legal Reviewer (if applicable): _______________________  Date: __________",
            ]),
        ],
    },

    "performance-review": {
        "title": "Performance Review",
        "sections": [
            ("Review Information", [
                "Employee: ____________________________",
                "Position: ____________________________",
                "Manager: ____________________________",
                "Review period: __________ to __________",
                "Type: [Annual / Mid-year / Quarterly / 90-day / Project-based]",
                "Date: ____________________________",
            ]),
            ("Self-Review (Employee Completes)", [
                "Top accomplishments this review period:",
                "1. _______________________________________________________________",
                "2. _______________________________________________________________",
                "3. _______________________________________________________________",
                "",
                "Areas where I want to grow: _______________________________________________________________",
                "",
                "What support do I need from my manager / company? _______________________________________________________________",
            ]),
            ("Goals From Last Review", [
                "Goal 1: ____________________________  Status: [Met / Partially met / Not met]   Notes: __________",
                "Goal 2: ____________________________  Status: [Met / Partially met / Not met]   Notes: __________",
                "Goal 3: ____________________________  Status: [Met / Partially met / Not met]   Notes: __________",
            ]),
            ("Competency Ratings (Manager Completes)", [
                "Rating scale: 1 — Below expectations | 2 — Approaching | 3 — Meets | 4 — Exceeds | 5 — Outstanding",
                "",
                "Job knowledge: __  Examples: __________",
                "Quality of work: __  Examples: __________",
                "Productivity: __  Examples: __________",
                "Collaboration: __  Examples: __________",
                "Communication: __  Examples: __________",
                "Initiative / ownership: __  Examples: __________",
                "Leadership (if applicable): __  Examples: __________",
                "",
                "Overall rating: __",
            ]),
            ("Strengths and Development Areas", [
                "Top 3 strengths:",
                "1. _______________________________________________________________",
                "2. _______________________________________________________________",
                "3. _______________________________________________________________",
                "",
                "Top 2 development areas:",
                "1. _______________________________________________________________",
                "2. _______________________________________________________________",
            ]),
            ("Goals for Next Review Period", [
                "Goal 1 (SMART): _______________________________________________________________",
                "Goal 2 (SMART): _______________________________________________________________",
                "Goal 3 (SMART): _______________________________________________________________",
                "",
                "Development plan / training: _______________________________________________________________",
            ]),
            ("Compensation Discussion (if applicable)", [
                "Recommendation: [No change / Merit increase / Promotion / Bonus]",
                "Details: _______________________________________________________________",
            ]),
            ("Sign-off", [
                "Employee: _______________________  Date: __________",
                "Manager:  _______________________  Date: __________",
                "Skip-level / HR: _______________________  Date: __________",
                "",
                "Employee comments (optional): _______________________________________________________________",
            ]),
        ],
    },

    "disciplinary-action": {
        "title": "Disciplinary Action Form",
        "sections": [
            ("Employee Information", [
                "Name: ____________________________",
                "Position: ____________________________",
                "Manager: ____________________________",
                "Date of incident: ____________________________",
                "Date of this action: ____________________________",
            ]),
            ("Type of Action", [
                "[ ] Verbal warning (documented)",
                "[ ] Written warning",
                "[ ] Final written warning",
                "[ ] Suspension (paid/unpaid; specify days): __________",
                "[ ] Demotion / pay reduction",
                "[ ] Termination",
            ]),
            ("Policy or Standard Violated", [
                "Cite specific company policy, employee handbook section, or job standard:",
                "_____________________________________________________________________",
            ]),
            ("Description of Incident / Conduct", [
                "Be specific. Include: who, what, when, where, and any witnesses or evidence. "
                "Stick to facts, not opinions.",
                "_____________________________________________________________________",
                "_____________________________________________________________________",
                "_____________________________________________________________________",
            ]),
            ("Prior Coaching / Discipline", [
                "Previous conversations, warnings, or corrective actions related to this issue:",
                "1. [Date]: _______________________________________________________________",
                "2. [Date]: _______________________________________________________________",
                "",
                "(If none, note that and proceed accordingly.)",
            ]),
            ("Expected Improvement", [
                "Specific behavior the employee must demonstrate going forward:",
                "_____________________________________________________________________",
                "_____________________________________________________________________",
                "",
                "Timeframe for improvement: _______________________________________________________________",
            ]),
            ("Consequences of Continued Issue", [
                "If the issue continues or recurs, the next step will be: _______________________________________________________________",
                "(May include further discipline up to and including termination.)",
            ]),
            ("Sign-off", [
                "I have read and understand this disciplinary action. My signature acknowledges receipt; "
                "it does not necessarily mean I agree with the contents.",
                "",
                "Employee: _______________________  Date: __________",
                "Manager:  _______________________  Date: __________",
                "HR / Witness: _______________________  Date: __________",
                "",
                "Employee comments (optional): _______________________________________________________________",
            ]),
        ],
    },

    "remote-work-agreement": {
        "title": "Remote Work Agreement",
        "sections": [
            ("Parties and Position", [
                "Employee: ____________________________",
                "Position: ____________________________",
                "Manager: ____________________________",
                "Effective date: ____________________________",
                "Remote work location (city, state): ____________________________",
                "Type: [Fully remote / Hybrid (X days remote per week)]",
            ]),
            ("Work Schedule and Availability", [
                "Standard work hours: __________ to __________ in [time zone].",
                "Core collaboration hours: __________ to __________ in [time zone] (employee must be available).",
                "Time tracking: [Required for non-exempt; not required for exempt]",
                "Manager approval required for changes to schedule.",
            ]),
            ("Equipment", [
                "[Company provides / Employee uses personal] equipment:",
                "- Laptop: ____________________________",
                "- Monitor(s): ____________________________",
                "- Keyboard / mouse: ____________________________",
                "- Headset / webcam: ____________________________",
                "Employee is responsible for safekeeping. Company-provided equipment must be returned at separation.",
            ]),
            ("Workspace and Internet", [
                "Employee is responsible for maintaining a safe, ergonomic workspace and reliable internet "
                "connection sufficient for video calls and the work performed.",
                "Reimbursement for home-office expenses (where required by state law — CA Labor Code §2802, IL, "
                "MA, MN, MT, NH, ND, PA, SD, IA, DC): [eligible categories and reimbursement process].",
            ]),
            ("Tax and Workers' Compensation", [
                "Working from a state other than the company's primary location may create tax nexus and "
                "workers' compensation obligations. Employee must notify HR before changing remote work "
                "location to a different state. Working from outside the U.S. is not permitted without "
                "prior written approval.",
            ]),
            ("Data Security and Confidentiality", [
                "Employee will:",
                "- Use company-approved security tools (VPN, MDM, password manager, antivirus)",
                "- Not store confidential data on personal devices",
                "- Lock devices when stepping away",
                "- Not allow non-employees to use company devices",
                "- Report any suspected breach within 24 hours",
            ]),
            ("Performance Expectations", [
                "Remote work does not change the employee's performance expectations or the company's "
                "right to require attendance at in-person meetings, training, or events with reasonable notice.",
            ]),
            ("Modification or Revocation", [
                "This agreement is not a guarantee of continued remote work. The company may modify or "
                "revoke remote work arrangements at any time, with or without cause, with reasonable notice. "
                "Employment remains at-will (except in Montana).",
            ]),
            ("Sign-off", [
                "Employee: _______________________  Date: __________",
                "Manager:  _______________________  Date: __________",
                "HR:       _______________________  Date: __________",
            ]),
        ],
    },

    "expense-reimbursement": {
        "title": "Expense Reimbursement Form and Policy",
        "sections": [
            ("Policy Overview", [
                "[Company Name] reimburses employees for reasonable, necessary, and pre-approved business "
                "expenses incurred in the performance of their job duties. This policy is structured as an "
                "IRS-compliant Accountable Plan under Treas. Reg. §1.62-2 — reimbursements are not taxable "
                "as wages provided the requirements below are met.",
            ]),
            ("Eligible Expenses", [
                "- Business travel: airfare (economy unless flights >6 hours), hotel, ground transportation",
                "- Meals during business travel (per diem or actual cost — see policy)",
                "- Client meals and entertainment (50% deductible to company; document attendees + business purpose)",
                "- Office supplies for remote workers (where state law requires reimbursement)",
                "- Mileage at the IRS standard rate for business use of personal vehicle",
                "- Conference and professional-development fees with manager pre-approval",
                "- Professional licenses, dues, and certifications required for the role",
            ]),
            ("Ineligible Expenses", [
                "- Personal items, gifts to colleagues, alcohol (unless client entertainment with approval)",
                "- Traffic violations, parking tickets, fines",
                "- Spousal travel (unless explicitly approved)",
                "- Commuting between home and primary work location",
                "- First-class or business-class travel (unless approved exception)",
                "- Entertainment with no business purpose",
            ]),
            ("Submission Requirements", [
                "1. Submit within 30 days of the expense (60 days max).",
                "2. Itemized receipts required for all expenses over $25.",
                "3. Document business purpose, attendees (for meals/entertainment), and date.",
                "4. Use company expense system or attach receipts to this form.",
                "5. Approval by direct manager required before reimbursement.",
            ]),
            ("Reimbursement Form", [
                "Employee: ____________________________   Department: ____________________________",
                "Period: __________ to __________",
                "",
                "Date  |  Description  |  Category  |  Business Purpose  |  Amount",
                "_______________________________________________________________________________",
                "_______________________________________________________________________________",
                "_______________________________________________________________________________",
                "_______________________________________________________________________________",
                "_______________________________________________________________________________",
                "",
                "Total: $__________",
                "",
                "Mileage (separate): __________ miles × IRS rate $__________ = $__________",
            ]),
            ("Sign-off", [
                "I certify the expenses above were incurred for legitimate business purposes and that "
                "receipts are attached for all expenses over $25.",
                "",
                "Employee: _______________________  Date: __________",
                "Manager approval: _______________________  Date: __________",
                "Finance / AP: _______________________  Date: __________",
            ]),
        ],
    },

    "severance-agreement": {
        "title": "Severance Agreement and Release",
        "sections": [
            ("IMPORTANT NOTICE", [
                "This template includes language designed for employees age 40 and older to comply with "
                "the Older Workers Benefit Protection Act (OWBPA). For group terminations, additional "
                "disclosures are required under OWBPA. Have qualified employment counsel review before use.",
            ]),
            ("Parties", [
                "This Severance Agreement and Release (\"Agreement\") is entered into between [Company Name] "
                "(\"Company\") and [Employee Name] (\"Employee\").",
            ]),
            ("Separation", [
                "Employee's employment with Company will end effective [Separation Date]. "
                "On the Separation Date, Company will pay Employee all earned wages through that date "
                "and any accrued, unused PTO required by law or policy.",
            ]),
            ("Severance Payment", [
                "In exchange for Employee's promises in this Agreement, including the release in Section 5, "
                "Company agrees to pay Employee severance in the amount of $[Amount], less applicable "
                "withholdings, payable [in a lump sum / over [X] pay periods] beginning after this Agreement "
                "becomes effective.",
                "",
                "Continued benefits: [Description, e.g., COBRA premium subsidy for X months]",
            ]),
            ("General Release of Claims", [
                "In exchange for the consideration above, Employee releases and forever discharges Company, "
                "its parents, subsidiaries, affiliates, officers, directors, employees, agents, and successors "
                "from any and all claims, demands, and causes of action of any kind that Employee has or "
                "may have against Company arising out of Employee's employment or separation, including but "
                "not limited to claims under Title VII, the ADA, the ADEA, the FMLA, the FLSA, ERISA, the "
                "WARN Act, [applicable state law], and any other federal, state, or local law.",
                "",
                "This release does NOT waive: (a) rights that cannot be waived as a matter of law; "
                "(b) the right to file a charge with the EEOC, NLRB, or similar agency (though Employee "
                "waives the right to recover monetary damages from such a charge); (c) workers' compensation "
                "or unemployment benefits; (d) vested retirement benefits.",
            ]),
            ("ADEA / OWBPA Notice (Employees Age 40+)", [
                "Employee acknowledges:",
                "- Employee is advised in writing to consult with an attorney before signing.",
                "- Employee has [21 days] (45 days for group RIF) from receipt to consider this Agreement.",
                "- Employee may revoke this Agreement within 7 days after signing by written notice to "
                "  [HR contact + email]. The Agreement does not take effect until the 7-day revocation period expires.",
                "- Employee is not waiving any ADEA claims that arise after signing.",
            ]),
            ("Confidentiality", [
                "Employee will keep confidential the existence, terms, and amount of this Agreement, except "
                "from Employee's spouse, attorney, accountant, or as required by law. Nothing in this Agreement "
                "prohibits Employee from communicating with the EEOC, NLRB, SEC, or other government agency, "
                "or from making disclosures protected by whistleblower laws.",
            ]),
            ("Return of Property", [
                "Employee has returned (or will return by [Date]) all Company property, including but not "
                "limited to laptop, phone, badges, keys, credit cards, and confidential information.",
            ]),
            ("Non-Disparagement", [
                "Employee and Company agree not to make any disparaging statements about each other. "
                "Truthful statements made to government agencies, regulators, or in legal proceedings are not disparagement.",
            ]),
            ("References", [
                "Inquiries about Employee's employment will be directed to HR, which will confirm dates of "
                "employment, position(s) held, and final salary, consistent with company policy.",
            ]),
            ("Entire Agreement", [
                "This Agreement, together with [Confidentiality and Invention Assignment Agreement / "
                "other surviving agreements], constitutes the entire agreement between the parties and "
                "supersedes all prior negotiations or agreements regarding its subject matter.",
            ]),
            ("Governing Law", [
                "This Agreement is governed by the laws of the State of [State], without regard to "
                "conflict-of-law principles.",
            ]),
            ("Sign-off", [
                "READ CAREFULLY BEFORE SIGNING. By signing below, Employee acknowledges that Employee has "
                "read this Agreement, understands its terms, has had the opportunity to consult counsel, "
                "and signs it voluntarily.",
                "",
                "Employee: _______________________  Date: __________",
                "Company:  _______________________  Date: __________",
                "          [Name, Title]",
            ]),
        ],
    },
}


async def main():
    load_settings()
    storage = get_storage()
    if not storage.s3_client or not storage.bucket:
        print("ERROR: S3 not configured.")
        sys.exit(1)

    print(f"Generating {len(TEMPLATES)} templates and uploading to S3 bucket: {storage.bucket}")
    print()

    results: dict[str, str] = {}
    for slug, spec in TEMPLATES.items():
        filename = f"{slug}.docx"
        print(f"  - Generating {filename}…", end=" ", flush=True)
        data = _build_doc(spec["title"], spec["sections"])
        try:
            url = await storage.upload_file(
                file_bytes=data,
                filename=filename,
                prefix="resources/templates",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            results[slug] = url
            print(f"OK ({len(data):,} bytes)")
        except Exception as e:
            print(f"FAIL: {e}")

    print()
    print("=" * 70)
    print("URL MAPPING (paste into ASSETS in resources.py):")
    print("=" * 70)
    for slug, url in results.items():
        title = TEMPLATES[slug]["title"]
        print(f'    "{slug}": {{"path": "{url}", "name": "{title}", "available": True}},')


if __name__ == "__main__":
    asyncio.run(main())
